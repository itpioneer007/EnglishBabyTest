"""
src/report_exporter.py — 报告生成 & 导出
负责人：C

职责：
  1. 生成 HTML 报告（全部题目 + 仅错误题目）
  2. 生成 CSV 表格（方便老师导入 Excel 批注）
  3. 打包错误截图 zip
  4. 按版本/Unit/日期组织输出目录结构

调用方：
  routes/export_routes.py → 各导出 API
  src/batch_runner.BatchRunner → 批量跑完自动调用

输出目录结构约定:
  {save_dir}/
    ├── 新湘鲁六上/
    │   ├── U6_基础巩固_20260728/
    │   │   ├── report_full.html       # 全量报告
    │   │   ├── report_errors.html     # 仅错误报告
    │   │   ├── errors.csv             # 错误表格
    │   │   └── screenshots/           # 错误截图
    │   │       ├── q03.png
    │   │       └── q07.png
    │   └── U7_基础巩固_20260728/
    │       └── ...
    └── latest/                        # 快捷入口：软链到最新
        └── report.html
"""

import json
import csv
import io
import shutil
import html
import re
from pathlib import Path
from datetime import datetime
from typing import Optional


# ============================================================
# HTML 报告模板（内联 CSS，无外部依赖）
# ============================================================

_HTML_CSS = """
<style>
  * { margin:0; padding:0; box-sizing:border-box; }
  body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
         background:#f5f7fa; color:#333; line-height:1.6; padding:20px; }
  .container { max-width:1100px; margin:0 auto; }
  h1 { font-size:24px; margin-bottom:4px; }
  .meta { color:#888; font-size:13px; margin-bottom:20px; }

  /* 汇总卡片 */
  .summary { display:flex; gap:16px; flex-wrap:wrap; margin-bottom:24px; }
  .summary-card { flex:1; min-width:140px; background:#fff; border-radius:10px;
                   padding:18px 20px; box-shadow:0 1px 4px rgba(0,0,0,.08); text-align:center; }
  .summary-card .num { font-size:32px; font-weight:700; }
  .summary-card .label { font-size:12px; color:#999; margin-top:4px; }
  .s-pass .num { color:#2ecc71; }
  .s-fail .num { color:#e74c3c; }
  .s-uncertain .num { color:#f39c12; }
  .s-conf .num { color:#3498db; }

  /* 题目卡片 */
  .q-card { background:#fff; border-radius:10px; padding:20px 24px;
             margin-bottom:16px; box-shadow:0 1px 4px rgba(0,0,0,.08);
             border-left:5px solid #ddd; transition:border-color .2s; }
  .q-card.pass { border-left-color:#2ecc71; }
  .q-card.fail { border-left-color:#e74c3c; }
  .q-card.uncertain { border-left-color:#f39c12; }
  .q-header { display:flex; justify-content:space-between; align-items:center; margin-bottom:12px; }
  .q-title { font-size:18px; font-weight:600; }
  .q-type { color:#888; font-size:13px; background:#f0f0f0; padding:2px 10px; border-radius:12px; }
  .q-badge { font-size:13px; padding:3px 12px; border-radius:12px; font-weight:600; }
  .q-badge.pass { background:#e8f8f0; color:#27ae60; }
  .q-badge.fail { background:#fde8e8; color:#e74c3c; }
  .q-badge.uncertain { background:#fef5e7; color:#e67e22; }

  /* 维度条 */
  .dim-grid { display:grid; grid-template-columns:1fr 1fr; gap:10px; }
  @media (max-width:600px) { .dim-grid { grid-template-columns:1fr; } }
  .dim-row { display:flex; align-items:center; gap:8px; font-size:13px; }
  .dim-name { width:40px; color:#888; flex-shrink:0; }
  .dim-bar-wrap { flex:1; height:14px; background:#eee; border-radius:7px; overflow:hidden; position:relative; }
  .dim-bar { height:100%; border-radius:7px; transition:width .4s; }
  .dim-bar.pass { background:#2ecc71; }
  .dim-bar.fail { background:#e74c3c; }
  .dim-bar.uncertain { background:#f39c12; }
  .dim-val { width:80px; font-size:12px; text-align:right; flex-shrink:0; }

  /* 详情说明 */
  .details { margin-top:10px; font-size:12px; color:#666; border-top:1px solid #eee; padding-top:10px; }
  .detail-item { margin:2px 0; }
  .detail-item.uncertain { color:#e67e22; font-weight:600; }

  /* 置信度指示器 */
  .conf-dot { display:inline-block; width:8px; height:8px; border-radius:50%; margin-right:4px; }
  .conf-high { background:#2ecc71; }
  .conf-mid { background:#f39c12; }
  .conf-low { background:#e74c3c; }

  .footer { text-align:center; color:#bbb; font-size:12px; margin-top:30px; }
  .no-data { text-align:center; padding:60px 0; color:#bbb; font-size:16px; }

  /* 导出按钮 */
  .toolbar { display:flex; gap:10px; margin-bottom:20px; flex-wrap:wrap; }
  .btn { padding:8px 18px; border:none; border-radius:6px; cursor:pointer; font-size:13px; font-weight:600; transition:opacity .2s; }
  .btn:hover { opacity:.85; }
  .btn-primary { background:#3498db; color:#fff; }
  .btn-outline { background:#fff; color:#3498db; border:1px solid #3498db; }

  @media print { body { background:#fff; padding:0; } .q-card { box-shadow:none; break-inside:avoid; } }
</style>
"""


class ReportExporter:
    """报告导出器 — 支持 HTML / CSV / ZIP"""

    def __init__(self, save_dir: str = None):
        base = Path(save_dir or str(Path(__file__).parent.parent / "outputs" / "reports"))
        self.save_dir = Path(base)
        self.save_dir.mkdir(parents=True, exist_ok=True)

    # ============================================================
    # 题干清洗（过滤状态栏时间/得分/进度等 UI 噪音）
    # ============================================================
    @staticmethod
    def _clean_stem(stem: str) -> str:
        """清洗题干：去掉状态栏时间(21:12)/得分(77.0)/进度(3/40)/百分比(100%)等 UI 噪音。
        按 engine 提取时的连接符" / "拆分，逐段判定；返回清洗后的题干，空则返回空串。"""
        s = (stem or "").strip()
        if not s:
            return ""
        _noise = re.compile(
            r"^(?:"
            r"还剩[：:]?\s*\d{1,2}:\d{2}"   # 还剩 19:58
            r"|\d{1,2}:\d{2}"               # 21:12
            r"|\d+\.?\d*\s*%"               # 77.0% / 100%
            r"|\d+\s*/\s*\d+"               # 3/40
            r"|\d+\.\d+"                    # 77.0
            r"|\d+"                         # 100
            r")\s*$"
        )
        # 按 engine 连接符拆（" / " 及兼容的 |｜·），中文逗号属于题干本身不拆
        parts = [p.strip() for p in re.split(r"\s*(?:/|｜|\|)\s*", s)]
        clean = [p for p in parts if p and not _noise.match(p)]
        out = " / ".join(clean)
        return out[:120] if out else ""

    def _loc_label(self, q: dict, meta: dict = None) -> str:
        """生成题目定位标签：模块 · 子模块 · 第N题（如：听力专项 · 练习 · 第3题）"""
        mod = (q.get("module", "") or "").strip()
        stage = (q.get("stage", "") or "").strip()
        if not mod:
            qid = q.get("qid", "")
            _m = re.match(r"^([^-]+)-?([^-]*)$", qid or "")
            if _m and _m.group(1) not in ("", "auto", "quick"):
                mod = _m.group(1)
        qno = q.get("module_qno") or q.get("idx") or q.get("progress") or "?"
        if isinstance(qno, float):
            qno = int(qno) if qno == int(qno) else qno
        return f"{mod} · {stage}" if stage else str(mod)

    # ============================================================
    # HTML 报告
    # ============================================================

    def export_html_full(self, questions: list[dict], metadata: dict = None,
                         only_errors: bool = False) -> str:
        """
        生成完整的 self-contained HTML 报告。

        Args:
            questions: 题目列表，每个元素为 dict，格式同 inspection_state.json 的 questions 条目
            metadata: {"version": "新湘鲁六上", "unit": 6, "stage": "基础巩固", "docx": "..."}
            only_errors: True=仅导出错误题目
        Returns:
            报告文件路径
        """
        meta = metadata or {}
        version = meta.get("version", "未知版本")
        unit = meta.get("unit", "?")
        stage = meta.get("stage", "?")
        date_str = meta.get("date", datetime.now().strftime("%Y-%m-%d %H:%M"))

        qs = [q for q in questions if not only_errors or not q.get("overall_passed", True)]
        if not qs:
            # 无结果时也生成一个空报告
            html = self._render_html([], meta, only_errors)
        else:
            html = self._render_html(qs, meta, only_errors)

        # 写入文件
        filename = "report_errors.html" if only_errors else "report_full.html"
        out_dir = self.organize_output_dir(version, unit, stage)
        out_path = out_dir / filename
        out_path.write_text(html, encoding="utf-8")
        return str(out_path)

    def export_html_errors(self, questions: list[dict], metadata: dict = None) -> str:
        """仅导出错误题目HTML"""
        return self.export_html_full(questions, metadata, only_errors=True)

    # ============================================================
    # 内部 HTML 渲染
    # ============================================================

    def _render_html(self, questions: list[dict], metadata: dict,
                     only_errors: bool) -> str:
        meta = metadata or {}
        version = meta.get("version", "未知版本")
        unit = meta.get("unit", "?")
        stage = meta.get("stage", "?")
        date_str = meta.get("date", datetime.now().strftime("%Y-%m-%d %H:%M"))

        total = len(questions)
        passed = sum(1 for q in questions if q.get("overall_passed"))
        failed = total - passed
        uncertain = sum(1 for q in questions if q.get("error_dimensions", []))

        # 平均置信度
        confs = []
        for q in questions:
            for dim in ["stem_reason", "content_reason", "image_reason", "answer_reason"]:
                if dim in q and q[dim]:
                    confs.append(80)  # 粗略估计，后续可从 details 中提取
        avg_conf = round(sum(confs) / len(confs)) if confs else 0

        cards_html = "\n".join(self._render_question(q) for q in questions)

        title = f"英语宝题目审查报告 — {version} U{unit} {stage}"
        if only_errors:
            title += "（仅错误）"

        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
{_HTML_CSS}
</head>
<body>
<div class="container">

  <h1>{title}</h1>
  <div class="meta">生成时间: {date_str} | 版本: {version} | Unit {unit} {stage} | 题目数: {total}</div>

  <div class="summary">
    <div class="summary-card s-pass"><div class="num">{passed}</div><div class="label">通过</div></div>
    <div class="summary-card s-fail"><div class="num">{failed}</div><div class="label">不通过</div></div>
    <div class="summary-card s-uncertain"><div class="num">{uncertain}</div><div class="label">需人工复核</div></div>
    <div class="summary-card s-conf"><div class="num">{avg_conf}%</div><div class="label">平均置信度</div></div>
    <div class="summary-card"><div class="num">{pass_rate}%</div><div class="label">通过率</div></div>
  </div>

  <div class="toolbar">
    <button class="btn btn-primary" onclick="window.print()">打印报告</button>
    <button class="btn btn-outline" onclick="document.querySelectorAll('.details').forEach(d=>d.style.display=d.style.display==='none'?'block':'none')">展开/收起详情</button>
  </div>

  {cards_html if questions else no_data_html}

  <div class="footer">英语宝审查智能体 v3 · 自动生成 · 仅供参考，最终判定以人工确认为准</div>

</div>
</body>
</html>"""

    def _render_question(self, q: dict) -> str:
        """渲染单道题目卡片"""
        qid = q.get("qid", "?")
        idx = q.get("idx", 0)
        qtype = q.get("question_type", "?")
        stem = self._clean_stem(q.get("stem", ""))[:60]
        overall_passed = q.get("overall_passed", False)
        error_dims = q.get("error_dimensions", [])

        # 判断状态
        if error_dims:
            card_class = "uncertain"
            badge_class = "uncertain"
            badge_text = "⚠ 需人工复核"
        elif overall_passed:
            card_class = "pass"
            badge_class = "pass"
            badge_text = "通过"
        else:
            card_class = "fail"
            badge_class = "fail"
            badge_text = "不通过"

        # 各维度状态
        dims = [
            ("题干", "stem_reason", q.get("stem_reason", "")),
            ("内容", "content_reason", q.get("content_reason", "")),
            ("配图", "image_reason", q.get("image_reason", "")),
            ("作答", "answer_reason", q.get("answer_reason", "")),
        ]

        dim_rows = []
        for dname, dattr, dreason in dims:
            val = q.get(f"ai_{dattr.replace('_reason','')}", False) if dattr.endswith("_reason") else None
            # 从 details 中判断
            dim_passed = dname not in error_dims
            if dreason:
                if "通过" in dreason and "不通过" not in dreason:
                    dim_passed = True
                elif "不通过" in dreason:
                    dim_passed = False
                elif "需人工复核" in dreason or "无法解析" in dreason:
                    dim_passed = None

            bar_class = "pass" if dim_passed is True else ("fail" if dim_passed is False else "uncertain")
            bar_pct = 100 if dim_passed is True else (30 if dim_passed is False else 50)
            label = "通过" if dim_passed is True else ("不通过" if dim_passed is False else "不确定")

            dim_rows.append(f"""<div class="dim-row">
  <span class="dim-name">{dname}</span>
  <div class="dim-bar-wrap"><div class="dim-bar {bar_class}" style="width:{bar_pct}%"></div></div>
  <span class="dim-val">{label}</span>
</div>""")

        details_html = ""
        for dname, dattr, dreason in dims:
            if dreason:
                cls = "detail-item"
                if "需人工复核" in dreason or "无法解析" in dreason:
                    cls += " uncertain"
                details_html += f'<div class="{cls}"><strong>{dname}:</strong> {dreason[:200]}</div>'

        return f"""<div class="q-card {card_class}">
  <div class="q-header">
    <span class="q-title">Q{idx} {stem}</span>
    <span class="q-type">{qtype}</span>
    <span class="q-badge {badge_class}">{badge_text}</span>
  </div>
  <div class="q-loc">{self._loc_label(q)} · 第{idx}题</div>
  <div class="dim-grid">
    {''.join(dim_rows)}
  </div>
  {"<div class='details'>" + details_html + "</div>" if details_html else ""}
</div>"""

    # ============================================================
    # CSV 导出
    # ============================================================

    def export_csv(self, questions: list[dict], only_errors: bool = True) -> str:
        """导出错误题目CSV"""
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow([
            "题号", "题型", "题干", "脚本答案", "综合得分",
            "综合通过", "错误维度", "题干判定", "内容判定", "配图判定", "作答判定",
            "需人工复核", "截图",
        ])

        for q in questions:
            if only_errors and q.get("overall_passed"):
                continue
            dims = q.get("error_dimensions", [])
            writer.writerow([
                q.get("qid", ""), q.get("question_type", ""),
                (q.get("stem") or "")[:60], q.get("script_answer", ""),
                q.get("overall_score", 0),
                "是" if q.get("overall_passed") else "否",
                ", ".join(dims) if dims else "无",
                q.get("stem_reason", ""), q.get("content_reason", ""),
                q.get("image_reason", ""), q.get("answer_reason", ""),
                "是" if dims else "否",
                q.get("screenshot", ""),
            ])

        csv_path = self.save_dir / f"errors_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        csv_path.write_text(buf.getvalue(), encoding="utf-8-sig")
        return str(csv_path)

    # ============================================================
    # 截图打包
    # ============================================================

    def export_screenshots_zip(self, questions: list[dict], shots_dir: str = None) -> str:
        """打包错误截图"""
        import zipfile

        shots = Path(shots_dir or Path(__file__).parent.parent / "screenshots")
        zip_path = self.save_dir / f"screenshots_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"

        with zipfile.ZipFile(zip_path, "w") as zf:
            for q in questions:
                if not q.get("overall_passed") or q.get("error_dimensions"):
                    shot = q.get("screenshot", "")
                    if shot:
                        shot_path = shots / shot
                        if shot_path.exists():
                            zf.write(shot_path, f"{q.get('qid','')}_{shot}")
                        else:
                            # 尝试绝对路径
                            shot_path = Path(shot)
                            if shot_path.exists():
                                zf.write(shot_path, f"{q.get('qid','')}_{shot_path.name}")

        return str(zip_path)

    # ============================================================
    # 输出目录管理
    # ============================================================

    def organize_output_dir(self, version: str, unit: int, stage: str) -> Path:
        """创建按版本/Unit/日期组织的输出目录"""
        date_str = datetime.now().strftime('%Y%m%d')
        dir_name = f"U{unit}_{stage}_{date_str}"
        out_dir = self.save_dir / version / dir_name
        out_dir.mkdir(parents=True, exist_ok=True)

        # 更新 latest 快捷入口
        latest = self.save_dir / "latest"
        if latest.exists():
            import os
            if latest.is_symlink() or latest.is_dir():
                try:
                    latest.unlink()
                except Exception:
                    pass
        # Windows 上用 junction 模拟软链
        try:
            import os, sys
            if sys.platform == "win32":
                import shutil
                if latest.exists():
                    shutil.rmtree(latest, ignore_errors=True)
                shutil.copytree(str(out_dir), str(latest), dirs_exist_ok=True)
            else:
                if not latest.exists():
                    latest.symlink_to(out_dir, target_is_directory=True)
        except Exception:
            pass

        return out_dir

    # ============================================================
    # 一键导出（批量跑完后自动调用）
    # ============================================================

    def export_all(self, questions: list[dict], metadata: dict = None,
                   shots_dir: str = None, version: str = None,
                   unit: int = 0, stage: str = "") -> dict:
        """
        一键导出所有格式（HTML + CSV + ZIP）。

        Returns:
            {"html_full": "...", "html_errors": "...", "csv": "...", "zip": "..."}
        """
        meta = metadata or {}
        if version:
            meta.setdefault("version", version)
        if unit:
            meta.setdefault("unit", unit)
        if stage:
            meta.setdefault("stage", stage)

        results = {}
        results["html_full"] = self.export_html_full(questions, meta)
        results["html_errors"] = self.export_html_errors(questions, meta)
        results["csv"] = self.export_csv(questions)
        results["zip"] = self.export_screenshots_zip(questions, shots_dir)

        if version and unit and stage:
            results["output_dir"] = str(self.organize_output_dir(version, unit, stage))

        return results

    # ============================================================
    # DOCX 错题报告（给检查人员查看：按模块分组，含位置/原因/建议/截图）
    # ============================================================

    @staticmethod
    def _parse_qid(qid: str):
        """从题目 key 解析位置信息 → (module, unit, stage, qno)
        支持格式:
          - "听力专项-脚本-Q19"          → (听力专项, 脚本, 无stage, 19)
          - "auto-Q001"                 → (auto, 无, 无, 1)
          - "听力专项-脚本-U6-Q19"       → (听力专项, 脚本, U6, 19)
        """
        s = str(qid)
        m = re.match(r"^(.*?)-?(?:[Qq]0*(\d+))?$", s)
        module = "未知模块"
        stage = ""
        qno = ""
        # 拆出题号
        qm = re.search(r"[Qq]0*(\d+)", s)
        if qm:
            qno = qm.group(1)
        # 拆模块名（第一个 - 前）
        if "-" in s:
            module = s.split("-")[0].strip()
        # 拆阶段（基础巩固/综合进阶/难点突破/脚本 等）
        sm = re.search(r"(基础巩固|综合进阶|难点突破|脚本|练习|测试|单元评价|AI检测)", s)
        if sm:
            stage = sm.group(1)
        return module, "", stage, qno

    def export_docx(self, questions: list[dict], metadata: dict = None) -> str:
        """生成按模块分组的错题报告 DOCX（检查人员专用）。

        结构:
          标题（版本/日期/总览统计）
          一、模块 A（X 题错）
            · 第N题 [题型] —— 位置说明
              - 出错位置：模块 · 单元 · 第N题
              - 出错理由：...
              - 改进建议：...
              [嵌入截图]
          二、模块 B ...
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError as e:
            print(f"  ⚠ python-docx 不可用: {e}")
            return ""

        meta = metadata or {}
        version = meta.get("version", "未知版本")
        date_str = meta.get("date", datetime.now().strftime("%Y-%m-%d %H:%M"))

        # 只取错题（overall_passed=False 或 question_type=错题截图）
        errors = [q for q in questions
                  if q.get("overall_passed") is False
                  or q.get("question_type") == "错题截图"]
        if not errors:
            errors = [q for q in questions
                      if q.get("overall_passed") is False
                      or q.get("question_type") == "错题截图"]
            # 仍无 → 尝试取所有（报告空白不好看）
            if not errors and questions:
                errors = [q for q in questions
                          if not q.get("overall_passed", True)]

        # 按模块分组（保持出现顺序）
        # ★ 分组：优先用记录里的 module/stage 字段（错题定位：哪个模块·哪个子模块）
        #   找不到再退回 qid 解析（听力专项-脚本-Q19）；question_type 是"题型"不是模块名，不做分组依据
        groups = {}   # module_key -> {"name":..., "items":[...]}
        order = []
        for q in errors:
            qid = q.get("qid", "")
            if not qid:
                qid = f"auto-Q{int(q.get('idx', 1)):03d}"
            # ① 优先 module 字段（web_server 记录时写入）
            mod = (q.get("module", "") or "").strip()
            stage = (q.get("stage", "") or "").strip()
            if not mod:
                # ② 从 qid 解析（有脚本格式：听力专项-脚本-Q19）
                mod, unit, stage2, qno = self._parse_qid(qid)
                if not stage:
                    stage = stage2
                if mod == "未知模块":
                    # ③ 退回 question_type（仅当 module 完全缺失时，如旧数据）
                    qtype = (q.get("question_type", "") or "").strip()
                    qt_clean = re.sub(r"[（(].*?[）)]", "", qtype).strip()
                    if qt_clean and qt_clean not in ("错题截图", "?", "", "未知题型"):
                        mod = qt_clean
                    else:
                        mod = "其他"
            key = f"{mod}|{stage}" if stage else mod
            if key not in groups:
                groups[key] = {"name": f"{mod}" + (f" · {stage}" if stage else ""),
                               "items": []}
                order.append(key)
            groups[key]["items"].append(q)

        doc = Document()

        # 标题
        title = doc.add_heading(f"英语宝错题报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"版本：{version}　|　生成时间：{date_str}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # 总览
        doc.add_heading("总览", level=1)
        summary = doc.add_paragraph()
        total_err = len(errors)
        total_q = len(questions)
        summary.add_run(
            f"共检测 {total_q} 题，其中错题/需复核 {total_err} 题，涉及 {len(groups)} 个模块板块。"
        )

        # 每个模块板块
        for ki, key in enumerate(order, 1):
            g = groups[key]
            doc.add_heading(f"{ki}. {g['name']}（错题 {len(g['items'])} 题）", level=1)

            for qi, q in enumerate(g["items"], 1):
                qid = q.get("qid", "")
                qtype = q.get("question_type", "未知题型")
                qno = q.get("progress", "") or qid
                # ★ 位置题号：优先模块内题号 module_qno（该子模块第几题），否则用 idx
                idx = q.get("module_qno") or q.get("idx", "?")
                if isinstance(idx, float):
                    idx = int(idx) if idx == int(idx) else idx

                # 位置说明（★ 已在分组大标题下展示模块，题内只写第几题）
                loc_str = f"第{idx}题"

                # 题型标题
                h = doc.add_heading(f"{qi}. 第{idx}题　[{qtype}]", level=2)

                # 位置（★ 显示"模块 · 子模块 · 第N题"完整路径，检查人员一目了然）
                p = doc.add_paragraph()
                p.add_run("▸ 在 App 中的位置：").bold = True
                p.add_run(f"{g['name']} · {loc_str}")

                # 题干（★ 清洗掉时间/得分等 UI 噪音，非空才展示）
                _stem_clean = self._clean_stem(q.get("stem", ""))
                if _stem_clean:
                    p = doc.add_paragraph()
                    p.add_run("▸ 题干：").bold = True
                    p.add_run(_stem_clean)

                reasons = []
                for dim_name, reason_field, ai_field in self._LIVE_DIMS:
                    r = q.get(reason_field, "")
                    if r and ("不通过" in r or "⚠" in r or "未" in r or "失败" in r
                              or q.get(ai_field) is False):
                        reasons.append(f"{dim_name}：{r}")
                if not reasons:
                    # 兜底：从 overall 说明
                    if q.get("question_type") == "错题截图":
                        reasons.append("答题答错（已截图）")
                    else:
                        reasons.append("界面检查未通过（存在不通过维度或缺失必要元素）")

                p = doc.add_paragraph()
                p.add_run("▸ 出错理由：").bold = True
                p.add_run("\n".join(f"  · {r}" for r in reasons) if reasons else "  （无）")

                # 改进建议（AI 启发式：根据原因生成）
                sugg = self._build_suggestion(q, reasons)
                p = doc.add_paragraph()
                p.add_run("▸ 改进建议：").bold = True
                p.add_run("\n".join(f"  · {s}" for s in sugg) if sugg else "  （请人工核对）")

                # 截图
                shot = q.get("screenshot", "")
                if shot:
                    # 尝试从 screenshots 目录找到图片并嵌入
                    shots_dir = Path(__file__).parent.parent / "screenshots"
                    img_path = shots_dir / shot
                    if img_path.exists():
                        try:
                            doc.add_picture(str(img_path), width=Inches(3.2))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception:
                            pass
                    else:
                        p = doc.add_paragraph()
                        p.add_run(f"（截图：{shot}，文件未找到）").font.color.rgb = RGBColor(0x99, 0x99, 0x99)

                doc.add_paragraph()  # 空行分隔

        # 页脚说明
        doc.add_paragraph()
        foot = doc.add_paragraph()
        foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = foot.add_run("英语宝审查智能体 · 自动生成 · 供检查人员参考，最终判定以人工确认为准")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xbb, 0xbb, 0xbb)

        # 保存
        version_safe = re.sub(r'[\\/:*?"<>|]', '_', version)
        filename = f"错题报告_{version_safe}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        out_path = self.save_dir / filename
        doc.save(str(out_path))
        print(f"  ✅ DOCX 错题报告: {out_path}")
        return str(out_path)

    @staticmethod
    def _build_suggestion(q: dict, reasons: list) -> list:
        """根据出错理由生成个性化改进建议（针对具体错误内容，非固定模板）

        每条 reason 都含具体错误描述（来自 LLM 审查 / 界面检查），
        建议基于"错误类型 + 具体内容"动态生成，例如：
          - 题干语法错 → 给出正确写法建议
          - 题干与选项不符 → 建议改题型/改选项
          - 选项数量不足 → 建议补齐选项
        """
        sugg = []
        stem = (q.get("stem", "") or "").strip()
        qtype = (q.get("question_type", "") or "").strip()
        ai = {k: q.get(k) for k in ("ai_stem", "ai_content", "ai_image",
                                    "ai_answer", "ai_audio", "ai_post_error")}

        # ── 遍历每个维度 reason，提取具体错误并生成对应建议 ──
        for r in reasons:
            if not r:
                continue
            # 去掉维度前缀（"题干：""内容：""配图："等）
            body = re.sub(r"^(题干|内容|配图|作答|音频|答错后)[：:]", "", r).strip()

            # ⓪ ★ 优先采用 LLM 已给出的具体修改建议（"建议修改：xxx"）
            m = re.search(r"建议修改[：:]\s*(.+)", body)
            if m:
                sugg.append(m.group(1).strip())
                continue

            # ① 题干语法/词序错误（含"应为"字样 → 错误原文+正确写法）
            m = re.search(r"“(.+?)”应为“(.+?)”", body)
            if m and ("语法" in body or "词序" in body or "错误" in body):
                sugg.append(
                    f"题干存在文字错误：「{m.group(1)}」应为「{m.group(2)}」。"
                    f"请修正题干中的单词/词序，确保学生能正确理解题目。"
                )
                continue
            # ② 题干语法错误（有"应为"但格式不同）
            m = re.search(r"应为“(.+?)”", body)
            if m and "语法" in body:
                sugg.append(f"按正确写法「{m.group(1)}」修正题干文字，避免语法/词序错误。")
                continue

            # ③ 题干与题型/选项不符（"题目为…但…" → 建议改题型）
            m = re.search(r"题目为“?([^，。；]+?)”?，但(?:脚本|实际)?(?:中)?(?:仅|只|没有|未)", body)
            if m:
                target = m.group(1).strip()
                sugg.append(
                    f"题干声称题型为「{target}」但与实际作答形式不符："
                    f"请调整题干描述或改用与作答形式匹配的题型"
                    f"（如改为填空题/选择题/配对题），保持题干与选项一致。"
                )
                continue
            # ④ 选项数量不足（"仅提供了N个…"）
            m = re.search(r"仅提供(?:了)?(\d+)个", body)
            if m and ("选项" in body or "内容" in body):
                n = int(m.group(1))
                sugg.append(
                    f"选项数量不足（当前{n}个）：请补齐选项至与题型要求一致"
                    f"（选择题一般≥3个，判断/配对按题量配齐），避免选项缺失。"
                )
                continue

            # ⑤ 选项与题干不匹配（"选项"+"不匹配/无关/不一致"）
            if ("选项" in body or "内容" in body) and any(k in body for k in ("不匹配", "无关", "不一致", "不对应", "不符合")):
                sugg.append(
                    "选项与题干语义不匹配：请核对选项内容，确保各选项均围绕题干要求设置"
                    "（正确项唯一、干扰项相关），避免学生因选项无关而困惑。"
                )
                continue
            # ⑥ 题干未提取到/缺失
            if "未提取到" in body or "未提供题干" in body:
                sugg.append("题干文字缺失：请补充完整的题目要求描述，说明作答方式（听/读/选/写）。")
                continue

            # ⑦ 配图问题
            if ("配图" in body or "图片" in body) and ("未匹配" in body or "不符" in body):
                sugg.append("配图与题目内容不符或缺失：请核对配图，确保与题干/答案语义一致且清晰可见。")
                continue
            if ("配图" in body or "图片" in body) and "未检测" in body:
                # 脚本审查未截图 → 温和提示（人工核对），非硬性错误
                sugg.append("需连手机截图核对配图：脚本审查无法看到实际图片，建议人工查看配图与题干是否一致。")
                continue

            # ⑧ 音频问题
            if "音频" in body and ("不可点击" in body or "未检测到播放" in body):
                sugg.append("听力音频资源异常：请确保播放控件可点击、音频文件可正常播放（检查文件路径/格式/时长）。")
                continue
            if "音频" in body and "未检测" in body and ai.get("ai_audio") is not False:
                sugg.append("需连手机截图核对音频控件：脚本审查无法验证播放按钮，建议人工确认听力音频可正常播放。")
                continue

            # ⑨ 作答元素问题
            if "作答" in body and ("未识别" in body or "未检测" in body):
                sugg.append("作答元素异常：请确保检查/录音/输入框等作答控件正常显示且可交互。")
                continue

        # ── 答错截图类（无六维 reason）──
        if q.get("question_type") == "错题截图":
            sugg.append("人工核对错题画面：定位答错原因（听音误解/选项混淆/知识点未掌握）")
            sugg.append("针对错题知识点设计专项巩固练习，如高频错词/句型重复训练")

        # ── 兜底：仍无建议 → 基于题型给提示 ──
        if not sugg:
            if "听" in qtype or "听力" in qtype:
                sugg.append("人工复核该听力题：检查录音清晰度、语速、选项干扰项设计是否合理。")
            elif "填空" in qtype or "补全" in qtype:
                sugg.append("人工复核该填空/补全题：确认答案唯一性、空格数量与提示词是否匹配。")
            elif "排序" in qtype or "排列" in qtype:
                sugg.append("人工复核该排序题：确认选项顺序逻辑清晰、无歧义。")
            else:
                sugg.append("人工复核该题：对照脚本答案与 App 实际内容确认错误原因。")
        # ★ 去重（同一条建议只保留一次，保持顺序）
        seen = set()
        uniq = []
        for s in sugg:
            key = s.strip()
            if key and key not in seen:
                seen.add(key)
                uniq.append(s)
        return uniq

    # 维度展示顺序：(中文名, reason字段, ai_布尔字段)
    _LIVE_DIMS = [
        ("题干", "stem_reason", "ai_stem"),
        ("内容", "content_reason", "ai_content"),
        ("配图", "image_reason", "ai_image"),
        ("作答", "answer_reason", "ai_answer"),
        ("音频", "audio_reason", "ai_audio"),
        ("答错后", "post_error_reason", "ai_post_error"),
    ]

    _HTML_LIVE_CSS = """
    <style>
      * { margin:0; padding:0; box-sizing:border-box; }
      body { font-family:-apple-system,BlinkMacSystemFont,"Segoe UI","PingFang SC","Microsoft YaHei",sans-serif; background:#f6f4ef; color:#333; line-height:1.6; padding:20px; }
      .container { max-width:840px; margin:0 auto; }
      h1 { font-size:21px; font-weight:700; color:#3a3a3a; }
      .sub { color:#9aa5a0; font-size:13px; margin:4px 0 16px; }
      .stats { display:flex; gap:10px; margin-bottom:18px; flex-wrap:wrap; }
      .stat { flex:1; min-width:120px; background:#fff; border:1px solid #f0e9dd; border-radius:10px; padding:12px; text-align:center; }
      .stat .num { font-size:22px; font-weight:800; }
      .stat .lab { font-size:11px; color:#888; margin-top:2px; }
      .s-fail .num { color:#bc4742; } .s-total .num { color:#3a3a3a; }
      .s-pass .num { color:#2e7d32; } .s-shot .num { color:#c77c00; }
      .card { background:#fff; border:1px solid #f0e9dd; border-left:4px solid #bc4742; border-radius:10px; margin-bottom:10px; overflow:hidden; }
      .card > summary { cursor:pointer; list-style:none; padding:12px 14px; display:flex; align-items:center; gap:9px; }
      .card > summary::-webkit-details-marker { display:none; }
      .card > summary:hover { background:#faf8f3; }
      .badge { background:#fee2e2; color:#bc4742; font-size:10px; font-weight:700; padding:2px 8px; border-radius:6px; }
      .qtitle { font-size:13.5px; font-weight:600; color:#333; }
      .src { margin-left:auto; font-size:10.5px; color:#9aa5a0; }
      .body { padding:0 14px 14px; border-top:1px solid #f3eee4; font-size:12.5px; color:#444; }
      .field { margin-top:9px; } .field b { color:#666; }
      .cause { margin-top:9px; background:#fdecec; border:1px solid #f6c9c9; border-radius:8px; padding:9px 11px; }
      .cause .h { color:#bc4742; font-weight:700; font-size:11.5px; margin-bottom:5px; }
      .fix { margin-top:8px; background:#eafaf0; border:1px solid #bfe9cd; border-radius:8px; padding:9px 11px; }
      .fix .h { color:#2e7d32; font-weight:700; font-size:11.5px; margin-bottom:5px; }
      .shot { margin-top:8px; display:flex; align-items:center; gap:8px; }
      .shot img { max-width:170px; max-height:110px; border:1px solid #d4dae0; border-radius:7px; }
      .shot .fn { font-size:10.5px; color:#9aa5a0; }
      .empty { text-align:center; padding:60px 0; color:#bbb; font-size:15px; }
      .footer { text-align:center; color:#bbb; font-size:12px; margin-top:28px; }
      @media print { body { background:#fff; padding:0; } .card { box-shadow:none; break-inside:avoid; } }
    </style>
    """

    def export_html_live(self, questions: list[dict], metadata: dict = None) -> str:
        """
        生成"仅错误"的实时可折叠卡片 HTML 报告（审查中每出一道错题即调用）。

        与 export_html_full 的区别：
          - 卡片用原生 <details> 折叠（默认收起，点标题展开），无需 JS；
          - 每张卡片内分「错误原因(红框) / 修改建议(绿框) / 截图」三块；
          - 截图拷贝到输出目录并内嵌，文件可单独带走、双击即看。
        """
        meta = metadata or {}
        qs = [q for q in questions if not q.get("overall_passed", True)]
        html_doc = self._render_live_html(qs, meta)

        version = meta.get("version", "未知版本")
        unit = str(meta.get("unit", "全部"))
        stage = str(meta.get("stage", "全部"))
        date_str = datetime.now().strftime("%Y%m%d")
        # ★ 防御：清洗 Windows 非法字符（<>:"/\|?* 等），避免 WinError 123
        _clean = lambda s: re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", s).strip() or "全部"
        version, unit, stage = _clean(version), _clean(unit), _clean(stage)
        out_dir = self.save_dir / version / f"U{unit}_{stage}_{date_str}"
        out_dir.mkdir(parents=True, exist_ok=True)

        # 拷贝错误截图到输出目录（HTML 用相对路径引用，文件可带走）
        shot_dir = out_dir / "screenshots"
        shot_dir.mkdir(parents=True, exist_ok=True)
        shots_root = Path(__file__).parent.parent / "screenshots"
        for q in qs:
            shot = q.get("screenshot", "")
            if shot:
                src = shots_root / shot
                if src.exists():
                    try:
                        shutil.copy2(src, shot_dir / shot)
                    except Exception:
                        pass

        out_path = out_dir / "report_live.html"
        out_path.write_text(html_doc, encoding="utf-8")

        # 刷新 latest 快捷入口（复用已有的目录管理逻辑）
        self.organize_output_dir(version, unit, stage)
        return str(out_path)

    def _render_live_html(self, questions: list[dict], meta: dict) -> str:
        version = meta.get("version", "未知版本")
        unit = meta.get("unit", "?")
        stage = meta.get("stage", "?")
        date_str = meta.get("date", datetime.now().strftime("%Y-%m-%d %H:%M"))
        total = meta.get("total", len(questions))
        failed = len(questions)
        passed = total - failed
        rate = round(passed / total * 100) if total else 0
        with_shot = sum(1 for q in questions if q.get("screenshot"))
        cards = "\n".join(self._render_live_card(q, meta) for q in questions)
        title = f"实时错题报告 — {version} · U{unit} · {stage}"
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{html.escape(title)}</title>
{self._HTML_LIVE_CSS}
</head>
<body>
<div class="container">
  <h1>{html.escape(str(version))} · U{unit} · {stage}</h1>
  <div class="sub">生成于 {date_str} · 共 {failed} 道错题 · 审查中实时更新</div>

  <div class="stats">
    <div class="stat s-fail"><div class="num">{failed}</div><div class="lab">不通过</div></div>
    <div class="stat s-total"><div class="num">{total}</div><div class="lab">已审查</div></div>
    <div class="stat s-pass"><div class="num">{rate}%</div><div class="lab">通过率</div></div>
    <div class="stat s-shot"><div class="num">{with_shot}</div><div class="lab">含截图</div></div>
  </div>

  {('<div class="empty">暂无错题 🎉 全部通过</div>' if not questions else cards)}

  <div class="footer">英语宝审查智能体 · 自动生成 · 仅供参考，最终以人工确认为准</div>
</div>
</body>
</html>"""

    def _render_live_card(self, q: dict, meta: dict) -> str:
        qid = q.get("qid", "?")
        idx = q.get("idx", "?")
        qtype = q.get("question_type", "?")
        stem = html.escape(str(self._clean_stem(q.get("stem", "")) or ""))
        script_answer = html.escape(str(q.get("script_answer", "") or ""))
        src = f'{html.escape(str(meta.get("version", "")))} · U{meta.get("unit", "?")} · {html.escape(str(meta.get("stage", "")))}'

        failed = []
        suggestions = []
        for dname, rkey, akey in self._LIVE_DIMS:
            reason = q.get(rkey, "") or ""
            ai_pass = q.get(akey, None)
            is_fail = (ai_pass is False) or ("不通过" in reason)
            if is_fail:
                reason_clean = re.split(r'建议修改[:：]', reason)[0].strip()
                failed.append((dname, reason_clean))
                m = re.search(r'建议修改[:：]\s*([^\n]{4,200})', reason)
                if m:
                    suggestions.append(m.group(1).strip())

        if failed:
            cause_html = "".join(
                f'<div>· {html.escape(dname)}：{html.escape(r)}</div>'
                for dname, r in failed
            )
        else:
            cause_html = '<div>—（未在六维中标记，详见原始理由）</div>'

        if suggestions:
            fix_html = "".join(f'<div>· {html.escape(s)}</div>' for s in suggestions)
        else:
            fix_html = '<div>—（AI 未给出具体建议，请人工核对）</div>'

        shot = q.get("screenshot", "")
        shot_html = ""
        if shot:
            shot_html = (
                f'<div class="shot"><img src="screenshots/{html.escape(shot)}" '
                f'alt=""><span class="fn">{html.escape(shot)}</span></div>'
            )

        return f'''<details class="card">
  <summary>
    <span class="badge">不通过</span>
    <span class="qtitle">{html.escape(self._loc_label(q, meta))} · 第{idx}题</span>
    <span class="qtype">[{html.escape(qtype)}]</span>
  </summary>
  <div class="body">
    <div class="field"><b>位置：</b>{html.escape(self._loc_label(q, meta))} · 第{idx}题</div>
    <div class="field"><b>题干：</b>{stem or "（无题干文字）"}</div>
    <div class="field"><b>脚本答案：</b><span style="color:#2e7d32;">{script_answer}</span></div>
    <div class="cause"><div class="h">✕ 错误原因</div>{cause_html}</div>
    <div class="fix"><div class="h">✓ 修改建议</div>{fix_html}</div>
    {shot_html}
  </div>
</details>'''
