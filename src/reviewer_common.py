"""
reviewer_common.py — 双人协作共享层

职责: 数据模型、LLM 客户端、题目加载器
原则: 两个人都不需要改这个文件（除非要加新数据类型）
"""

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
from difflib import SequenceMatcher


# ============================================================
# 数据模型（全项目统一）
# ============================================================

@dataclass
class CheckItem:
    """单项检查结果 — 贯穿四项检查的通用格式"""
    name: str                          # 检查项名称（中文）
    passed: bool = False               # 是否通过
    screenshot: str = ""               # 截图路径
    actual_text: str = ""              # 实际检测到的文字
    expected_text: str = ""            # 脚本中的预期文字
    details: list = field(default_factory=list)  # 详细描述
    error: str = ""                    # 出错信息
    similarity: float = 0.0            # 相似度（0-1）

    def to_dict(self):
        return {
            "name": self.name,
            "passed": self.passed,
            "actual_text": self.actual_text,
            "expected_text": self.expected_text,
            "similarity": round(self.similarity, 3),
            "details": self.details,
            "error": self.error,
            "screenshot": self.screenshot,
        }


@dataclass
class Question:
    """单道题目数据结构"""
    idx: int                           # 题号
    stem: str = ""                     # 题干文字
    content: str = ""                  # 题目内容（选项、说明等）
    correct_answer: str = ""           # 正确答案
    question_type: str = ""            # 题型
    knowledge_points: list = field(default_factory=list)
    image_paths: list = field(default_factory=list)  # 配图路径
    expected_image_desc: str = ""      # 脚本中的配图描述（用于对比实际截图）

    def to_dict(self):
        return {
            "idx": self.idx,
            "stem": self.stem,
            "content": self.content,
            "correct_answer": self.correct_answer,
            "question_type": self.question_type,
            "knowledge_points": self.knowledge_points,
            "image_paths": self.image_paths,
        }


# ============================================================
# 题目文件加载器（支持多种格式）
# ============================================================

class QuestionLoader:
    """加载公司提供的题目文件（JSON / Excel / CSV）"""

    @staticmethod
    def from_json(filepath: str) -> list[Question]:
        """从 JSON 文件加载题目列表"""
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)

        # 自动识别 JSON 结构：可能是 { "questions": [...] } 或直接是数组
        if isinstance(data, dict):
            raw = data.get("questions", data.get("data", []))
        elif isinstance(data, list):
            raw = data
        else:
            raise ValueError(f"不支持的 JSON 结构: {type(data)}")

        questions = []
        for item in raw:
            q = Question(
                idx=item.get("id", item.get("idx", item.get("question_id", 0))),
                stem=item.get("stem", item.get("stem_text", item.get("question", ""))),
                content=item.get("content", item.get("content_text", "")),
                correct_answer=item.get("answer", item.get("correct_answer", item.get("key", ""))),
                question_type=item.get("type", item.get("question_type", "")),
                knowledge_points=item.get("knowledge_points", item.get("knowledge", [])),
                image_paths=item.get("images", item.get("image_paths", [])),
            )
            questions.append(q)

        return questions

    @staticmethod
    def from_docx(filepath: str) -> list[Question]:
        """从 DOCX 文件加载题目（自动识别表格/段落格式）"""
        from docx import Document
        doc = Document(filepath)

        # 方式1: 尝试表格模式
        questions = QuestionLoader._parse_docx_tables(doc)
        if questions:
            return questions

        # 方式2: 段落模式（Q1: xxx \n A: xxx 格式）
        return QuestionLoader._parse_docx_paragraphs(doc)

    @staticmethod
    def _parse_docx_tables(doc) -> list[Question]:
        """从 DOCX 表格中解析题目"""
        questions = []
        for table in doc.tables:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            if not any(k in headers for k in ['id','题号','题目','题干','stem','question']):
                continue

            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) < 2 or not cells[0]:
                    continue
                data = dict(zip(headers, cells))
                q = QuestionLoader._map_docx_row(data)
                if q:
                    questions.append(q)
        return questions

    @staticmethod
    def _parse_docx_paragraphs(doc) -> list[Question]:
        """从 DOCX 段落中解析题目（Q1/Q01/1. 格式）"""
        import re as regex
        questions = []
        current = None
        buf = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检测新题目开始: Q1 / 1. / 第1题 等
            m = regex.match(r'^(?:Q|q)?\s*(\d+)[\.\、\s]|^第\s*(\d+)\s*题', text)
            if m:
                if current:
                    current.stem = '\n'.join(buf[:1]) if buf else ''
                    current.content = '\n'.join(buf[1:]) if len(buf) > 1 else ''
                    questions.append(current)
                qid = int(m.group(1) or m.group(2))
                current = Question(idx=qid)
                buf = [text]
            elif current is not None:
                # 检测答案行
                if regex.match(r'^(答案|Answer|Key)[:：]', text):
                    current.correct_answer = text.split('：')[-1].split(':')[-1].strip()
                elif regex.match(r'^(题型|类型|Type)[:：]', text):
                    current.question_type = text.split('：')[-1].split(':')[-1].strip()
                else:
                    buf.append(text)

        if current:
            current.stem = '\n'.join(buf[:1]) if buf else ''
            current.content = '\n'.join(buf[1:]) if len(buf) > 1 else ''
            questions.append(current)

        return questions

    @staticmethod
    def _map_docx_row(data: dict) -> Question:
        """将 DOCX 表格行映射为 Question（支持中英文列名）"""
        def get(*keys):
            for k in keys:
                if k in data:
                    return data[k]
            return ""
        return Question(
            idx=int(get('id', 'idx', '题号', '序号', 'question_id', '0') or 0),
            stem=get('stem', '题干', '题目', 'question', 'stem_text', '内容'),
            content=get('content', '选项', '内容', 'options', 'content_text'),
            correct_answer=get('answer', '答案', 'key', 'correct_answer', '正确答案'),
            question_type=get('type', '题型', 'question_type', '类型'),
            expected_image_desc=get('image_desc', '配图描述', '图片描述', '配图', 'image', '图片'),
            knowledge_points=get('knowledge', '知识点', 'knowledge_points').split(';') if get('knowledge', '知识点', 'knowledge_points') else [],
        )
        """从 CSV 文件加载"""
        import csv
        questions = []
        with open(filepath, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                q = Question(
                    idx=int(row.get("id", row.get("idx", 0))),
                    stem=row.get("stem", row.get("question", "")),
                    content=row.get("content", row.get("options", "")),
                    correct_answer=row.get("answer", row.get("key", "")),
                    question_type=row.get("type", ""),
                )
                questions.append(q)
        return questions

    @staticmethod
    def from_excel(filepath: str) -> list[Question]:
        """从 Excel 文件加载"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath)
            ws = wb.active
            headers = [cell.value for cell in ws[1]]
            questions = []
            for row in ws.iter_rows(min_row=2, values_only=True):
                data = dict(zip(headers, row))
                q = Question(
                    idx=int(data.get("id", data.get("idx", 0))),
                    stem=data.get("stem", data.get("question", "")),
                    content=data.get("content", data.get("options", "")),
                    correct_answer=data.get("answer", data.get("key", "")),
                    question_type=data.get("type", ""),
                )
                questions.append(q)
            return questions
        except ImportError:
            raise ImportError("读取 Excel 需要安装 openpyxl: pip install openpyxl")

    @staticmethod
    def describe(filepath: str) -> str:
        """预览文件内容，帮助确认解析是否正确"""
        questions = QuestionLoader.load(filepath)
        lines = [f"文件: {filepath}", f"题目数: {len(questions)}", ""]
        for q in questions:
            lines.append(f"Q{q.idx:02d} | {q.question_type or '?'} | {q.stem[:40]}")
            if q.correct_answer:
                lines.append(f"     答案: {q.correct_answer}")
            if q.expected_image_desc:
                lines.append(f"     配图: {q.expected_image_desc}")
        return "\n".join(lines)

    @staticmethod
    def load(filepath: str) -> list[Question]:
        """自动识别文件格式并加载"""
        ext = Path(filepath).suffix.lower()
        if ext == ".json":
            return QuestionLoader.from_json(filepath)
        elif ext == ".csv":
            return QuestionLoader.from_csv(filepath)
        elif ext in (".xlsx", ".xls"):
            return QuestionLoader.from_excel(filepath)
        elif ext in (".docx", ".doc"):
            return QuestionLoader.from_docx(filepath)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")


# ============================================================
# LLM 客户端（统一接口 — 你配置你的 LLM）
# ============================================================

class LLMClient:
    """
    LLM 统一调用接口

    使用方法（三种，任选一种）：
        # 方式1: 自动从 llm_config.json 读取（推荐）
        client = LLMClient.from_config()

        # 方式2: 直接传参数
        client = LLMClient(api_key="sk-xxx", model="deepseek-chat",
                           base_url="https://api.deepseek.com/v1")

        # 方式3: 从环境变量读取
        export OPENAI_API_KEY="sk-xxx"
        client = LLMClient(model="deepseek-chat")

    调用：
        result = client.ask("请检查这段文字是否有错别字")
        result = client.ask("请分析这张图片", image_path="q01.png")
    """

    # 配置文件路径（相对于项目根目录）
    CONFIG_PATH = "llm_config.json"

    def __init__(self, api_key: str = None, model: str = "deepseek-chat",
                 base_url: str = "https://api.deepseek.com/v1",
                 vision_api_key: str = "", vision_model: str = "",
                 vision_base_url: str = ""):
        self.api_key = api_key or os.environ.get("OPENAI_API_KEY", "")
        self.model = model
        self.base_url = base_url.rstrip("/")
        # 视觉模型（处理图片输入）
        self.vision_api_key = vision_api_key or self.api_key
        self.vision_model = vision_model
        self.vision_base_url = vision_base_url.rstrip("/") if vision_base_url else ""

    @classmethod
    def from_config(cls, config_path: str = None) -> "LLMClient":
        """
        从 llm_config.json 自动读取配置

        优先级：环境变量 > config 文件 > 默认值
        """
        path = config_path or cls.CONFIG_PATH
        config = {}

        if Path(path).exists():
            with open(path, "r", encoding="utf-8") as f:
                config = json.load(f)

        return cls(
            api_key=os.environ.get("OPENAI_API_KEY") or config.get("api_key", ""),
            model=os.environ.get("LLM_MODEL") or config.get("model", "deepseek-chat"),
            base_url=os.environ.get("LLM_BASE_URL") or config.get("base_url", "https://api.deepseek.com/v1"),
            vision_api_key=config.get("vision_api_key", ""),
            vision_model=config.get("vision_model", ""),
            vision_base_url=config.get("vision_base_url", ""),
        )

    def ask(self, prompt: str, image_path: str = None, image_paths: list = None) -> str:
        """
        智能路由：
          - 无图片 → 主模型 (DeepSeek)
          - 有图片 + 有视觉模型 → 视觉模型 (qwen3.7-plus)
          - 有图片 + 无视觉模型 → 主模型 (OCR 降级)
          
        参数:
          prompt: 提示词
          image_path: 单张图片路径 (兼容)
          image_paths: 多张图片路径列表 (优先于 image_path)
        """
        paths = image_paths if image_paths else ([image_path] if image_path else None)

        if paths and self.vision_model:
            # 有视觉模型：直接发图
            try:
                return self._call_api(prompt, paths,
                    model=self.vision_model,
                    base_url=self.vision_base_url or self.base_url,
                    api_key=self.vision_api_key)
            except Exception as e:
                return f"[视觉模型失败: {e}] → 降级为纯文本分析\n{self._call_api(prompt, image_paths=None)}"

        if paths:
            # 无视觉模型：OCR 降级
            return self._ask_with_ocr(prompt, paths[0])

        # 纯文本
        try:
            return self._call_api(prompt, image_path=None)
        except Exception as e:
            return f"[LLM 调用失败] {e}"

    def _ask_with_ocr(self, prompt: str, image_path: str) -> str:
        """
        两步走：本地 OCR 提取截图文字 → 主模型分析
        """
        ocr_text = ""
        try:
            # 优先用项目自带的 OCR 引擎
            from src.ocr_engine import OCREngine
            engine = OCREngine()
            result = engine.extract(image_path)
            texts = [b.text for b in result if hasattr(b, 'text') and b.text]
            ocr_text = " ".join(texts)
        except Exception:
            pass

        if not ocr_text:
            try:
                import subprocess
                # 兜底：Windows 上的 tesseract
                subprocess.run(["tesseract", image_path, "stdout"],
                             capture_output=True, timeout=10)
            except Exception:
                pass

        enhanced_prompt = prompt
        if ocr_text and len(ocr_text) > 5:
            enhanced_prompt += (
                f"\n\n[截图 OCR 文字]\n{ocr_text[:3000]}"
                f"\n\n请结合以上从截图中提取的文字进行分析。"
            )
        else:
            enhanced_prompt += (
                "\n\n(截图中的文字无法自动提取，请仅基于题目信息分析。"
                "如果题目是配图类题目，标注为[无法判断图文匹配--缺少视觉模型]。)"
            )

        return self._call_api(enhanced_prompt, image_path=None)

    def _call_api(self, prompt: str, image_paths: list = None,
                  model: str = None, base_url: str = None,
                  api_key: str = None) -> str:
        """底层 API 调用 (支持多图)"""
        import base64
        from urllib.request import Request, urlopen

        use_model = model or self.model
        use_url = (base_url or self.base_url).rstrip("/")
        use_key = api_key or self.api_key

        messages = [{"role": "user", "content": []}]
        messages[0]["content"].append({"type": "text", "text": prompt})

        if image_paths:
            for img_path in image_paths:
                if img_path and Path(img_path).exists():
                    with open(img_path, "rb") as f:
                        img_b64 = base64.b64encode(f.read()).decode("utf-8")
                    messages[0]["content"].append({
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{img_b64}"}
                    })

        body = json.dumps({
            "model": use_model,
            "messages": messages,
            "max_tokens": 1000,
        }).encode("utf-8")

        req = Request(
            f"{use_url}/chat/completions",
            data=body,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {use_key}",
            },
        )
        resp = urlopen(req, timeout=60)
        result = json.loads(resp.read())
        return result["choices"][0]["message"]["content"]


# ============================================================
# 文本相似度工具
# ============================================================

def text_similarity(a: str, b: str) -> float:
    """计算两个字符串的相似度（0-1）"""
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a.lower().strip(), b.lower().strip()).ratio()


def find_diff_positions(expected: str, actual: str) -> list[str]:
    """找出两个字符串之间的差异位置"""
    diffs = []
    matcher = SequenceMatcher(None, expected, actual)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag != "equal":
            diffs.append(f"[{tag}] 预期'{expected[i1:i2]}' vs 实际'{actual[j1:j2]}'")
    return diffs

import os
