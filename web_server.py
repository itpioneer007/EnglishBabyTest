#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
英语宝模块检测 - Web 控制面板
Flask 后端，提供 REST API 和前端页面

启动:
    cd 英语宝模块检测
    python web_server.py

打开浏览器访问: http://localhost:5000
"""

import os
import sys
import io

# Windows 控制台默认 GBK 无法输出 emoji(✅❌⚠ 等)，全局切换为 UTF-8
# ★ 必须用 reconfigure 而非替换对象：替换 sys.stdout 会让原 wrapper 被 GC
#   时连带关闭共享 buffer，导致 "I/O operation on closed file" 崩溃
if sys.platform == "win32":
    for _stream in (sys.stdout, sys.stderr):
        try:
            _stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass
import json
import time
import threading
import re
import subprocess as sp
from pathlib import Path
from datetime import datetime

# 添加 src 到路径
PROJECT_ROOT = Path(__file__).parent.absolute()
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from flask import Flask, jsonify, request, send_from_directory, render_template, Response
try:
    from adb_controller import ADBController
except ImportError:
    # 依赖缺失时占位，保证服务可启动（听力专项新引擎不依赖 ADBController）
    class ADBController:
        def __init__(self, *a, **kw): pass
from config_loader import load_config

app = Flask(__name__)

# ============================================================
# 三人协作路由注册（按角色拆分）
#   A: routes/trace_routes.py  → 溯源系统
#   B: routes/batch_routes.py  → 批量自动化
#   C: routes/export_routes.py → 报告输出分发
# ============================================================
from routes import register_trace, register_batch, register_export
register_trace(app)
register_batch(app)
# ★ 导出优先读内存 _inspection_state（多模块检测含全部模块；文件会被新任务覆盖）
register_export(app, state_provider=lambda: _inspection_state)
# ★ 借鉴队友：实时错题报告路由（/api/errors/live、/api/errors/live-status）
import routes.error_log_routes as _error_log_routes
_error_log_routes.register(app)

# ============================================================
# 全局状态
# ============================================================
task_status = {
    "running": False,
    "current_task": "",
    "progress": [],
    "start_time": "",
    "end_time": "",
    "log": [],
}

_last_screenshot = ""

# ★ 协作式停止标志（多模块检测用）：前端点停止 → True → 模块循环/答题循环中断
_STOP_REQUESTED = False
_STOP_LOCK = threading.Lock()

# ★ 当前活动任务线程 id（用于"立即停止"：停止时向该线程注入异常）
_CURRENT_TASK_THREAD_ID = None
_TASK_THREAD_LOCK = threading.Lock()


def _register_task_thread():
    """记录当前任务线程 id（各 run 接口启动线程时调用）"""
    global _CURRENT_TASK_THREAD_ID
    with _TASK_THREAD_LOCK:
        _CURRENT_TASK_THREAD_ID = threading.get_ident()


def _force_stop_thread():
    """强制中断当前任务线程：向线程注入 SystemExit（立即停止，不再执行脚本内容）

    Python 线程无法直接 kill，但可通过 PyThreadState_SetAsyncExc 向目标线程
    注入异常——异常在线程下一条 Python 字节码处抛出（若阻塞在 HTTP/sleep，
    则等待其返回后立即抛出，最长受 u2 HTTP_TIMEOUT=15s 限制）。
    """
    global _CURRENT_TASK_THREAD_ID
    with _TASK_THREAD_LOCK:
        tid = _CURRENT_TASK_THREAD_ID
    if not tid:
        return False
    try:
        import ctypes
        res = ctypes.pythonapi.PyThreadState_SetAsyncExc(
            ctypes.c_long(tid), ctypes.py_object(SystemExit))
        if res > 1:  # 注入失败（异常被吞）时撤销
            ctypes.pythonapi.PyThreadState_SetAsyncExc(ctypes.c_long(tid), None)
        return res == 1
    except Exception:
        return False


def _is_stop_requested():
    """线程安全读取停止标志（供 scheduler/engine/模块循环调用）"""
    global _STOP_REQUESTED
    with _STOP_LOCK:
        return _STOP_REQUESTED


def _request_stop():
    """统一停止入口：设置协作式停止标志 + 向任务线程注入 SystemExit 强制中断

    所有停止接口（modules/stop、inspect/stop、audio/stop 等）必须走这里，
    否则只改 running=False 无法真正停止脚本（线程内部不检查 running）。
    返回 True 表示已向任务线程发送强制中断信号。
    """
    global _STOP_REQUESTED
    if not task_status["running"]:
        return False, "没有正在运行的任务"
    # 1. 设协作式标志（scheduler/engine/模块循环内的 should_stop() 立即感知）
    with _STOP_LOCK:
        _STOP_REQUESTED = True
    # 2. 立即强制中断线程（下一条 Python 字节码处抛 SystemExit；若阻塞在
    #    u2 HTTP/time.sleep，最长等 HTTP_TIMEOUT=15s 返回后立即抛出）
    forced = _force_stop_thread()
    log_msg("⏹ 收到停止请求，立即中断任务…", "warning")
    if not forced:
        log_msg("⚠ 强制中断未生效，任务将在当前步骤结束后停止", "warning")
    return forced, ("停止信号已发送，任务已中断" if forced
                    else "停止信号已发送，等待当前步骤结束")

# 模块坐标（从 uiautomator dump 获取的精确坐标）
MODULE_COORDS = {
    # 教材精学 (人教版下的坐标)
    "课本点读(左)": (203, 1191),
    "课本点读(中)": (540, 1191),
    "巧记单词":     (876, 1191),
    "语音评测":     (203, 1358),
    # 专项突破 (可见)
    "听课文":    (161, 1792),
    "课文动画":  (414, 1792),
    "基础训练":  (666, 1792),
    "一课一练":  (919, 1792),
    "课文配音":  (161, 2033),
    "口语训练":  (414, 2033),
    "复习回顾":  (666, 2033),
    "全脑记词":  (919, 2033),
    # 专项突破 (需滚动)
    "单元自检":      (414, 746),
    "单元练习计划":  (666, 746),
    # 听力专项 (需滚动更深)
    "基础巩固":   (180, 1250),   # 听力专项下的三个阶段
    "综合进阶":   (540, 1250),
    "难点突破":   (900, 1250),
    "听力专项_五上": (414, 1250),   # 五上入口
    "听力专项_五下": (414, 950),    # 五下入口
    "听力专项_六上": (414, 1150),   # 六上入口
    "听力专项_六下": (414, 1050),   # 六下入口
}

# 需要先滚动才能看到的模块
DEEP_MODULES = {
    "单元自检",
    "单元练习计划",
    "教材同步题库",
    "听力训练",
    "你听一刻",
    # 听力专项下的内容
    "基础巩固",
    "综合进阶",
    "难点突破",
    "听力专项_五上",
    "听力专项_五下",
    "听力专项_六上",
    "听力专项_六下",
}

# 底部导航坐标
TABS = {
    "英语": (108, 2233),
    "我":   (972, 2220),
}

SETTINGS_ICON = (1000, 170)
AD_CLOSE = (540, 1821)

# ---- 模块自动识别 ----

# 文件名关键词 → 模块名 + 导航方式
MODULE_PATTERNS = {
    # (关键词列表, 模块名, 区域, 点击坐标)
    ("巧记单词", "巧记",): ("巧记单词", "教材精学", (876, 1191)),
    ("听力专项", "听力",): ("听力专项", "专项突破", None),  # None = 滚动搜索
    ("知识过关", "知识",): ("知识过关", "专项突破", None),
    ("基础训练",): ("基础训练", "专项突破", (666, 1792)),
    ("一课一练",): ("一课一练", "专项突破", (919, 1792)),
    ("复习回顾",): ("复习回顾", "专项突破", (666, 2033)),
    ("全脑记词",): ("全脑记词", "专项突破", (919, 2033)),
}

def detect_module_from_filename(filename: str):
    """从文件名提取模块名和导航方式, 返回 (module_name, section, coords) 或 None"""
    for keywords, (name, section, coords) in MODULE_PATTERNS.items():
        for kw in keywords:
            if kw in filename:
                return name, section, coords
    return None, None, None


# ============================================================
# 工具函数
# ============================================================


# ============================================================
# 自动坐标缩放 — 适配不同分辨率手机
# ============================================================
REFERENCE_RES = (1080, 2400)
_detected_res = None
_scale_x = 1.0
_scale_y = 1.0

def detect_screen_resolution(adb_serial: str = ""):
    """通过 ADB 检测手机实际分辨率，更新缩放比"""
    global _detected_res, _scale_x, _scale_y
    try:
        # ★ 跨机器自动定位 adb（复用 ADBController 的探测逻辑，不写死任何用户路径）
        adb_path = ADBController._find_adb()
        cmd = [adb_path]
        if adb_serial:
            cmd.extend(["-s", adb_serial])
        cmd.extend(["shell", "wm", "size"])
        r = sp.run(cmd, capture_output=True, text=True, timeout=5,
                   encoding="utf-8", errors="replace")
        m = re.search(r"(\d+)x(\d+)", r.stdout)
        if m:
            w, h = int(m.group(1)), int(m.group(2))
            _detected_res = (w, h)
            _scale_x = w / REFERENCE_RES[0]
            _scale_y = h / REFERENCE_RES[1]
            print(f"  [坐标缩放] 检测到分辨率 {w}x{h}, 缩放比 X={_scale_x:.3f} Y={_scale_y:.3f}")
        else:
            print(f"  [坐标缩放] 无法解析分辨率: {r.stdout}")
    except Exception as e:
        print(f"  [坐标缩放] 检测失败: {e}, 使用默认1:1")

def sc(x: int, y: int) -> tuple:
    """缩放单个坐标点"""
    return (int(x * _scale_x + 0.5), int(y * _scale_y + 0.5))

def sc_xy(xy: tuple) -> tuple:
    """缩放一个 (x, y) 元组"""
    return sc(xy[0], xy[1])

def scale_all_coords():
    """原地缩放所有坐标字典，使之适配当前手机分辨率"""
    global MODULE_COORDS, TABS, SETTINGS_ICON, AD_CLOSE, MODULE_PATTERNS
    if _scale_x == 1.0 and _scale_y == 1.0:
        return
    MODULE_COORDS = {k: sc_xy(v) for k, v in MODULE_COORDS.items()}
    TABS = {k: sc_xy(v) for k, v in TABS.items()}
    SETTINGS_ICON = sc_xy(SETTINGS_ICON)
    AD_CLOSE = sc_xy(AD_CLOSE)
    # ★ 修复：MODULE_PATTERNS 中的坐标也需要缩放
    MODULE_PATTERNS = {k: (v[0], v[1], sc_xy(v[2]) if v[2] else None)
                       for k, v in MODULE_PATTERNS.items()}
    print(f"  [坐标缩放] 所有坐标已缩放 ({REFERENCE_RES[0]}x{REFERENCE_RES[1]} -> {_detected_res})")


def get_adb():
    """获取 ADBController 实例 (截图保存到 screenshots/)"""
    config = load_config()
    # 用相对路径避免 adb pull 在 Windows 上拼接出错
    return ADBController(serial=config.device.serial, screenshot_dir="screenshots")




def _detect_content_dimension(module: str = ""):
    """自动补充内容检查（LLM 批量知识性审查）

    ★ 修复：必须按当前模块匹配脚本，否则无脚本模块会拿任意最新脚本错配比对
      - module 为空 → 不比对（无模块上下文）
      - 该模块在 docx_map 中无匹配脚本 → 跳过（仅基础完整性）
      - 有匹配脚本 → 用该模块的脚本比对
    """
    try:
        if not module:
            return
        # ★ 该模块的匹配脚本（多模块检测时 docx_map 已构建；单模块/快速检查无映射则跳过）
        _docx = (_GLOBAL_DOCX_MAP or {}).get(module, "")
        if not _docx:
            return
        docx_path = PROJECT_ROOT / "uploads" / _docx
        if not docx_path.exists():
            return
        # ★ 按当前模块单元过滤脚本（用户要求：脚本含多单元，只比对所选单元）
        #   从 _inspection_state 读当前单元（多模块检测启动时记录）
        # ★ 从 _inspection_state 读当前单元（多模块检测启动时记录）——防御"NONE"/非数字
        try:
            _unit_cur = int(str(_inspection_state.get("unit", 0) or 0).split("-")[0])
        except Exception:
            _unit_cur = 0
        from src.review_agent import ReviewAgent, ReviewConfig
        cfg = ReviewConfig(docx_path=str(docx_path), unit=_unit_cur, screenshot_dir="", verbose=False)
        agent = ReviewAgent(cfg)
        if not agent.script_questions: return
        log_msg(f"📄 {module} 自动补充内容检查 (批量, {_docx}, 单元{_unit_cur or '全部'}, {len(agent.script_questions)}题)", "info")
        batch = []; qid_map = {}
        # ★ 听力专项多子模块：只审查当前 stage 的题（用户反馈：错题日志错误记录
        #   练习/测试的题——之前不过滤，练习的题也拿测试脚本去比对→脚本错配）
        _cur_mod, _cur_stage = "", ""
        try:
            from common.logger import get_current_module
            _cur_mod, _cur_stage = get_current_module()
        except Exception:
            pass
        for qid, q in list(_inspection_state.get("questions", {}).items()):
            # ★ stage 过滤：仅审查当前回调对应的 stage 题目
            _q_stage = (q or {}).get("stage", "") or ""
            if _cur_stage and _q_stage and _q_stage != _cur_stage:
                continue
            idx = (q or {}).get("idx", 0)
            if not idx: continue
            # ★ 口语训练题卡对齐（2026-08-24 修复"脚本审查乱报"根因二）：
            #   口语训练模块记录 idx=小题号(1-5)+big(1-4)，而脚本 global_idx 是全局题号(1-20)。
            #   ⚠ 不能先按 global_idx 匹配：第3大题·第2小题(idx=2)会错配到脚本 Q02（朗读单词），
            #   必须优先按 (unit, big, stage_idx) 精确定位；无 big 的题卡（听力专项等）再走 global_idx。
            _q_big = (q or {}).get("big") or 0
            _q_no = (q or {}).get("module_qno") or 0
            sqs = []
            if _q_big and _q_no:
                sqs = [s for s in agent.script_questions
                       if s.unit == _unit_cur and s.big == _q_big and s.stage_idx == _q_no]
            if not sqs:
                sqs = [s for s in agent.script_questions if s.global_idx == idx]
            if not sqs: continue
            sq = sqs[0]
            stem = (getattr(sq, "stem", "") or "").strip()
            opts = "; ".join([o for o in (sq.options or []) if o]) or "(无)"
            ans = (getattr(sq, "answer", "") or "").strip() or "(无)"
            typ = (getattr(sq, "type_2", "") or "").strip() or "未知"
            # ★ 口语训练：定位描述用"第N大题·第M小题"（与 App 证据卡一致），非口语题保持 Q{idx}
            _pos = f"第{sq.big}大题·第{sq.stage_idx}小题" if getattr(sq, "big", 0) else f"Q{idx}"
            _caption = (getattr(sq, "caption", "") or "").strip()
            # ★ 唯一序号：口语训练多题共用 idx=小题号(1-5)，qid_map 不能用 idx 当 key（会互相覆盖）
            _batch_idx = len(batch) + 1
            batch.append(f"{_batch_idx}. {_pos}: 题型={typ}, 大题说明={(_caption or '(无)')[:45]}, 题干={stem or '(无)'}, 答案={ans}")
            qid_map[_batch_idx] = qid
        if not batch: return
        rules = ("规则:对每题做内容检查-1.知识性(PASS/REVIEW/REJECT)必须有依据才判错;"
                 "2.语言质量(拼写/语法/选项/搭配/语义)。"
                 '输出JSON数组:[{"idx":1,"verdict":"PASS/REVIEW/REJECT","basis":"判断依据","lang":["问题描述"]},...]'
                 "其中 idx 为上面列表中每题的序号(1,2,3...)。")
        prompt_text = f"批量内容审查:共{len(batch)}题\n" + "\n".join(batch) + "\n\n" + rules + "\n只输出JSON数组。"
        try:
            raw = (agent.llm.ask(prompt_text) or "").strip()
            if not raw:
                raise ValueError("LLM 返回空内容（可能请求超时/密钥失效）")
            # ★ 健壮解析：LLM 可能返回 markdown 围栏 / 前后说明文字 / 部分对象
            #   依次尝试：①整串JSON ②去markdown围栏 ③取 [] 区间 ④逐行对象
            data = None
            # ① 整串
            try:
                data = json.loads(raw)
            except Exception:
                pass
            # ② 去 markdown ```json ... ```
            if data is None:
                _m = re.search(r"```(?:json)?\s*(.*?)```", raw, re.S)
                if _m:
                    try:
                        data = json.loads(_m.group(1).strip())
                    except Exception:
                        pass
            # ③ 取 [ 到 ] 区间
            if data is None:
                _l, _r = raw.find("["), raw.rfind("]")
                if _l >= 0 and _r > _l:
                    try:
                        data = json.loads(raw[_l:_r+1])
                    except Exception:
                        pass
            # ④ 逐行找 { "idx":... } 对象
            if data is None:
                _objs = re.findall(r"\{[^{}]*\"idx\"\s*:\s*\d+[^{}]*\}", raw)
                if _objs:
                    try:
                        data = [json.loads(o) for o in _objs]
                    except Exception:
                        pass
            if data is None:
                raise ValueError(f"LLM 返回内容无法解析为 JSON: {raw[:80]!r}…")
            count = 0
            # ★ 查重：idx+stage 已存在时复用 qid（避免 LLM 审查与 engine.py 重复记录同一题）
            for item in data:
                qi = item.get("idx", 0)
                qid = qid_map.get(qi)
                if not qid or not _inspection_state["questions"].get(qid):
                    # 兜底：按 stage+idx 找已有 qid（避免 LLM qid 失效时丢失写入）
                    _existing_qid = None
                    for _k, _v in list(_inspection_state["questions"].items()):
                        if (_v.get("idx") == qi
                                and (_v.get("stage") or "") == (_cur_stage or "")):
                            _existing_qid = _k
                            break
                    if _existing_qid:
                        qid = _existing_qid
                    else:
                        continue
                q_obj = _inspection_state["questions"][qid]
                v = str(item.get("verdict","REVIEW")).upper()
                ls = item.get("lang") or []
                # ★ 修复矛盾：ai_content(是否通过) 与 content_reason(解释词) 必须一致
                #   PASS 但带 lang 语言问题 → 整体判不通过，reason 显示"有语言问题"而非"通过"
                has_lang = len(ls) > 0
                # ★ REVIEW=待审（信息不足无法判断，如选项内容为空）→ 计未检(None)而非不通过
                #   用户实测：待审被当不通过 → 明明其他维度过却综合判定不通过
                if v == "REVIEW":
                    passed = None
                    vv = "待审"
                else:
                    passed = (v == "PASS" and not has_lang)
                    if passed:
                        vv = "通过"
                    elif v == "REJECT":
                        vv = "不通过"
                    elif v == "PASS":
                        vv = "通过但有问题"  # PASS 但带语言问题 → 不通过（与 ai_content=False 一致）
                    else:
                        vv = "待审"
                b = (item.get("basis") or "")[:200]
                d = f"{vv}: {b}"
                for li in ls[:3]:
                    d += f"; [{li}]"
                q_obj["ai_content"] = passed
                q_obj["content_reason"] = d[:300]
                count += 1
            log_msg(f"🧪 内容检查已补充: {count}题 ({'通过' if count > 0 else ''})", "success")
        except Exception as e:
            _err = str(e)
            # ★ 友好识别常见 LLM 错误（配额/密钥/网络），给出可操作提示而非笼统报错
            if "403" in _err or "Free quota" in _err or "AllocationQuota" in _err:
                log_msg("⚠ 批量内容检查失败: 阿里云百炼免费额度已用完(403 Forbidden)。请在百炼控制台充值或关闭『仅使用免费额度』模式后重试", "warning")
            elif "401" in _err or "Invalid" in _err or "Authentication" in _err:
                log_msg("⚠ 批量内容检查失败: LLM API 密钥无效或过期(401)。请检查 .env 中的 LLM_API_KEY", "warning")
            elif "timed out" in _err.lower() or "Timeout" in _err:
                log_msg("⚠ 批量内容检查失败: LLM 请求超时。可稍后重试或分批检查", "warning")
            elif "返回空内容" in _err:
                log_msg(f"⚠ 批量内容检查失败: {_err[:100]}", "warning")
            else:
                log_msg(f"⚠ 批量内容检查失败: {_err[:80]}", "warning")
    except Exception as e:
        log_msg(f"⚠ 内容补充异常: {e}", "warning")




def log_msg(msg: str, level: str = "info", evidence: list = None):
    """添加日志, 可选附带结构化证据(审查差异高亮展示)"""
    entry = {
        "time": datetime.now().strftime("%H:%M:%S"),
        "msg": msg,
        "level": level,
    }
    if evidence:
        entry["evidence"] = evidence
    task_status["log"].append(entry)

    if any(k in msg for k in ("完成",)):
        # ★ 传当前模块：_detect_content_dimension 只在该模块有匹配脚本时比对
        _cm_cur = ""
        try:
            from common.logger import get_current_module as _gcm
            _cm_cur, _ = _gcm()
        except Exception:
            pass
        threading.Thread(target=_detect_content_dimension, args=(_cm_cur,), daemon=True).start()
    try:
        print(f"[{entry['time']}] [{level}] {msg}")
    except UnicodeEncodeError:
        # Windows GBK 控制台无法编码 emoji(✅❌ 等)，降级输出
        safe_msg = msg.encode("gbk", errors="replace").decode("gbk")
        print(f"[{entry['time']}] [{level}] {safe_msg}")

    # ★ 每题界面级完整性检查证据 → 同步写入审查结果区（多模块检测也能看到 AI 判断）
    #   触发条件：消息含"第N题 检查"且带 evidence（来自 engine.py _collect_ui_evidence）
    if evidence and isinstance(evidence, list) and evidence:
        try:
            # ★ 修复：兼容实际消息格式——"第1题 检查"(可能带前导空格) 与
            #   "第1题 完整性检查"(测试循环) 都能命中
            #   ★ 口语训练："第2大题·第3小题 完整性检查"（大题·小题精确定位）
            #   ★ 2026-08-25 修复：口语训练实际消息是
            #     "第1大题·第2小题（句型提示：xxx）完整性检查"——"小题"和"检查"之间
            #     夹着（题干内容）！旧正则要求紧跟导致匹配失败 → 界面证据从未写入，
            #     前端只有脚本审查结果（用户实测"没有实际测试记录，全是脚本审查"）
            m = re.match(
                r"\s*(?:第(\d+)大题·)?第(\d+)(?:题|小题)\s*"
                r"(?:（[^）]*）)?\s*(?:完整性)?\s*检查", msg)
            if m:
                _big = int(m.group(1)) if m.group(1) else None
                qidx = int(m.group(2))
                _record_module_evidence(qidx, msg, evidence, big=_big)
        except Exception:
            pass

    # ★ 答错题目截图 → 同步到审查结果区（前端「最近截图」展示）
    #   触发条件：消息含"第N题 答错截图"且 evidence 带 type="wrong_shot" + screenshot 文件名
    if evidence and isinstance(evidence, list) and evidence:
        try:
            m = re.match(r"\s*(?:第(\d+)大题·)?第(\d+)(?:题|小题)\s*答错截图", msg)
            if m:
                qidx = int(m.group(2))
                shot = ""
                for e in evidence:
                    if e.get("type") == "wrong_shot":
                        shot = e.get("screenshot", "") or ""
                        break
                if shot:
                    key = f"auto-Q{qidx:03d}"
                    qs = _inspection_state.setdefault("questions", {})
                    # ★ 模块上下文（错题定位用）
                    try:
                        from common.logger import get_current_module
                        _cm, _cs = get_current_module()
                    except Exception:
                        _cm, _cs = "", ""
                    if not _cm:
                        _cm = _inspection_state.get("module", "")
                    if key in qs:
                        qs[key]["screenshot"] = shot
                    else:
                        # 该题无完整性检查记录（答题循环未发证据）→ 建一条错题记录
                        qs[key] = {
                            "idx": qidx,
                            "total": len(qs) + 1,
                            "module": _cm,
                            "stage": _cs,
                            "module_qno": qidx,
                            "question_type": "错题截图",
                            "screenshot": shot,
                            "progress": f"Q{qidx}",
                            "ai_stem": None, "ai_content": None, "ai_image": None,
                            "ai_answer": None, "ai_audio": None, "ai_post_error": None,
                            "overall_passed": None, "overall_score": 0.0,
                            "stem_reason": "", "content_reason": "", "image_reason": "",
                            "answer_reason": "", "audio_reason": "", "post_error_reason": "",
                            "stem": f"第{qidx}题（答错，已截图）", "options": "",
                            "script_answer": "", "note": "",
                        }
                    try:
                        _save_inspection_state()
                    except Exception:
                        pass
        except Exception:
            pass

# 注入共享日志通道：让 scheduler 和模块内部的流程日志也能送到前端
sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
from common.logger import set_log_callback, set_stop_check
set_log_callback(log_msg)
# 注入停止检查：让 engine/模块答题循环能感知前端"停止"请求
set_stop_check(lambda: _is_stop_requested())


# ★ 设备连接缓存：_connect_device 每次可能 5-20s（u2 初始化），错题自动截图
#   每题调用会严重拖慢 → 缓存成功连接，断连后自动重连
_DEV_CACHE = {"d": None}

def _connect_device():
    """连接当前选中的设备（带就绪检查 + 短超时 + 连接缓存）

    - 设置 u2 全局超时（默认 300s，设备断连时每个操作挂 5 分钟 = 看起来卡死）
    - 检查设备在线，未连接/离线 → 抛 RuntimeError（上层记 error 日志）
    - 整个连接过程硬超时 20s：u2.connect 初始化（装ATX/起server）可能很慢
    """
    import threading
    # ★ 缓存复用：已连接且设备在线 → 直接返回（错题截图每秒调用时避免 5-20s 重连）
    try:
        _cached = _DEV_CACHE.get("d")
        if _cached is not None:
            try:
                if _cached.info:
                    return _cached
            except Exception:
                _DEV_CACHE["d"] = None  # 连接失效 → 重连
    except Exception:
        pass
    result = {}
    def _do_connect():
        try:
            import uiautomator2 as u2
            u2.HTTP_TIMEOUT = 15
            u2.WAIT_FOR_DEVICE_TIMEOUT = 10
            try:
                from common.device import device_ok
                if not device_ok():
                    result["err"] = "设备未连接或离线，请先连接设备"
                    return
            except ImportError:
                pass
            d = u2.connect()
            if not d.info:
                result["err"] = "设备连接失败（u2 无响应）"
                return
            result["d"] = d
            _DEV_CACHE["d"] = d  # ★ 缓存成功连接
        except SystemExit:
            log_msg("⏹ 任务已被立即停止", "warning")
        except Exception as e:
            result["err"] = f"设备连接异常: {e}"
    t = threading.Thread(target=_do_connect, daemon=True)
    t.start()
    t.join(timeout=20)
    if t.is_alive():
        raise RuntimeError("设备连接超时（20s），请检查设备 uiautomator2 服务是否正常")
    if "err" in result:
        raise RuntimeError(result["err"])
    return result["d"]


def clear_status():
    """清空状态"""
    task_status["running"] = False
    task_status["current_task"] = ""
    task_status["progress"] = []
    task_status["start_time"] = ""
    task_status["end_time"] = ""
    task_status["log"] = []


def set_running(task_name: str):
    """设置为运行中"""
    clear_status()
    task_status["running"] = True
    task_status["current_task"] = task_name
    task_status["start_time"] = datetime.now().strftime("%H:%M:%S")


def set_done():
    """设置为完成"""
    task_status["running"] = False
    task_status["end_time"] = datetime.now().strftime("%H:%M:%S")


def _cleanup_task_state():
    """任务线程统一收尾：清 running 状态 + 重置停止标志（供正常/停止/异常所有路径调用）"""
    global _STOP_REQUESTED
    with _STOP_LOCK:
        _STOP_REQUESTED = False
    set_done()


def update_progress(step: int, total: int, msg: str):
    """更新进度"""
    entry = {"step": step, "total": total, "msg": msg}
    task_status["progress"].append(entry)


# ============================================================
# 后台任务函数
# ============================================================

def run_login_task():
    """后台执行登录"""
    try:
        set_running("login")
        adb = get_adb()
        config = load_config()

        log_msg("启动英语宝APP...")
        update_progress(1, 6, "启动APP")
        adb.launch_app(config.app.package)
        time.sleep(5)

        log_msg("关闭启动广告...")
        update_progress(2, 6, "关闭广告")
        adb.tap(*AD_CLOSE)
        time.sleep(2)

        log_msg("勾选协议...")
        update_progress(3, 6, "勾选协议")
        adb.click_element(text="我已阅读并同意", exact=False)
        time.sleep(1)

        log_msg("点击登录...")
        update_progress(4, 6, "点击登录")
        # 先按 exact 尝试
        if not adb.click_element(text="登录", exact=True):
            adb.tap(*sc(540, 1232))  # fallback: 登录按钮已知坐标
        time.sleep(3)

        log_msg("处理协议弹窗...")
        update_progress(5, 6, "协议弹窗")
        if adb.wait_for_element(text="同意", timeout=3):
            adb.click_element(resource_id="agree_tv", exact=False)
            time.sleep(3)

        log_msg("关闭广告弹窗...")
        adb.tap(*AD_CLOSE)
        time.sleep(2)

        adb.screenshot("login_done.png")
        log_msg("✅ 登录完成！", "success")
        update_progress(6, 6, "登录成功")

    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        log_msg(f"❌ 登录失败: {e}", "error")
    finally:
        set_done()


def run_version_detect_task():
    """后台检测可用版本"""
    try:
        set_running("version_detect")
        config = load_config()
        adb = get_adb()

        log_msg("导航到版本选择页...")
        update_progress(1, 5, "进入'我'页")
        adb.tap(*TABS["我"])
        time.sleep(3)

        update_progress(2, 5, "进入设置")
        adb.tap(*SETTINGS_ICON)
        time.sleep(2)

        update_progress(3, 5, "进入个人信息")
        adb.click_element(text="个人信息", exact=True)
        time.sleep(2)

        update_progress(4, 5, "点击英语所学教材版本")
        adb.click_element(text="英语所学教材版本", exact=True)
        time.sleep(2)

        log_msg("截图版本选择页...")
        adb.screenshot("versions_page.png")
        update_progress(5, 5, "版本页已打开")

        # dump UI 获取版本列表
        elements = adb.dump_ui()
        versions = []
        for elem in elements:
            if elem.text and ("版" in elem.text or "年级" in elem.text):
                versions.append({
                    "text": elem.text,
                    "center": list(elem.center),
                    "bounds": list(elem.bounds),
                })

        # 保存到临时文件供 API 读取
        versions_file = PROJECT_ROOT / "outputs" / "web" / "versions.json"
        with open(versions_file, "w", encoding="utf-8") as f:
            json.dump(versions, f, ensure_ascii=False, indent=2)

        log_msg(f"✅ 检测到 {len(versions)} 个版本元素", "success")
        log_msg(f"  可用版本: {[v['text'] for v in versions]}")

    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        log_msg(f"❌ 版本检测失败: {e}", "error")
    finally:
        set_done()


def run_full_task(version: str, grade: str, modules: list):
    """一键全流程：登录 → 切换版本 → 选年级 → 测模块 (加速版)"""
    try:
        set_running("full")
        config = load_config()
        adb = get_adb()

        total = 8 + len(modules) * 3
        cur = 0

        # 1. 启动APP
        cur += 1
        log_msg(f"{cur}/{total} 启动APP")
        adb.launch_app(config.app.package)
        time.sleep(4)

        # 2. 关广告
        cur += 1
        log_msg(f"{cur}/{total} 关广告")
        adb.tap(*AD_CLOSE)
        time.sleep(1)

        # 3. 登录 (仅当需要时)
        cur += 1
        log_msg(f"{cur}/{total} 检查登录状态")
        elements = adb.dump_ui()
        needs_login = any('登录' in (e.text or '') for e in elements)
        if needs_login:
            log_msg(f"  需要登录")
            adb.click_element(text="我已阅读并同意", exact=False)
            time.sleep(0.5)
            adb.click_element(text="登录", exact=True)
            time.sleep(2)
            if adb.wait_for_element(text="同意", timeout=3):
                adb.tap(*sc(540, 1550))
                time.sleep(1.5)
        else:
            log_msg(f"  已登录, 跳过", "success")
        adb.tap(*AD_CLOSE)
        time.sleep(1)

        # 4. 切版本 (仅当需要时)
        cur += 1
        log_msg(f"{cur}/{total} 切换版本: {version}")
        # 先检查主页面上的版本文字
        elements = adb.dump_ui()
        on_correct_version = False
        for e in elements:
            t = (e.text or "").replace('（', '(').replace('）', ')')
            v = version.replace('（', '(').replace('）', ')')
            if v in t:
                on_correct_version = True
                break
        if on_correct_version:
            log_msg(f"  已是 {version}, 跳过切换", "success")
        else:
            adb.tap(*TABS["我"]); time.sleep(2)
            adb.tap(*SETTINGS_ICON); time.sleep(1.5)
            adb.click_element(text="个人信息", exact=True); time.sleep(1.5)
            adb.click_element(text="英语所学教材版本", exact=True); time.sleep(1.5)
            adb.click_element(text=version, exact=True); time.sleep(1.5)

        # 5. 回首页
        cur += 1
        log_msg(f"{cur}/{total} 返回首页")
        for _ in range(3):
            adb.press_back(); time.sleep(1)
        adb.tap(*TABS["英语"]); time.sleep(3)
        adb.tap(*AD_CLOSE); time.sleep(1)

        # 6. 选年级
        if grade:
            cur += 1
            log_msg(f"{cur}/{total} 选择年级: {grade}")
            adb.tap(*sc(346, 275)); time.sleep(1.5)

            found = False
            for attempt in range(4):
                elements = adb.dump_ui()
                for e in elements:
                    if grade in (e.text or ""):
                        adb.tap(e.center[0], e.center[1])
                        log_msg(f"  ✅ 已选 {grade}", "success")
                        found = True
                        break
                if found: break
                adb.swipe(*sc(540, 1700), *sc(540, 900), 300); time.sleep(1)

            # 关闭弹窗
            adb.press_back(); time.sleep(1)
            adb.press_back(); time.sleep(1)
            adb.tap(*TABS["英语"]); time.sleep(3)
            adb.tap(*AD_CLOSE); time.sleep(1)

        # 7. 稳定
        cur += 1
        log_msg(f"{cur}/{total} 页面就绪")
        time.sleep(1)

        # 8. 测试模块
        for i, module in enumerate(modules):
            if module not in MODULE_COORDS:
                log_msg(f"⚠ 跳过: {module}", "warning")
                continue

            cx, cy = MODULE_COORDS[module]
            
            if module in DEEP_MODULES:
                log_msg(f"  ⏬ 滚动到 {module}")
                found_deep = False
                for scroll_step in range(5):
                    adb.swipe(*sc(200, 1600), *sc(200, 1200), 400)
                    time.sleep(1)
                    elements = adb.dump_ui()
                    for e in elements:
                        if e.text and e.text.strip() == module:
                            cx, cy = e.center
                            log_msg(f"    滚{scroll_step+1}次找到 {module}", "success")
                            found_deep = True; break
                    if found_deep: break
                    if not any('专项突破' in (e.text or '') for e in elements):
                        break
                if not found_deep:
                    log_msg(f"    ⚠ 未找到 {module}", "warning")
                    continue

            cur += 1
            log_msg(f"[{i+1}/{len(modules)}] 进入: {module} ({cur}/{total})")
            adb.tap(cx, cy); time.sleep(1.5)

            cur += 1
            safe_name = f"mod_{i+1:02d}.png"
            adb.screenshot(safe_name)
            log_msg(f"  ✅ {module} 截图 ({cur}/{total}) → {safe_name}", "success")

            cur += 1
            adb.press_back(); time.sleep(1.5)
            adb.tap(*AD_CLOSE); time.sleep(1)

        log_msg(f"✅ 完成! {version} {grade} {len(modules)}模块", "success")
        adb.screenshot("99_complete.png")

    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        log_msg(f"❌ 失败: {e}", "error")
        import traceback
        traceback.print_exc()
    finally:
        set_done()


def run_grade_scan_task():
    """全版本年级扫描（调用 grade_scanner.py）"""
    try:
        set_running("grade_scan")
        log_msg("开始全版本年级扫描...")
        scanner = PROJECT_ROOT / "grade_scanner.py"
        r = sp.run(
            [sys.executable, str(scanner)],
            capture_output=True, text=True, timeout=600,
            cwd=str(PROJECT_ROOT),
        )
        if r.returncode == 0:
            log_msg("✅ 全版本年级扫描完成", "success")
        else:
            log_msg(f"⚠ 扫描部分失败: {r.stderr[-200:]}", "warning")
    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        log_msg(f"❌ 扫描失败: {e}", "error")
    finally:
        set_done()
    try:
        set_running("full")
        config = load_config()
        adb = get_adb()

        total = 8 + len(modules) * 3
        cur = 0

        # 1. 启动APP
        cur += 1
        log_msg(f"1/{total} 启动APP")
        update_progress(cur, total, "启动APP")
        adb.launch_app(config.app.package)
        time.sleep(5)

        # 2. 关闭广告
        cur += 1
        log_msg(f"2/{total} 关闭广告")
        update_progress(cur, total, "关闭广告")
        adb.tap(*AD_CLOSE)
        time.sleep(2)

        # 3. 登录流程
        cur += 1
        log_msg(f"3/{total} 自动登录")
        update_progress(cur, total, "自动登录")
        # 勾协议
        adb.click_element(text="我已阅读并同意", exact=False)
        time.sleep(1)
        # 点登录
        adb.click_element(text="登录", exact=True)
        time.sleep(3)
        # 处理弹窗
        if adb.wait_for_element(text="同意", timeout=3):
            adb.click_element(resource_id="agree_tv", exact=False)
            time.sleep(3)
        adb.tap(*AD_CLOSE)
        time.sleep(2)
        adb.screenshot("01_login_done.png")

        # 4. 切到版本选择页 → 选版本
        cur += 1
        log_msg(f"4/{total} 切换版本: {version}")
        update_progress(cur, total, f"切换版本: {version}")
        adb.tap(*TABS["我"])
        time.sleep(3)
        adb.tap(*SETTINGS_ICON)
        time.sleep(2)
        adb.click_element(text="个人信息", exact=True)
        time.sleep(2)
        adb.click_element(text="英语所学教材版本", exact=True)
        time.sleep(2)
        adb.click_element(text=version, exact=True)
        time.sleep(3)

        # 5. 返回主页面
        cur += 1
        log_msg(f"5/{total} 返回主页面")
        update_progress(cur, total, "返回主页面")
        for _ in range(3):
            adb.press_back()
            time.sleep(1)
        adb.tap(*TABS["英语"])
        time.sleep(6)
        adb.tap(*AD_CLOSE)
        time.sleep(2)
        adb.screenshot("02_back_home.png")

        # 6. 选择年级（在主页面顶部点年级切换器）
        if grade:
            cur += 1
            log_msg(f"6/{total} 选择年级: {grade}")
            update_progress(cur, total, f"选择年级: {grade}")
            adb.tap(*sc(346, 275))
            time.sleep(3)

            # 动态查找年级文字（支持滚动）
            found = False
            for attempt in range(4):
                elements = adb.dump_ui()
                for elem in elements:
                    if grade in elem.text:
                        adb.tap(elem.center[0], elem.center[1])
                        log_msg(f"  ✅ 已选 \"{grade}\" at {elem.center}", "success")
                        found = True
                        break
                if found:
                    break
                # 没找到，滚屏
                adb.swipe(*sc(540, 1700), *sc(540, 900), 300)
                time.sleep(2)

            if not found:
                log_msg(f"  ⚠ 未找到年级 \"{grade}\"", "warning")
            time.sleep(3)
            # 关闭弹窗（点左上角）
            adb.tap(*sc(108, 100))
            time.sleep(1)
            adb.screenshot("03_grade_selected.png")

        # 7. 等待页面稳定
        cur += 1
        log_msg(f"{'7' if grade else '6'}/{total} 页面稳定")
        update_progress(cur, total, "页面已就绪")
        time.sleep(3)

        # 8. 逐模块检测
        for i, module in enumerate(modules):
            if module not in MODULE_COORDS:
                log_msg(f"⚠ 跳过未知模块: {module}", "warning")
                continue

            cx, cy = MODULE_COORDS[module]

            # 深度模块需先滚到专项突破底部
            if module in DEEP_MODULES:
                log_msg(f"  ⏬ 滚动到深度模块 {module}")
                for _ in range(2):
                    adb.swipe(*sc(540, 1900), *sc(540, 600), 500)
                    time.sleep(2)
                # 重新读取位置 (滚动后坐标可能变)
                elements = adb.dump_ui()
                for elem in elements:
                    if elem.text and elem.text.strip() == module:
                        cx, cy = elem.center
                        log_msg(f"  📍 找到 {module} at ({cx},{cy})", "success")
                        break

            cur += 1
            log_msg(f"[{i+1}/{len(modules)}] 进入: {module} ({cur}/{total})")
            update_progress(cur, total, f"进入 {module}")
            adb.tap(cx, cy)
            time.sleep(3)

            cur += 1
            adb.screenshot(f"module_{module}.png")
            log_msg(f"  ✅ {module} 截图完成 ({cur}/{total})", "success")
            update_progress(cur, total, f"{module} 截图完成")

            cur += 1
            adb.press_back()
            time.sleep(3)
            adb.tap(*AD_CLOSE)
            time.sleep(2)
            log_msg(f"  返回首页 ({cur}/{total})")
            update_progress(cur, total, "返回首页")

        log_msg(f"✅ 全流程完成！版本={version}, 年级={grade}, 模块={len(modules)}个", "success")
        adb.screenshot("99_complete.png")

    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        log_msg(f"❌ 流程失败: {e}", "error")
        import traceback
        traceback.print_exc()
    finally:
        set_done()


def run_inspect_loop():
    """
    逐题巡检循环 (稳定版)
    
    核心逻辑:
      1. 导航到 AI检测 考试页
      2. 对每题: 截图→点右侧选项(970, y)→点底部按钮×2(检查+下一题)
      3. 遇到问题就停，不卡住
    
    按钮约定:
      - 底部按钮 (540, 2174): 选完选项后点一次→显示答案, 再点一次→跳下一题
      - 选项位置: 右侧选项标签区 (x≈900-1000) 不是图片中央
    """
    try:
        set_running("inspect_loop")
        config = load_config()
        adb = get_adb()

        # 1. 启动 + 导航
        log_msg("启动APP")
        adb.launch_app(config.app.package)
        time.sleep(4)
        adb.tap(*AD_CLOSE)  # 关广告
        time.sleep(0.5)
        adb.tap(*TABS["英语"])  # 英语tab
        time.sleep(3)
        adb.tap(*AD_CLOSE)

        # 2. 滚到单元自检
        log_msg("滚动到单元自检")
        for i in range(4):
            adb.swipe(*sc(200, 1600), *sc(200, 1200), 400)
            time.sleep(0.5)
            elements = adb.dump_ui()
            for e in elements:
                if e.text and '单元自检' in e.text:
                    adb.tap(e.center[0], e.center[1])
                    break
            if any('单元自检' in (e.text or '') for e in elements):
                break
        time.sleep(3)
        adb.tap(*AD_CLOSE)

        # 3. 等加载 + 点AI检测去答题
        log_msg("进入AI检测")
        time.sleep(6)
        elements = adb.dump_ui()
        for e in elements:
            if e.text == '去答题' and e.clickable and e.center[1] < 800:
                adb.tap(e.center[0], e.center[1])
                break
        time.sleep(5)

        # 4. 关规则弹窗 + 开始答题
        adb.tap(*sc(540, 1578))
        time.sleep(0.5)
        adb.tap(*sc(554, 2116))
        time.sleep(10)

        # 5. 逐题巡检
        log_msg("逐题巡检")
        questions = []
        last = 0
        
        for step in range(80):
            if not task_status["running"]:
                break

            elements = adb.dump_ui(retries=2)
            cur = None
            for e in elements:
                m = re.match(r'^(\d+)/(\d+)$', (e.text or "").strip())
                if m:
                    cur = int(m.group(1))
                    break
            
            if not cur:
                log_msg("不在考题页，结束", "warning")
                break
            if cur == last:
                time.sleep(0.3)
                continue
            last = cur

            adb.screenshot(f"q{cur:02d}.png")
            questions.append({"idx": cur, "screenshot": f"q{cur:02d}.png"})
            log_msg(f"Q{cur}", "success")
            update_progress(cur, 40, f"Q{cur}")

            # 点右侧选项或中央（取决于布局）
            for e in elements:
                if e.clickable and 700 < e.bounds[1] < 1700:
                    if e.center[0] > 450:
                        adb.tap(e.center[0], e.bounds[1] + 30)
                    else:
                        adb.tap((e.bounds[0] + e.bounds[2]) // 2, (e.bounds[1] + e.bounds[3]) // 2)
                    break
            time.sleep(1)

            # 连点两次底部按钮 (检查→下一题)
            adb.tap(*sc(540, 2174))
            time.sleep(1.5)
            adb.tap(*sc(540, 2174))
            time.sleep(1.5)

            if cur >= 40:
                log_msg("✅ 40题完成!", "success")
                break

        # 6. 保存报告
        out = PROJECT_ROOT / "outputs" / "questions" / "inspect_report.json"
        with open(out, "w", encoding="utf-8") as f:
            json.dump(questions, f, ensure_ascii=False, indent=2)
        log_msg(f"✅ 报告: {out.name} ({len(questions)}题)", "success")

    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        log_msg(f"❌ 异常: {e}", "error")
        import traceback
        traceback.print_exc()
    finally:
        set_done()



def run_quick_inspect_task(docx_file: str = "", unit: int = 0):
    """⚡ 快速检查：从手机当前页面直接开始逐题巡检（跳过启动/关广告/登录/导航）

    前提: 用户已手动将手机调到题目页 (如 听力专项-基础巩固-第1题)
    流程: 截图 → AI六维审查(若提供脚本) → 写入巡检状态 → 点选项 → 点底部按钮x2 → 循环
    坐标自动缩放 (适配不同分辨率手机)
    """
    try:
        set_running("quick_inspect")
        config = load_config()
        adb = get_adb()

        # 清空旧巡检状态
        _inspection_state["questions"] = {}
        _inspection_state["workflow_steps"] = []
        _inspection_state["current_question_idx"] = 0
        _save_inspection_state()

        # 加载脚本 + AI审查agent (可选)
        agent = None
        if docx_file:
            docx_path = str(UPLOAD_DIR / docx_file)
            if Path(docx_path).exists():
                try:
                    from src.review_agent import ReviewAgent, ReviewConfig
                    cfg = ReviewConfig(docx_path=docx_path, unit=int(unit or 0),
                                       screenshot_dir=str(PROJECT_ROOT / "screenshots"),
                                       verbose=False)
                    agent = ReviewAgent(cfg)
                    log_msg(f"📄 脚本已加载: {docx_file} (共{len(agent.script_questions)}题) → AI六维审查开启", "success")
                except Exception as e:
                    log_msg(f"⚠ 脚本加载失败: {e}，降级为仅截图", "warning")
            else:
                log_msg(f"⚠ 脚本不存在: {docx_path}，降级为仅截图", "warning")
        else:
            log_msg("⚡ 未指定脚本，仅截图+推进（选脚本可开启AI六维审查）", "info")

        log_msg("⚡ 快速检查启动（从当前页面开始，跳过导航）", "success")
        time.sleep(1)

        # 1. 尝试检测当前题目进度 (如 1/40)
        elements = adb.dump_ui(retries=2)
        cur = None
        total_q = 0
        for e in elements:
            m = re.match(r'^(\d+)/(\d+)$', (e.text or "").strip())
            if m:
                cur = int(m.group(1))
                total_q = int(m.group(2))
                break

        if cur:
            log_msg(f"✅ 检测到题目进度 {cur}/{total_q or 40}，开始逐题检查", "success")
        else:
            log_msg("⚠ 未检测到题目进度(如 1/40)。若您已在题目页将正常检查，", "warning")
            log_msg("   否则请先手动进入题目页再点「快速检查」", "warning")
            time.sleep(2)

        # 2. 逐题巡检
        last = 0
        q_count = 0
        for step in range(80):
            if not task_status["running"]:
                break

            idx = cur if cur else (last + 1)

            # 截图 (保存到 outputs/web, 前端可预览)
            shot = f"q{idx:02d}.png"
            adb.screenshot(shot)
            log_msg(f"📸 Q{idx} 已截图", "success")
            update_progress(idx, total_q or 40, f"Q{idx}")
            q_count += 1

            # ====== AI 六维审查 (若有脚本) ======
            ai = {
                "ai_stem": None, "ai_content": None, "ai_image": None,
                "ai_answer": None, "ai_audio": None, "ai_post_error": None,
                "overall_passed": None, "overall_score": None,
                "stem_reason": "", "content_reason": "", "image_reason": "",
                "answer_reason": "", "audio_reason": "", "post_error_reason": "",
                "question_type": "快速检查",
                "script_answer": "", "stem": "",
            }
            if agent:
                try:
                    # ★ 修复：截图保存在 screenshots/（adb.screenshot），
                    #   原代码读 outputs/web/ 导致找不到图 → 降级纯文字模式
                    shot_path = str(PROJECT_ROOT / "screenshots" / shot)
                    matching = [q for q in agent.script_questions if q.global_idx == idx]
                    if matching:
                        script_q = matching[0]
                        r = agent._review_one(script_q, shot_path)
                        ai = {
                            "ai_stem": r.stem_check.passed,
                            "ai_content": r.content_check.passed,
                            "ai_image": r.image_check.passed,
                            "ai_answer": r.answer_check.passed,
                            "ai_audio": r.audio_check.passed if r.audio_check is not None else None,
                            "ai_post_error": r.post_error_check.passed if r.post_error_check is not None else None,
                            "overall_passed": r.overall_passed,
                            "overall_score": round(r.overall_score, 2),
                            "stem_reason": r.stem_check.details[0][:100] if r.stem_check.details else "",
                            "content_reason": r.content_check.details[0][:100] if r.content_check.details else "",
                            "image_reason": r.image_check.details[0][:100] if r.image_check.details else "",
                            "answer_reason": r.answer_check.details[0][:100] if r.answer_check.details else "",
                            "audio_reason": r.audio_check.details[0][:100] if r.audio_check and r.audio_check.details else "",
                            "post_error_reason": r.post_error_check.details[0][:100] if r.post_error_check and r.post_error_check.details else "",
                            "question_type": script_q.type_2 or "快速检查",
                            "script_answer": script_q.answer or "",
                            "stem": script_q.stem[:60] if script_q.stem else "",
                        }
                        log_msg(f"  AI: 题干{ai['ai_stem']} 内容{ai['ai_content']} 配图{ai['ai_image']} 作答{ai['ai_answer']} 音频{ai['ai_audio']} 答错后{ai['ai_post_error']} 得分{ai['overall_score']}")
                    else:
                        log_msg(f"  ⚠ 脚本中无 Q{idx}，跳过AI审查", "warning")
                except Exception as e:
                    log_msg(f"  ⚠ AI审查失败 Q{idx}: {e}", "warning")

            # ====== 写入巡检状态 → 前端中栏显示 ======
            qid = f"quick-Q{idx:02d}"
            _inspection_state["questions"][qid] = {
                "idx": idx,
                "total": total_q or 40,
                "question_type": ai["question_type"],
                "screenshot": shot,
                "progress": f"{idx}/{total_q or 40}",
                "ai_stem": ai["ai_stem"], "ai_content": ai["ai_content"],
                "ai_image": ai["ai_image"], "ai_answer": ai["ai_answer"],
                "ai_audio": ai["ai_audio"], "ai_post_error": ai["ai_post_error"],
                "overall_passed": ai["overall_passed"],
                "overall_score": ai["overall_score"],
                "stem_reason": ai["stem_reason"], "content_reason": ai["content_reason"],
                "image_reason": ai["image_reason"], "answer_reason": ai["answer_reason"],
                "audio_reason": ai["audio_reason"], "post_error_reason": ai["post_error_reason"],
                "script_answer": ai["script_answer"],
                "stem": ai["stem"],
                "note": "",
                "human_label": None,
                "human_note": "",
                "timestamp": datetime.now().isoformat(),
            }
            _inspection_state["current_question_idx"] = idx
            _save_inspection_state()

            # 点选项: 优先找UI元素，找不到则点默认右侧选项区(缩放)
            elements = adb.dump_ui(retries=2)
            clicked_option = False
            for e in elements:
                if e.clickable and 700 < e.bounds[1] < 1700:
                    if e.center[0] > 450:
                        adb.tap(e.center[0], e.bounds[1] + 30)
                    else:
                        adb.tap((e.bounds[0] + e.bounds[2]) // 2,
                                (e.bounds[1] + e.bounds[3]) // 2)
                    clicked_option = True
                    break
            if not clicked_option:
                tx, ty = sc(970, 1500)   # 默认点右侧选项区
                adb.tap(tx, ty)
            time.sleep(1)

            # 点底部按钮 x2 (检查答案 → 下一题)
            bx, by = sc(540, 2174)
            adb.tap(bx, by)
            time.sleep(1.5)
            adb.tap(bx, by)
            time.sleep(1.5)

            # 检测进度变化
            elements_after = adb.dump_ui(retries=2)
            new_cur = None
            new_total = 0
            for e in elements_after:
                m = re.match(r'^(\d+)/(\d+)$', (e.text or "").strip())
                if m:
                    new_cur = int(m.group(1))
                    new_total = int(m.group(2))
                    break

            if new_cur and new_cur == last and cur:
                log_msg(f"  ⚠ Q{new_cur} 进度未变，重试一次", "warning")
                time.sleep(2)
                adb.tap(bx, by)
                time.sleep(1.5)
                continue

            if new_cur:
                last = new_cur
                cur = new_cur
                total_q = new_total or total_q
                if new_cur >= (total_q or 40):
                    log_msg(f"✅ {new_cur}题完成!", "success")
                    break
            else:
                last += 1
                log_msg(f"  未检测到进度文字, 继续下一题 (Q{last})", "info")

        # 3. 保存报告
        out = PROJECT_ROOT / "outputs" / "questions" / "quick_report.json"
        out.parent.mkdir(parents=True, exist_ok=True)
        with open(out, "w", encoding="utf-8") as f:
            json.dump({"questions": list(_inspection_state["questions"].values())},
                      f, ensure_ascii=False, indent=2)
        log_msg(f"✅ 快速检查完成: 共 {q_count} 题", "success")

    except Exception as e:
        log_msg(f"❌ 异常: {e}", "error")
        import traceback
        traceback.print_exc()
    finally:
        set_done()


def run_inspect_questions_task():
    """自动化巡检 AI检测 题目：
    1. 启动APP, 关广告, 进英语tab
    2. 滚到专项突破底部 → 单元自检
    3. 点 AI检测的 去答题 (870, 756)
    4. 关规则弹窗 (540, 1578) → 点开始答题 (554, 2116)
    5. 逐题截图+解析, 推进直到没有"下一题"按钮
    """
    report = {"questions": []}
    try:
        set_running("inspect_questions")
        config = load_config()
        adb = get_adb()

        # 1. 启动APP
        log_msg("启动APP...")
        sp.run(['C:/Users/bunana/AppData/Local/Microsoft/WinGet/Packages/Google.PlatformTools_Microsoft.Winget.Source_8wekyb3d8bbwe/platform-tools/adb.exe', '-s', config.device.serial, 'shell', 'am', 'force-stop', 'com.dinoenglish.yyb'])
        time.sleep(2)
        sp.run(['C:/Users/bunana/AppData/Local/Microsoft/WinGet/Packages/Google.PlatformTools_Microsoft.Winget.Source_8wekyb3d8bbwe/platform-tools/adb.exe', '-s', config.device.serial, 'shell', 'am', 'start', '-n', 'com.dinoenglish.yyb/.base.SplashActivity'])
        time.sleep(5)
        adb.tap(*AD_CLOSE)  # 关启动广告
        time.sleep(3)
        adb.tap(*sc(948, 1821))  # 关可能的第二个
        time.sleep(2)

        # 2. 滚到专项突破底部 → 单元自检
        log_msg("滚动到单元自检...")
        adb.tap(*TABS["英语"])  # 英语tab
        time.sleep(5)
        for _ in range(2):
            adb.swipe(*sc(540, 1500), *sc(540, 800), 400)
            time.sleep(2)
        # 小幅调整
        adb.swipe(*sc(540, 1500), *sc(540, 1100), 400)
        time.sleep(2)

        # 动态找 单元自检
        elements = adb.dump_ui()
        for elem in elements:
            if elem.text and '单元自检' in elem.text:
                adb.tap(elem.center[0], elem.center[1])
                log_msg(f"  ✅ 进入单元自检 at {elem.center}", "success")
                break
        time.sleep(4)

        # 3. 点 AI检测的 去答题
        log_msg("点 AI检测 去答题...")
        adb.tap(*sc(870, 756))
        time.sleep(6)

        # 4. 关规则弹窗
        log_msg("关闭训练规则弹窗...")
        adb.tap(*sc(540, 1578))
        time.sleep(2)

        # 5. 点开始答题
        log_msg("点开始答题...")
        adb.tap(*sc(554, 2116))
        time.sleep(8)  # AI生成题目

        # 6. 逐题巡检
        log_msg("开始逐题巡检...")
        out_dir = PROJECT_ROOT / "outputs" / "questions"
        out_dir.mkdir(parents=True, exist_ok=True)

        for q_idx in range(50):  # 最多 50 题
            elements = adb.dump_ui()
            # 找当前进度
            progress = ""
            question_text = ""
            for e in elements:
                t = e.text or ""
                if re.match(r'^\d+/\d+$', t.strip()):
                    progress = t
                if ('听' in t or '看' in t or '选' in t or '读' in t or '写' in t or '听音' in t) and len(t) > 5 and len(t) < 80:
                    if not question_text:
                        question_text = t

            # 截图
            shot_path = out_dir / f"q_{q_idx+1:02d}.png"
            sp.run(['C:/Users/bunana/AppData/Local/Microsoft/WinGet/Packages/Google.PlatformTools_Microsoft.Winget.Source_8wekyb3d8bbwe/platform-tools/adb.exe', '-s', config.device.serial, 'shell', 'screencap', '-p', '/sdcard/_q.png'])
            sp.run(['C:/Users/bunana/AppData/Local/Microsoft/WinGet/Packages/Google.PlatformTools_Microsoft.Winget.Source_8wekyb3d8bbwe/platform-tools/adb.exe', '-s', config.device.serial, 'pull', '/sdcard/_q.png', str(shot_path)])
            sp.run(['C:/Users/bunana/AppData/Local/Microsoft/WinGet/Packages/Google.PlatformTools_Microsoft.Winget.Source_8wekyb3d8bbwe/platform-tools/adb.exe', '-s', config.device.serial, 'shell', 'rm', '/sdcard/_q.png'])

            # 检测题目类型
            question_type = "未知"
            if any('听录音' in e.text or '🔊' in e.text for e in elements):
                question_type = "听力题"
            elif any('A.' in e.text or 'B.' in e.text or 'C.' in e.text for e in elements):
                question_type = "选择题"
            elif any('读' in e.text and '单词' in e.text for e in elements):
                question_type = "朗读题"
            elif any('写' in e.text and ('单词' in e.text or '字母' in e.text) for e in elements):
                question_type = "拼写题"

            has_image = any(e.text in ['A', 'B', 'C', 'D'] and e.clickable for e in elements)

            report["questions"].append({
                "idx": q_idx + 1,
                "progress": progress,
                "text": question_text,
                "type": question_type,
                "screenshot": str(shot_path.name),
            })

            log_msg(f"题 {q_idx+1}: {progress} | {question_type} | {question_text[:40]}", "info")

            # 检查是否有"下一题"按钮
            has_next = any('下一' in e.text or '提交' in e.text for e in elements)
            if not has_next:
                # 直接看截图/UI找下一题按钮
                # 通常右下角有"下一题"按钮
                for e in elements:
                    if e.clickable and e.text:
                        # 猜的"下一题"按钮位置
                        pass

            # 假设"下一题"按钮在右下角 (约 950, 1700~1900)
            adb.tap(*sc(960, 1850))
            time.sleep(2)

            # 检查是否还在题目页 (有 X/40)
            elements_after = adb.dump_ui()
            on_question = any(re.match(r'^\d+/\d+$', (e.text or '').strip()) for e in elements_after)
            if not on_question:
                log_msg(f"  题目已结束 (q_idx={q_idx+1})", "info")
                break

        # 保存报告
        report_path = PROJECT_ROOT / "outputs" / "questions" / "report.json"
        with open(report_path, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

        log_msg(f"✅ 巡检完成: 共 {len(report['questions'])} 题", "success")
        log_msg(f"报告: {report_path}", "info")

    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        log_msg(f"❌ 巡检失败: {e}", "error")
        import traceback
        traceback.print_exc()
    finally:
        set_done()


# ============================================================
# Flask 路由
# ============================================================


@app.after_request
def _no_cache(resp):
    """禁用页面缓存：确保浏览器始终加载最新 HTML/JS"""
    resp.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


@app.route("/")
def index():
    """前端页面"""
    return render_template("index.html")


@app.route("/trace")
def trace_page():
    """错题溯源页面"""
    return render_template("trace.html")


@app.route("/api/status")
def api_status():
    """设备状态 + 任务状态"""
    config = load_config()
    # 优先用动态选择的设备序列号，未选择则用配置文件兜底
    cur_serial = os.environ.get("ANDROID_SERIAL") or config.device.serial
    device_ok = False
    try:
        sys.path.insert(0, str(Path(__file__).parent / "scripts"))
        from common.device import device_ok as _check_ok
        device_ok = _check_ok(cur_serial)
    except Exception:
        pass

    return jsonify({
        "device_connected": device_ok,
        "device_serial": cur_serial,
        "devices": _device_list(),
        "current_version": _get_current_version_from_config(),
        "task_status": task_status,
    })


def _device_list():
    """列出所有已连接设备（供前端下拉选择）"""
    try:
        sys.path.insert(0, str(Path(__file__).parent / "scripts"))
        from common.device import list_devices
        return list_devices()
    except Exception:
        return []


@app.route("/api/devices", methods=["GET"])
def api_devices():
    """获取所有已连接的 adb 设备列表"""
    return jsonify({"devices": _device_list()})


@app.route("/api/device/select", methods=["POST"])
def api_device_select():
    """选择当前设备序列号（写 ANDROID_SERIAL，全局生效）"""
    data = request.get_json() or {}
    serial = (data.get("serial") or "").strip()
    try:
        sys.path.insert(0, str(Path(__file__).parent / "scripts"))
        from common.device import set_device, device_ok
        set_device(serial or None)
        ok = device_ok()
        return jsonify({"status": "ok", "serial": serial or "", "connected": ok})
    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/login", methods=["POST"])
def api_login():
    """触发自动登录"""
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409
    t = threading.Thread(target=run_login_task, daemon=True)
    t.start()
    return jsonify({"status": "started", "task": "login"})


@app.route("/api/errors/summary", methods=["GET"])
def api_errors_summary():
    """返回当前检测中的错题汇总（供前端 renderServe 展示）

    ★ 2026-08-25 统计口径修复：只统计【实际测试的题】（auto-Q* 界面证据题），
      不统计【脚本全量审查的题】（*-脚本-Q*，LLM 把整个 docx 所有题都审一遍，
      很多题根本没在手机上测过）。
      → 分母 = 实际测试的题数；分子 = 其中判定不通过(overall_passed is False)的题数。
      ★ 同时修复 not overall_passed 把 None(未审查) 误计为不通过的问题。

    ★ 2026-08-25 数据源修复：优先读内存 _inspection_state（本次运行的实时数据）。
      旧逻辑只读磁盘 data/inspection_state.json —— 多模块入口启动时只清内存不清磁盘，
      磁盘残留上次运行的题（如五年级20题+4不通过）→ 用户跑三年级未上线0题，
      日志显示0题但前端 summary 却显示"不通过4/20"，前后矛盾。
      内存为空（服务重启等）才读磁盘兜底。
    """
    questions = _inspection_state.get("questions") or {}
    if not questions:
        state_path = PROJECT_ROOT / "data" / "inspection_state.json"
        if not state_path.exists():
            return jsonify({"error": "暂无检测数据"}), 404
        try:
            with open(state_path, "r", encoding="utf-8") as f:
                state = json.load(f)
            questions = state.get("questions", {})
        except Exception as e:
            return jsonify({"error": str(e)}), 500
    try:
        if isinstance(questions, dict):
            # ★ 保留 qid(key) 才能区分 auto-Q(实际测试) / *-脚本-Q*(脚本全量审查)
            q_items = list(questions.items())
        else:
            q_items = [(str(q.get("qid") or q.get("_qid") or f"q{i}"), q)
                       for i, q in enumerate(questions)]

        # ★ 只保留"实际测试"的题：qid 以 auto-Q 开头（界面证据题 = 引擎真正跑到的题）。
        #   排除 *-脚本-Q*（LLM 脚本全量审查：整个 docx 的题都审，含未实际测试的）
        tested = [q for _qid, q in q_items
                  if str(_qid).startswith("auto-Q") or
                  (str(_qid).startswith("Q") and "脚本-Q" not in str(_qid))]

        failed = []
        # ★ 2026-08-25 修复：key 必须与 questions 里的字段名完全一致（ai_stem/ai_content/ai_image/
        #   ai_answer/ai_audio/ai_post_error）。旧代码用 key[:3]（stem→ste、content→con、post_error→pos）
        #   全部对不上 → failed_dims 永远是空数组 → 错题报告不显示"哪些维度不通过"。
        dim_keys = [
            ("题干", "ai_stem"), ("内容", "ai_content"), ("配图", "ai_image"),
            ("作答", "ai_answer"), ("答错后", "ai_post_error"), ("音频", "ai_audio")
        ]
        for q in tested:
            # ★ 修复：只看严格不通过(False)；None(未审查/信息不足)不误计
            if q.get("overall_passed") is False:
                failed_dims = [label for label, key in dim_keys
                               if q.get(key) is False]
                failed.append({
                    "qid": q.get("qid", ""),
                    "idx": q.get("idx", 0),
                    "question_type": q.get("question_type", ""),
                    "overall_score": q.get("overall_score", 0),
                    "failed_dims": failed_dims,
                })

        return jsonify({
            "total_questions": len(tested),     # ★ 分母 = 实际测试题数
            "failed_count": len(failed),        # ★ 分子 = 其中严格不通过数
            "failed_items": failed,
            "has_errors": len(failed) > 0,
            # ★ 附加信息：脚本审查另有 N 题（供前端展示"另有脚本审查错题"）
            "script_review_count": len(q_items) - len(tested),
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/versions/available")
def api_versions_available():
    """可用版本列表 = App 实际展示的5个2024审定版本（真机 2026-08-26 确认），
    + 知识库中可能存在的其他版本（兜底），去重"""
    # ★ 2026-08-26：App 当前只展示这5个版本（带2024审定），preset 改为完整版本名
    #   旧 preset "新湘鲁五上"等拆分形式已停用（与 App 实际版本不一致）
    preset = [
        "湘少版(2024审定)", "湘鲁版(2024审定)", "人教版(PEP)(2024审定)",
        "教科版(2024审定)", "教研版(2024审定)",
    ]
    versions = list(preset)
    try:
        from src.knowledge_base import KnowledgeBase
        kb = KnowledgeBase()
        summary = kb.summary()
        for s in summary:
            key = s.get("key", "")
            # 知识库key形如 "湘少版:五上" / "湘鲁版:六上" → 转成 "湘少版五上"
            label = key.replace(":", "") if ":" in key else key
            if label and label not in versions:
                versions.append(label)
    except Exception:
        pass
    return jsonify({"versions": versions, "preset": preset})


@app.route("/api/versions", methods=["GET", "POST"])
def api_versions():
    """检测可用版本"""
    if request.method == "GET":
        # 返回已缓存的版本列表
        versions_file = PROJECT_ROOT / "outputs" / "web" / "versions.json"
        if versions_file.exists():
            with open(versions_file, "r", encoding="utf-8") as f:
                versions = json.load(f)
            return jsonify({"versions": versions})
        return jsonify({"versions": [], "msg": "尚未检测，请先 POST /api/versions 触发检测"})

    # POST: 触发版本检测
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409
    t = threading.Thread(target=run_version_detect_task, daemon=True)
    t.start()
    return jsonify({"status": "started", "task": "version_detect"})


@app.route("/api/version-grades")
def api_version_grades():
    """返回版本→年级配置表（scan_versions_grades.py 生成的缓存数据，秒回）"""
    # 优先新版配置表 versions_grades.json（版本 → {grades, current}）
    vg_file = PROJECT_ROOT / "outputs" / "web" / "versions_grades.json"
    if vg_file.exists():
        try:
            with open(vg_file, "r", encoding="utf-8") as f:
                data = json.load(f)
            # 兼容两种结构：{table:{...}} 或 直接 {版本:{...}}
            if isinstance(data, dict) and "table" in data and isinstance(data["table"], dict):
                table = data["table"]
            else:
                table = data
            if table:
                return jsonify({"table": table, "source": "table"})
        except Exception:
            pass
    # 回退旧版 all_grades.json
    grades_file = PROJECT_ROOT / "outputs" / "web" / "all_grades.json"
    if grades_file.exists():
        with open(grades_file, "r", encoding="utf-8") as f:
            data = json.load(f)
        return jsonify({"table": data, "source": "all_grades"})
    return jsonify({"error": "尚未生成配置表，请先运行 scripts/scan_versions_grades.py 或扫描年级"}), 404


@app.route("/api/version-grades/current", methods=["POST"])
def api_version_grades_current():
    """★ 实时从 App 读取当前版本的年级列表（前端切换版本后联动年级下拉）

    与 App 实际内容保持一致：打开「切换课本」页 dump 解析，而非写死。
    body: {"version": "湘少版"}（可选；传入且与当前不同 → 先在 App 内切版本再读）
    返回: {"version": "湘少版", "grades": ["五年级上册", ...], "current_grade": "五年级上册"}
    """
    data = request.get_json(silent=True) or {}
    target = (data.get("version") or "").strip()
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409
    try:
        d = _connect_device()
    except Exception as e:
        return jsonify({"error": f"设备未连接: {e}"}), 400
    try:
        import importlib
        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        setup = importlib.import_module("common.setup")

        def _homebar_text():
            """主页顶部版本+年级栏文本（resource-id=switch_textbook_tv，属性顺序不固定）"""
            try:
                _xml = d.dump_hierarchy()
                _m = re.search(r'<node[^>]*resource-id="[^"]*switch_textbook_tv"[^>]*>', _xml)
                if _m:
                    _tm = re.search(r'text="([^"]*)"', _m.group(0))
                    if _tm:
                        return _tm.group(1)
            except Exception:
                pass
            return ""

        # 1) 当前版本（主页顶部栏文本，取"X版"部分）
        homebar = _homebar_text()
        cur_version = ""
        vm = re.search(r'([\u4e00-\u9fa5]+版)', homebar) if homebar else None
        if vm:
            cur_version = vm.group(1)
        # 2) 指定版本且与当前不同 → 先切版本（走"我的"页路径）
        if target and target not in (homebar or ""):
            log_msg(f"🔄 切换版本 → {target}", "step")
            if not setup.switch_version(d, target):
                log_msg("❌ 版本切换失败", "error")
                return jsonify({"error": f"版本切换失败: {target}"}), 500
            time.sleep(1)
        # 3) 打开「切换课本」页实时读取年级列表
        grades = setup.get_grades_from_app(d)
        if not grades:
            return jsonify({"error": "未能从 App 读取年级列表（请确认设备已连接且在主页）"}), 500
        # 4) 当前年级（主页栏文本中提取 X年级上/下册）
        homebar = _homebar_text()
        cur_grade = ""
        gm = re.search(r'([一二三四五六]年级(?:上|下)册)', homebar) if homebar else None
        if gm:
            cur_grade = gm.group(1)
        if not cur_version:
            vm2 = re.search(r'([\u4e00-\u9fa5]+版)', homebar) if homebar else None
            if vm2:
                cur_version = vm2.group(1)
        log_msg(f"✅ 实时读取年级列表: {len(grades)} 个（当前: {cur_grade or '未知'}）", "success")
        return jsonify({"version": target or cur_version, "grades": grades, "current_grade": cur_grade})
    except Exception as e:
        log_msg(f"❌ 读取年级列表异常: {e}", "error")
        return jsonify({"error": str(e)}), 500


@app.route("/api/version-grades/scan", methods=["POST"])
def api_scan_grades():
    """触发全版本年级扫描（慢，约3分钟）"""
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409
    t = threading.Thread(target=run_grade_scan_task, daemon=True)
    t.start()
    return jsonify({"status": "started"})


@app.route("/api/run-full", methods=["POST"])
def api_run_full():
    """一键全流程运行：登录 → 切换版本 → 选年级 → 测模块"""
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409

    data = request.get_json()
    if not data:
        return jsonify({"error": "请提供参数"}), 400

    version = data.get("version", "")
    grade = data.get("grade", "")
    modules = data.get("modules", [])

    if not version:
        return jsonify({"error": "请选择版本"}), 400
    if not modules:
        return jsonify({"error": "请选择至少一个模块"}), 400

    t = threading.Thread(target=run_full_task, args=(version, grade, modules), daemon=True)
    t.start()
    return jsonify({
        "status": "started",
        "task": "full",
        "version": version,
        "grade": grade,
        "modules": modules,
    })


@app.route("/api/inspect-questions", methods=["POST"])
def api_inspect_questions():
    """AI检测单元自检题目巡检：一题一截图+识别"""
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409
    t = threading.Thread(target=run_inspect_questions_task, daemon=True)
    t.start()
    return jsonify({"status": "started", "task": "inspect_questions"})


@app.route("/api/questions/report")
def api_questions_report():
    """获取最近一次巡检报告"""
    report_path = PROJECT_ROOT / "outputs" / "questions" / "report.json"
    if report_path.exists():
        with open(report_path, "r", encoding="utf-8") as f:
            return jsonify(json.load(f))
    return jsonify({"questions": [], "msg": "尚未巡检"})


# ============================================================
# 四步检查 API
# ============================================================


_current_screenshot = ""



# ============================================================
# 听力专项 专属API — 与CLI脚本一致的工作流
# ============================================================

@app.route("/api/log/clear", methods=["POST"])
def api_log_clear():
    """清空操作日志"""
    task_status["log"] = []
    return jsonify({"status": "cleared"})

@app.route("/api/inspect/listening-run", methods=["POST"])
def api_inspect_listening_run():
    """
    听力专项自动巡检 (与CLI脚本 inspect_u6.py 行为一致)
    
    请求参数:
        version: "新湘鲁六上" / "新湘鲁五上" / ...
        unit: 6 (单元号)
        stage: "基础巩固" / "综合进阶" / "难点突破"   (兼容旧版单阶段)
        stages: ["基础巩固","综合进阶",...]            (新版多阶段，优先)
        docx: 脚本文件名 (已在uploads目录的)
    """
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409

    data = request.get_json() or {}
    version_label = data.get("version", "新湘鲁六上")
    unit = data.get("unit", 6)
    stage = data.get("stage", "基础巩固")
    docx_file = data.get("docx", "")
    # ★ 多阶段支持：stages 优先；兼容旧版单 stage
    stages = data.get("stages") or []
    if isinstance(stages, str):
        stages = [s.strip() for s in stages.split(",") if s.strip()]
    if not stages:
        stages = [stage] if stage else ["基础巩固"]

    # 新巡检自动清空日志和检查状态（多阶段只清一次，后面各阶段结果累积）
    task_status["log"] = []
    _inspection_state["questions"] = {}
    _inspection_state["workflow_steps"] = []
    _inspection_state["current_question_idx"] = 0
    _save_inspection_state()

    # 启动线程
    t = threading.Thread(
        target=_run_listening_inspect_stages,
        args=(version_label, unit, stages, docx_file),
        daemon=True,
    )
    t.start()
    return jsonify({
        "status": "started",
        "task": "listening_inspect",
        "params": {"version": version_label, "unit": unit, "stages": stages},
    })


def _run_listening_inspect_stages(version_label: str, unit: int, stages: list, docx_file: str):
    """多阶段巡检：依次跑每个子模块，结果按 stage 前缀累积到六维面板。
    最后一阶段才 set_done（避免中间阶段提前让前端显示"完成"）。"""
    n = len(stages)
    for i, st in enumerate(stages):
        log_msg(f"📌 巡检阶段 [{i+1}/{n}]: {st}", "step")
        try:
            run_listening_inspect(version_label, unit, st, docx_file,
                                  keep_running=(i < n - 1))
        except SystemExit:
            log_msg("⏹ 任务已被立即停止", "warning")
            break
        except Exception as e:
            log_msg(f"⚠ 阶段 {st} 巡检异常: {e}", "error")
        if _is_stop_requested():
            break
    set_done()
    log_msg(f"✅ 多阶段巡检结束: {'、'.join(stages)}", "success")


def run_listening_inspect(version_label: str, unit: int, stage: str, docx_file: str,
                          keep_running: bool = False):
    """
    通用模块巡检主循环
    
    自动从文件名识别模块(巧记单词/听力专项/知识过关...),导航到对应区域,逐题检查
    """
    def record_step(step, status, detail=""):
        """记录流程步骤到后端 (用urllib避免依赖requests)"""
        try:
            import urllib.request, json
            data = json.dumps({"step": step, "status": status, "detail": detail}).encode()
            req = urllib.request.Request("http://127.0.0.1:5000/api/inspect/workflow-step",
                                        data=data, headers={"Content-Type": "application/json"})
            urllib.request.urlopen(req, timeout=2)
        except Exception:
            pass

    def record_q_result(qid, **kwargs):
        """记录每题结果到后端"""
        try:
            import urllib.request, json
            data = json.dumps({"qid": qid, **kwargs}).encode()
            req = urllib.request.Request("http://127.0.0.1:5000/api/inspect/question-result",
                                        data=data, headers={"Content-Type": "application/json"})
            urllib.request.urlopen(req, timeout=2)
        except Exception:
            pass

    try:
        # 记录开始（★ 题目结果由 API 入口统一清空一次；此处不再 reset，
        #   否则多阶段遍历时后一阶段会清掉前一阶段的题）
        record_step("初始化", "running", f"{version_label} U{unit} {stage}")

        set_running("listening_inspect")
        config = load_config()
        adb = get_adb()

        # ===== B: 通用导航引擎 =====
        from src.universe_navigator import UniverseNavigator
        nav = UniverseNavigator(adb)

        log_msg(f"🚀 巡检: {version_label} Unit{unit} {stage}")
        record_step("1.启动APP", "running", "")

        # 强制回首页
        log_msg("强制回到首页")
        nav.universal_reset()
        record_step("1.启动APP", "done")
        time.sleep(1)

        # 通用导航: 首页 → 答题页
        record_step("2.导航到模块", "running")
        ok = nav.navigate_to("question_page", {
            "version": version_label,
            "unit": unit,
            "stage": stage,
        })
        if not ok:
            log_msg("⚠ 导航失败，尝试通用重置后重试...", "warning")
            nav.universal_reset()
            time.sleep(3)
            nav.navigate_to("question_page", {"version": version_label, "unit": unit, "stage": stage})
        record_step("2.导航到模块", "done")
        record_step("0.回到首页", "running", "")
        for _ in range(3):
            adb.press_back()
            time.sleep(0.5)
        time.sleep(1)
        record_step("0.回到首页", "done", "")

        # 1. 启动APP + 导航到首页
        log_msg("启动APP")
        adb.launch_app(config.app.package)
        time.sleep(4)
        record_step("1.启动APP", "done", "")

        # 1.5 检测并关闭教程页(如果有的话)
        log_msg("检查教程页")
        record_step("1.5关闭教程", "running", "")
        for _ in range(8):
            elements = adb.dump_ui()
            # 教程页特征: 出现"眺望远眺" / "下一关" / "知道了" / "开始体验"
            tutorial = False
            for e in elements:
                t = (e.text or '').strip()
                if t in ['下一关', '知道了', '开始体验', '我知道了', 'Next', '跳过']:
                    adb.tap(e.center[0], e.center[1])
                    log_msg(f"  关教程: '{t}' at {e.center}")
                    tutorial = True
                    time.sleep(1.5)
                    break
            if not tutorial:
                # 兜底: 点击屏幕中部右侧(教程"下一步"按钮区)
                # 但要先确认不是已经进入APP首页
                has_main = any('英语' in (e.text or '') for e in elements)
                if has_main:
                    break  # 已在APP首页
                adb.tap(*sc(900, 2200))  # 教程页"下一步"通常在右下
                time.sleep(1.5)
        record_step("1.5关闭教程", "done", "")

        # 2. 关闭广告 (用UI识别, 找到就停)
        log_msg("关闭广告")
        record_step("2.关闭广告", "running", "")
        ad_found = False
        for _ in range(3):
            if ad_found:
                break  # 已经关掉了, 不再跑后续兜底
            elements = adb.dump_ui(retries=1)
            found_close = False
            for e in elements:
                t = (e.text or '').strip()
                if t in ['关闭', '关闭广告', '×', 'X', 'Close', '跳过', '忽略', '我知道了', '取消']:
                    adb.tap(e.center[0], e.center[1])
                    log_msg(f"  关闭广告: '{t}' at {e.center}")
                    found_close = True
                    ad_found = True
                    time.sleep(1)
                    break
                if 'close' in e.resource_id.lower() or 'cancel' in e.resource_id.lower():
                    adb.tap(e.center[0], e.center[1])
                    found_close = True
                    ad_found = True
                    time.sleep(1)
                    break
            if not found_close:
                # 只在屏幕顶部边缘点击, 绝不到底部(会触发HOME)
                adb.tap(*sc(540, 40))
                time.sleep(0.5)
                adb.tap(*sc(540, 80))
            time.sleep(1)
        record_step("2.关闭广告", "done", "")

        # 3. 确保在英语tab + 检查是否跑到了"消息"页
        record_step("3.英语Tab", "running", "")
        adb.tap(*TABS["英语"])
        time.sleep(1.5)
        elements = adb.dump_ui()
        if any('消息' in (e.text or '') and e.center[1] < 300 for e in elements):
            log_msg("⚠ 进入了消息页, 按返回", "warning")
            adb.press_back()
            time.sleep(2)
            adb.tap(*TABS["英语"])
            time.sleep(1.5)
        record_step("3.英语Tab", "done", "")

        # 3.5 检查当前APP显示的版本/年级是否与用户选择的一致
        record_step("3.5版本检查", "running", f"用户选: {version_label}")
        log_msg("检查版本是否匹配")
        time.sleep(1)
        elements = adb.dump_ui()
        page_text = ' '.join([(e.text or '') for e in elements])
        # 提取APP显示的年级文本 (如"五年级上册")
        current_grade = ""
        for grade_text in ['一年级上册','一年级下册','二年级上册','二年级下册',
                           '三年级上册','三年级下册','四年级上册','四年级下册',
                           '五年级上册','五年级下册','六年级上册','六年级下册']:
            if grade_text in page_text:
                current_grade = grade_text
                break
        # 提取用户期望的年级 (从version_label如"新湘鲁六上"中提取)
        target_grade_map = {'五上':'五年级上册','五下':'五年级下册',
                            '六上':'六年级上册','六下':'六年级下册'}
        target_grade = ""
        for k, v in target_grade_map.items():
            if k in version_label:
                target_grade = v
                break

        log_msg(f"  APP当前: {current_grade or '未知'} | 用户期望: {target_grade or '未知'}")
        if current_grade and target_grade and current_grade != target_grade:
            log_msg(f"  ⚠ 版本不匹配! 需要切换到 {target_grade}", "warning")
            record_step("3.5版本检查", "running", f"切换 {current_grade} → {target_grade}")
            for e in elements:
                if current_grade in (e.text or ''):
                    adb.tap(e.center[0], e.center[1])
                    log_msg(f"  点版本标签 at {e.center}")
                    break
            time.sleep(2)
            elements2 = adb.dump_ui()
            for e in elements2:
                t = (e.text or '').strip()
                if target_grade in t:
                    adb.tap(e.center[0], e.center[1])
                    log_msg(f"  选择 {target_grade} at {e.center}")
                    break
            time.sleep(2)
            for e in adb.dump_ui():
                t = (e.text or '').strip()
                if t in ['确认', '确定', '完成', 'OK', '保存']:
                    adb.tap(e.center[0], e.center[1])
                    log_msg(f"  确认选择 at {e.center}")
                    break
            time.sleep(2)
        record_step("3.5版本检查", "done", current_grade or "已对齐")

        # ---- 模块自动识别 ----
        module_name, module_section, module_coords = detect_module_from_filename(docx_file or "")
        if module_name:
            log_msg(f"🎯 识别到模块: {module_name}")
        else:
            module_name, module_section, module_coords = "听力专项", "专项突破", None

        # 4. 根据模块区域分流
        record_step("4.专项突破", "running", "")
        skip_special_nav = False
        if module_section == "教材精学" and module_coords:
            # 教材精学区: 不进入专项突破,直接点模块
            log_msg(f"导航到教材精学 → {module_name} at {module_coords}")
            adb.tap(module_coords[0], module_coords[1])
            time.sleep(3)
            record_step("4.专项突破", "done", f"教材精学: {module_name}")
            # 跳过专项突破+听力专项导航, 直接到阶段选择
            skip_special_nav = True
        else:
            log_msg("导航到专项突破")

        # ---- 专项突破区: 导航+找模块 (教材精学区直接跳过) ----
        if not skip_special_nav:
            found = False
            for attempt in range(3):
                elements = adb.dump_ui()
                for e in elements:
                    if e.text and '专项突破' in e.text:
                        if 400 < e.center[1] < 2000:
                            adb.tap(e.center[0], e.center[1])
                            log_msg(f"  点'专项突破' at {e.center}")
                            found = True
                            break
                if found:
                    break
                adb.swipe(*sc(540, 1500), *sc(540, 1000), 300)
                time.sleep(1)
            if not found:
                adb.tap(*sc(414, 1700))
            time.sleep(3)
            record_step("4.专项突破", "done", "")

            # 5. 滚动找模块 (动态匹配,不限定"听力专项")
            record_step("5.听力专项", "running", "")
            log_msg(f"寻找模块: {module_name}")
            found_module = False
            for scroll_i in range(10):
                elements = adb.dump_ui()
                for e in elements:
                    t = (e.text or '')
                    # 支持多种匹配: 听力专项/巧记单词闯关/知识过关
                    if module_name in t or any(kw in t for kw in module_name.split()):
                        cx, cy = e.center
                        if 200 < cy < 2200 and cx > 50 and cx < 1030:
                            log_msg(f"  发现'{t.strip()}' at {e.center}")
                            adb.tap(cx, cy)
                            found_module = True
                            break
                if found_module:
                    break
                adb.swipe(*sc(540, 1600), *sc(540, 600), 500)
                time.sleep(1)
            if not found_module:
                log_msg(f"  ⚠ 未找到模块 '{module_name}', 尝试搜索全部文本", "warning")
                # 最后尝试: dump所有文本 + 模糊匹配
                elements = adb.dump_ui()
                texts = [(e.center, e.text) for e in elements if e.text and e.clickable]
                log_msg(f"  可点击元素: {[(c,t) for c,t in texts[:10]]}")
            time.sleep(3)
            record_step("5.听力专项", "done" if found_module else "error", module_name)
        else:
            # 教材精学区: 跳过专项突破+模块搜索
            record_step("5.听力专项", "done", f"教材精学: {module_name}")

        # 5.5 关"已完成"弹窗 (如果之前跑过该单元)
        time.sleep(1)
        elements = adb.dump_ui()
        for e in elements:
            t = (e.text or '').strip()
            if t in ['先走一步', '继续练习', '已完成', '完成']:
                adb.tap(e.center[0], e.center[1])
                log_msg(f"  关完成弹窗: '{t}' at {e.center}")
                time.sleep(1.5)
                break

        # 6. 版本检查: 如果页面上已经显示正确版本, 跳过
        log_msg(f"检查版本: {version_label}")
        time.sleep(2)
        elements = adb.dump_ui()
        # 检查是否跑到消息页/错误页
        if any('消息' in (e.text or '') and e.center[1] < 200 for e in elements):
            adb.press_back()
            time.sleep(2)
            elements = adb.dump_ui()
        
        # 从version_label提取年级关键词
        grade_keywords = {"五上": "五上", "五下": "五下", "六上": "六上", "六下": "六下"}
        target_kw = ""
        full_grade_map = {"五上":"五年级上册","五下":"五年级下册","六上":"六年级上册","六下":"六年级下册"}
        target_grade_full = ""
        for k, v in grade_keywords.items():
            if k in version_label:
                target_kw = v
                target_grade_full = full_grade_map.get(k, "")
                break
        
        # 检查当前页面是否已经显示正确版本
        page_text = ' '.join([(e.text or '') for e in elements])
        version_already_set = False
        if target_grade_full and target_grade_full in page_text:
            version_already_set = True
            log_msg(f"  版本已正确: {target_grade_full}")
        elif target_kw and target_kw in page_text:
            version_already_set = True
            log_msg(f"  版本已正确(含 {target_kw})")
        
        if target_kw and not version_already_set:
            # 需要手动点击切换版本
            found_grade = False
            for e in elements:
                if e.text and target_kw in (e.text or ''):
                    cx, cy = e.center
                    if 200 < cy < 2200:
                        adb.tap(cx, cy)
                        log_msg(f"  选中 {target_kw} at ({cx},{cy})")
                        found_grade = True
                        break
            if not found_grade:
                log_msg(f"  ⚠ 未找到 {target_kw}, 尝试兜底", "warning")
                adb.tap(*sc(540, 1200))
            time.sleep(3)
        elif version_already_set:
            log_msg(f"  版本一致, 跳过选择")
            time.sleep(1)

        # 7. 选择Unit - 先关可能出现的完成弹窗
        time.sleep(1)
        elements = adb.dump_ui()
        for e in elements:
            t = (e.text or '').strip()
            if t in ['先走一步', '继续练习', '已完成', '完成']:
                adb.tap(e.center[0], e.center[1])
                log_msg(f"  关完成弹窗: '{t}' at {e.center}")
                time.sleep(1.5)
                break

        log_msg(f"选择 Unit {unit}")
        elements = adb.dump_ui()
        # 找到Unit文本的位置 (如果不在当前屏幕, 滚动查找)
        unit_y = None
        for scroll_attempt in range(6):  # 最多滚动5次
            for e in elements:
                if e.text and re.match(rf'^Unit {unit}\b', (e.text or '').strip()):
                    unit_y = e.center[1]
                    log_msg(f"  Unit {unit} 文本 at y={unit_y} (scroll={scroll_attempt})")
                    break
            if unit_y:
                break
            # 没找到, 向下滚动
            if scroll_attempt < 5:
                adb.swipe(*sc(540, 1800), *sc(540, 500), 400)  # 大幅下滚
                time.sleep(0.8)
                elements = adb.dump_ui()
        # 兜底: 模糊匹配 (针对Unit数字小但文本格式不同)
        if not unit_y:
            for e in elements:
                if e.text and f"Unit {unit}" in (e.text or ''):
                    unit_y = e.center[1]
                    log_msg(f"  Unit {unit} 文本(模糊) at y={unit_y}")
                    break

        found_unit = False
        if unit_y:
            # Unit行的"去练习"按钮固定在该unit文本下方~100px处
            # 直接找 y 在 (unit_y + 50, unit_y + 150) 范围内的"去练习"
            btn_y_min = unit_y + 40
            btn_y_max = unit_y + 160
            for e in elements:
                t = (e.text or '').strip()
                if '去练习' in t or '去答题' in t:
                    cy = e.center[1]
                    if btn_y_min < cy < btn_y_max and e.center[0] > 0:
                        adb.tap(e.center[0], e.center[1])
                        log_msg(f"  点'去练习' for Unit {unit} at {e.center}")
                        found_unit = True
                        break
            if not found_unit:
                # 兜底: 直接用固定偏移量点
                btn_x = 881  # 常见右侧按钮x
                log_msg(f"  兜底: 点({btn_x}, {unit_y + 99})", "warning")
                adb.tap(btn_x, unit_y + 99)
                found_unit = True
        else:
            log_msg(f"  ⚠ 没找到 Unit {unit} 文本", "warning")
            adb.tap(*sc(540, 1400))
        time.sleep(4)

        # 8. 现在应该在阶段选择页 (基础巩固/综合进阶/难点突破)
        log_msg(f"选择阶段: {stage}")
        found_stage = False
        for wait_i in range(5):
            elements = adb.dump_ui()
            # 直接找阶段按钮, 不判断"是否还在Unit列表"(避免误判)
            for e in elements:
                t = (e.text or '').strip()
                if not t: continue
                # 方式1: 精确匹配
                # 方式2: stage是t的子串(如"基础巩固"在"基础巩固模式"中)
                # 方式3: 尝试用关键字符找(基础/综合/难点、巩固/进阶/突破)
                matched = False
                if t == stage: matched = True
                elif stage in t or t in stage: matched = True
                elif any(kw in t for kw in stage.replace(' ','').split('·')):
                    # 对付"基础·巩固"这类分隔符,拆开后分别匹配
                    for seg in stage.replace('·','').split():
                        if seg in t: matched = True; break
                if matched and e.clickable and 50 < e.center[1] < 2200:
                    adb.tap(e.center[0], e.center[1])
                    log_msg(f"  点'{t}' at {e.center}")
                    found_stage = True
                    break
            if found_stage:
                break
            log_msg(f"  未找到 {stage} ({wait_i+1}/5)")
            time.sleep(1)
        if not found_stage:
            log_msg(f"  ⚠ 未找到 {stage},可能已是默认阶段,继续执行", "warning")
            # 阶段选择失败不致命 — 可能该页面没有阶段选择,直接进入题目
            # 或者"基础巩固"是默认已选中的
        time.sleep(3)

        # 9. 开始答题 → 轮流尝试 "开始答题"/"去答题"→ 关弹窗 → 等待考题加载
        log_msg("开始答题")
        for attempt in range(5):
            elements = adb.dump_ui()
            # 先检查是否已经在考题页
            in_question = any(re.match(r'^\d+/\d+$', (e.text or "").strip()) for e in elements)
            if in_question:
                log_msg("  已在考题页!", "success")
                break
            # 找"开始答题"/"去答题"按钮
            found_btn = False
            for e in elements:
                t = (e.text or '').strip()
                if t in ['开始答题', '去答题', '开始', 'Start']:
                    if e.clickable and e.center[1] < 2000:
                        adb.tap(e.center[0], e.center[1])
                        log_msg(f"  点'{t}' at {e.center}")
                        found_btn = True
                        break
            if not found_btn:
                # 可能已经进入题目但还没加载, 或弹窗挡住了
                # 尝试点屏幕底部中间 (有些APP的确认在底部)
                adb.tap(*sc(540, 2100))
                # 关可能存在的弹窗
                adb.tap(*sc(540, 1800))
            time.sleep(2)

        # 再多等一会儿, 确保题目渲染完成
        time.sleep(2)

        # 10. 逐题巡检循环
        log_msg(f"逐题巡检开始: 脚本={docx_file or '无'}")
        record_step("6.逐题巡检", "running", "")
        last_q = 0
        questions_reviewed = []
        
        # 加载脚本数据 (用于LLM审查)
        from src.review_agent import ReviewAgent, ReviewConfig
        from src.parse_yingyubao_docx import parse

        agent = None
        docx_path = ""
        if docx_file:
            docx_path = str(UPLOAD_DIR / docx_file)
            if Path(docx_path).exists():
                cfg = ReviewConfig(docx_path=docx_path, unit=unit, stage=stage,
                                   screenshot_dir=str(PROJECT_ROOT / "screenshots"), verbose=False)
                agent = ReviewAgent(cfg)
                log_msg(f"  脚本已加载: {docx_file} (共{len(agent.script_questions)}题)", "success")
            else:
                log_msg(f"  ⚠ 脚本不存在: {docx_path}", "warning")

        for loop in range(80):
            if not task_status["running"]:
                log_msg("任务被手动停止", "warning")
                break

            elements = adb.dump_ui(retries=2)
            all_texts = [(e.text or '').strip() for e in elements]

            # ★ 实时题型识别（不依赖脚本，从 UI 元素判断当前是什么题）
            detected = None
            try:
                from src.type_detector import TypeDetector
                detected = TypeDetector(verbose=False).detect(elements)
                if detected and detected.type_1 != "未知":
                    log_msg(f"  🔍 实时题型: {detected.describe()} (conf={detected.confidence:.0%})", "info")
            except Exception:
                detected = None

            # ====== 检测是否在结果页 (有"正确答案"文本) ======
            if any('正确答案' in t for t in all_texts):
                log_msg("[结果页] 刚答完一题, 找'下一题'", "info")
                found_next = False
                for e in elements:
                    t = (e.text or '').strip()
                    if t in ['下一题', '完成', '继续', 'Next']:
                        adb.tap(e.center[0], e.center[1])
                        log_msg(f"  点'{t}' at {e.center}")
                        found_next = True
                        break
                if not found_next:
                    adb.tap(*sc(540, 2100))  # 底部兜底
                time.sleep(1.5)
                continue

            # ====== 读题号 ======
            cur = None
            total_q = 0
            for e in elements:
                m = re.match(r'^(\d+)/(\d+)$', (e.text or "").strip())
                if m:
                    cur = int(m.group(1))
                    total_q = int(m.group(2))
                    break
            if not cur:
                log_msg("不在考题页,结束", "warning")
                break
            if cur == last_q:
                time.sleep(0.3)
                continue
            last_q = cur

            # 截图
            shot_name = f"q{cur:02d}.png"
            adb.screenshot(shot_name)
            shot_path = str(PROJECT_ROOT / "screenshots" / shot_name)
            log_msg(f"Q{cur:02d}", "success")
            update_progress(cur, total_q, f"Q{cur}")

            # 用ReviewAgent做四维审查 (如果脚本已加载)
            review_result = None
            if agent:
                try:
                    # 找对应脚本题目 — 用列表索引而非global_idx
                    # APP题号是相对位置的(1/15), 脚本按unit+stage过滤后也是有序的
                    script_q = None
                    if cur and 1 <= cur <= len(agent.script_questions):
                        script_q = agent.script_questions[cur - 1]
                    if script_q:
                        # 从UI dump提取屏幕文字 (供文字题走文本模型加速)
                        ui_texts = [e.text or '' for e in elements if e.text and e.text.strip()]
                        r = agent._review_one(script_q, shot_path, ui_texts=ui_texts, detected=detected)
                        review_result = {
                            "stem": "✅" if r.stem_check.passed else "❌",
                            "content": "✅" if r.content_check.passed else "❌",
                            "image": "✅" if r.image_check.passed else "❌",
                            "answer": "✅" if r.answer_check.passed else "❌",
                            "score": r.overall_score,
                            "stem_reason": r.stem_check.details[0] if r.stem_check.details else "",
                            "content_reason": r.content_check.details[0] if r.content_check.details else "",
                            "image_reason": r.image_check.details[0] if r.image_check.details else "",
                            "answer_reason": r.answer_check.details[0] if r.answer_check.details else "",
                        }
                        log_msg(f"  审查: 题干{review_result['stem']} 内容{review_result['content']} 配图{review_result['image']} 作答{review_result['answer']} 得分{review_result['score']:.2f}")

                        # ====== 发送每题结果到后端 ======
                        # ★ 多阶段：qid 带 stage 前缀，避免不同子模块同题号互相覆盖
                        qid = f"{version_label}-U{unit}-{stage}-Q{cur:02d}"
                        det_type = (detected.full_type if detected and detected.full_type else script_q.type_2)
                        record_q_result(
                            qid,
                            idx=cur,
                            total=total_q,
                            question_type=det_type,
                            stem=(detected.stem if detected and detected.stem else script_q.stem)[:40] + "..." if len((detected.stem if detected and detected.stem else script_q.stem)) > 40 else (detected.stem if detected and detected.stem else script_q.stem),
                            recording=script_q.recording,
                            script_answer=script_q.answer,
                            ai_stem=r.stem_check.passed,
                            ai_content=r.content_check.passed,
                            ai_image=r.image_check.passed,
                            ai_answer=r.answer_check.passed,
                            overall_passed=r.overall_passed,
                            overall_score=round(r.overall_score, 2),
                            stem_reason=r.stem_check.details[0][:100] if r.stem_check.details else "",
                            content_reason=r.content_check.details[0][:100] if r.content_check.details else "",
                            image_reason=r.image_check.details[0][:100] if r.image_check.details else "",
                            answer_reason=r.answer_check.details[0][:100] if r.answer_check.details else "",
                            ai_reason=r.content_check.details[0][:100] if r.content_check.details else "",
                            screenshot=shot_name,
                            knowledge_check=r.knowledge_check,
                        )

                        # 加入反馈 (临时先用AI判断)
                        from src.feedback_loop import FeedbackSample
                        agent.feedback.add(FeedbackSample(
                            question_id=qid,
                            check_dimension="all",
                            human_judgment="通过" if r.overall_passed else "不通过",
                            ai_judgment="通过" if r.overall_passed else "不通过",
                            ai_reason=r.content_check.details[0] if r.content_check.details else "",
                            question_type=script_q.type_2,
                            screenshot=shot_name,
                        ))
                except SystemExit:
                    log_msg("⏹ 任务已被立即停止", "warning")
                    raise  # ★ 必须重新抛出，让注入的 SystemExit 继续传播到任务线程外层收尾
                except Exception as e:
                    log_msg(f"  审查异常: {e}", "warning")

            questions_reviewed.append({
                "idx": cur, "screenshot": shot_name, "review": review_result
            })

            # 点击选项 — 根据脚本答案选择正确选项
            found_option = False
            if script_q and script_q.answer in 'ABC':
                # 找对应选项字母的文本 (A/B/C)
                ans_letter = script_q.answer
                for e in elements:
                    t = (e.text or '').strip()
                    # 匹配 "A." / "A、" / "A" 等格式的选项文本
                    if re.match(rf'^{ans_letter}[\.\、\)\s]', t) and e.clickable and 50 < e.center[1] < 2200:
                        adb.tap(e.center[0], e.center[1])
                        log_msg(f"  点选项 {ans_letter} at {e.center}")
                        found_option = True
                        break
                if not found_option:
                    # 按选项位置推定: A左/B中/C右
                    opt_x = {'A': 250, 'B': 500, 'C': 750, 'D': 950}.get(ans_letter, 500)
                    candidates = [e for e in elements if e.clickable and 400 < e.bounds[1] < 2000
                                 and abs(e.center[0] - opt_x) < 200 and (e.bounds[2]-e.bounds[0]) > 60]
                    if candidates:
                        # 选最靠上的一个
                        best = min(candidates, key=lambda e: e.center[1])
                        adb.tap(best.center[0], best.center[1])
                        log_msg(f"  点选项(位置推定) {ans_letter} at {best.center}")
                        found_option = True
            
            if script_q and script_q.answer in 'TF':
                # 判断题: 点T或F
                ans_word = {'T': 'T', 'F': 'F'}.get(script_q.answer, '')
                for e in elements:
                    t = (e.text or '').strip().upper()
                    if t == ans_word and e.clickable:
                        adb.tap(e.center[0], e.center[1])
                        log_msg(f"  点 {'正确' if ans_word=='T' else '错误'} at {e.center}")
                        found_option = True
                        break
            
            if not found_option:
                # 兜底: 点中间偏左单选区域
                for e in elements:
                    if e.clickable and 400 < e.bounds[1] < 2000 and (e.bounds[2]-e.bounds[0]) > 100:
                        if e.center[0] < 150: continue
                        adb.tap(e.center[0], e.center[1])
                        found_option = True
                        break
            if not found_option:
                adb.tap(*sc(540, 800))
            time.sleep(1.5)

            # 找"检查"按钮, 点它
            found_check = False
            for _ in range(3):
                new_elems = adb.dump_ui()
                for e in new_elems:
                    t = (e.text or '').strip()
                    if t in ['检查', 'Check', '提交', '下一题', '完成']:
                        adb.tap(e.center[0], e.center[1])
                        log_msg(f"  '检查'点中 ({e.center[0]},{e.center[1]})")
                        found_check = True
                        break
                if found_check:
                    break
                adb.tap(*sc(540, 2100))
                time.sleep(0.5)
            time.sleep(1.5)

            # 完成
            if cur >= total_q:
                log_msg(f"✅ 全部 {total_q} 题完成!", "success")
                record_step("6.逐题巡检", "done", f"完成 {len(questions_reviewed)}/{total_q}")
                record_step("7.生成报告", "running", "")
                break

        # 保存审查结果
        if agent and questions_reviewed:
            agent.export_report(str(PROJECT_ROOT / "outputs" / f"review_{docx_file.replace('.docx','')}.md"))
            agent.export_json(str(PROJECT_ROOT / "outputs" / f"review_{docx_file.replace('.docx','')}.json"))
            record_step("7.生成报告", "done", f"已保存到 outputs/")

        if not keep_running:
            set_done()
        log_msg(f"✅ 听力专项巡检完成! {len(questions_reviewed)}题", "success")

    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        import traceback
        log_msg(f"❌ 巡检失败: {e}", "error")
        log_msg(traceback.format_exc(), "error")
        if not keep_running:
            set_done()



@app.route("/api/inspect/quick-run", methods=["POST"])
def api_inspect_quick_run():
    """⚡ 快速检查：从当前页面直接开始逐题巡检 (跳过导航)

    可选参数: docx(脚本文件名), unit(单元号) — 提供后逐题做 AI 六维审查
    """
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409
    data = request.get_json() or {}
    docx_file = data.get("docx", "")
    unit = data.get("unit", 0)
    t = threading.Thread(target=run_quick_inspect_task,
                         args=(docx_file, unit), daemon=True)
    t.start()
    return jsonify({"status": "started", "task": "quick_inspect",
                    "docx": docx_file, "unit": unit})


@app.route("/api/inspect/run", methods=["POST"])
def api_inspect_run():
    """启动逐题巡检 (任务版)"""
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409
    t = threading.Thread(target=run_inspect_loop, daemon=True)
    t.start()
    return jsonify({"status": "started", "task": "inspect_loop"})


@app.route("/api/inspect/stop", methods=["POST"])
def api_inspect_stop():
    """停止正在运行的任务（设置停止标志 + 强制中断线程，与 modules/stop 一致）"""
    if not task_status["running"]:
        return jsonify({"status": "idle", "message": "没有正在运行的任务"})
    # ★ 必须真正停止：只设 running=False 无法中断脚本线程（线程内部不检查 running），
    #   先执行统一停止（标志 + 注入 SystemExit），再清 running 状态
    forced, message = _request_stop()
    task_status["running"] = False
    return jsonify({"status": "stopping", "message": message})


# ============================================================
# 🔵 审查反馈 路由区 — 存储AI判断+人工标注,让审查智能体学习
# ============================================================

_inspection_state = {
    "current_question_idx": 0,
    "questions": {},           # {qid: {ai_result, human_label, timestamp, ...}}
    "workflow_steps": [],     # 当前运行的步骤
    "docx": "",
    "version": "未知版本",
    "unit": "全部",
    "stage": "全部",   # ★ 不能用 "?"（Windows 文件名非法字符）
    "live_report_path": "",    # ★ 实时错题报告(report_live.html)路径（借鉴队友）
    "_reuse_info": "",         # ★ 审查体系改造：复用标记（如"巧记单词 U1 已于2026-08-26审查，本次复用结论")
}

# ★ 实时报告重生成锁（避免多线程下并发写同一文件）
_live_report_lock = threading.Lock()

# ============================================================
# ★ 审查体系改造：脚本审查索引（script_review_index.json）
#   每个（模块+单元）只认真审一次，结论沉淀复用。
#   索引 key = "<版本>/<年级>/<模块>/<单元>"
# ============================================================
SCRIPT_REVIEW_INDEX_PATH = PROJECT_ROOT / "data" / "script_review_index.json"

def load_script_review_index() -> dict:
    """加载脚本审查索引"""
    import json
    try:
        if SCRIPT_REVIEW_INDEX_PATH.exists():
            return json.loads(SCRIPT_REVIEW_INDEX_PATH.read_text(encoding="utf-8") or "{}")
    except Exception:
        pass
    return {}

def save_script_review_index(index: dict):
    """保存脚本审查索引"""
    import json
    SCRIPT_REVIEW_INDEX_PATH.parent.mkdir(parents=True, exist_ok=True)
    SCRIPT_REVIEW_INDEX_PATH.write_text(
        json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")

def _script_file_hash(docx_path) -> str:
    """计算脚本文件 md5 hash，用于检测脚本是否变更"""
    import hashlib
    try:
        return hashlib.md5(open(docx_path, "rb").read()).hexdigest()
    except Exception:
        return ""

# ★ 自动生成脚本的存放目录（与手动上传的 uploads/ 区分）
GENERATED_SCRIPTS_DIR = PROJECT_ROOT / "uploads" / "generated"
GENERATED_SCRIPTS_DIR.mkdir(parents=True, exist_ok=True)

def _save_inspection_state():
    """保存巡检状态到文件, 审查智能体下次学习"""
    import json
    p = PROJECT_ROOT / "data" / "inspection_state.json"
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(
        json.dumps(_inspection_state, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def _save_inspection_state_merge():
    """★ 合并保存：保留旧文件里【其他模块】的题目（多模块检测时，
    新模块结果不与旧模块互相覆盖——否则导出的错题报告只会剩最后一个模块）。
    仅多模块检测的 LLM 审查完成后调用（此时内存含本次全部模块结果）。"""
    import json
    p = PROJECT_ROOT / "data" / "inspection_state.json"
    p.parent.mkdir(parents=True, exist_ok=True)
    try:
        old = {}
        if p.exists():
            old = json.loads(p.read_text(encoding="utf-8") or "{}")
    except Exception:
        old = {}
    # 合并 questions：本次内存的覆盖同 qid；旧文件里不冲突的保留
    _qs = dict(old.get("questions") or {})
    _qs.update(dict(_inspection_state.get("questions") or {}))
    _merged = dict(old)
    _merged["questions"] = _qs
    # 时间/模块信息以最新为准
    for k in ("version", "grade", "module", "start_time", "end_time"):
        if _inspection_state.get(k):
            _merged[k] = _inspection_state[k]
    p.write_text(json.dumps(_merged, ensure_ascii=False, indent=2), encoding="utf-8")


def _live_regen_error_report():
    """
    ★ 实时重新生成"仅错误"的可折叠卡片 HTML 报告（审查中每出一道错题即调用）。
    报告写入 outputs/reports/<版本>/U<单元>_<阶段>_<日期>/report_live.html，
    并刷新 latest 入口。路径记录在 _inspection_state["live_report_path"]。
    """
    _live_report_lock.acquire()
    try:
        from src.report_exporter import ReportExporter
        state = _inspection_state
        qs_all = list(state.get("questions", {}).values())
        meta = {
            "version": state.get("version", "未知版本"),
            "unit": state.get("unit", "全部"),
            "stage": state.get("stage", "全部"),   # ★ "?" 是 Windows 非法字符，会 WinError 123
            "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "total": len(qs_all),
        }
        exp = ReportExporter()
        path = exp.export_html_live(qs_all, meta)
        state["live_report_path"] = path
        return path
    except Exception as e:
        log_msg(f"⚠ 实时错题报告生成失败: {e}", "warning")
        return ""
    finally:
        _live_report_lock.release()


# ★ 多模块检测每题界面级证据 → 审查结果区（无脚本也能展示 AI 通过/不通过）
_SCRIPT_CACHE = {}  # ★ 脚本解析缓存 {docx名: script_questions}，避免每题重复解析100题

def _record_module_evidence(qidx: int, msg: str, evidence: list, big: int = None):
    """把 engine.py 每题收集的 5 维界面证据（题型/题干/选项/音频/作答）
    映射成前端六维卡片字段写入 _inspection_state["questions"]。

    判定规则：
      - text_ok → 通过(true)；text_mismatch → 不通过(false)；skip → 未检(None)
      - overall: 可查维度≥3且无不通过 → 通过；否则 → 不通过；全是 skip → 未审
    big: 大题号（口语训练"第2大题·第3小题"精确定位；None=无大题概念）
    """
    global _inspection_state
    # ★ 模块上下文（scheduler 设置）：模块名/子模块/模块内题号
    try:
        from common.logger import get_current_module
        _current_module_name, _current_stage_name = get_current_module()
    except Exception:
        _current_module_name, _current_stage_name = "", ""
    if not _current_module_name:
        _current_module_name = _inspection_state.get("module", "")
        _current_stage_name = _inspection_state.get("stage", "")
    # ★ 模块内题号：优先直接用 engine 传来的题号 qidx（engine 的"第N题"就是
    #   当前子模块内题号，App 页面显示 9/10 的 9），不再用跨子模块累加的计数器
    _current_module_qno = qidx
    # 证据维度 → 前端六维映射（"题型"不映射到任何六维，仅作为元信息）
    field_map = {
        "题干": "stem",
        "选项": "content",
        "音频": "audio",
        "作答": "answer",
        "作答匹配": "answer",   # ★ 题型-作答一致性（判断题给ABC等）→ 作答维度
    }
    dims = {"stem": None, "content": None, "image": None,
            "answer": None, "audio": None, "post_error": None}
    reasons = {"stem": "", "content": "", "image": "",
               "answer": "", "audio": "", "post_error": ""}

    # 收集实际显示文字（题干+选项，供前端卡展示）
    stem_text = ""
    option_text = ""
    question_type = "界面检查"  # 默认

    for e in evidence:
        field = e.get("field", "")
        # ★ 分值等附加信息：写入 note（不参与六维判定），供检查人员参考
        if field == "分值":
            _sc = (e.get("actual") or e.get("diff") or "").strip()
            if _sc:
                _inspection_state["_last_score_info"] = _sc
            continue
        # ★ 2026-08-26 扬声器播放内容（听音题）：存到状态供脚本比对/题干后展示
        if field == "扬声器":
            _spk = (e.get("actual") or "").strip()
            if _spk:
                _inspection_state["_last_speaker_word"] = _spk
            continue
        # 题型识别 → 只更新 question_type
        if field == "题型":
            diff = e.get("diff") or ""
            if diff and ("[" in diff):
                question_type = diff.split("[")[1].split("]")[0] if "]" in diff else diff
            continue

        dim = field_map.get(field)
        if not dim:
            continue
        etype = e.get("type", "")
        # ★ mismatch 优先：同一维度多条证据时，只要有一条不通过就判不通过
        #   （口语题麦克风不可点 + 通用"作答元素存在"同时写入 answer 维度时，
        #     不通过不能被后写的通过覆盖）
        if etype == "text_ok":
            if dims[dim] is not False:
                dims[dim] = True
        elif etype == "text_mismatch":
            dims[dim] = False
        else:
            if dims[dim] is None:
                dims[dim] = None

        # ★ 提取实际文字：题干/选项的实际内容
        actual = e.get("actual") or ""
        if field == "题干" and actual and actual != "(无题干文字)" and actual != "(无)":
            stem_text = actual[:150]
        if field == "选项" and actual:
            option_text = actual[:120]

        # ★ 生成自然语言原因（检查人员一眼能看懂；不通过原因优先保留）
        diff_str = e.get("diff") or ""
        if etype == "text_mismatch":
            reasons[dim] = _fmt_reason(field, diff_str, "不通过")
        elif etype == "text_ok":
            # 仅当该维度最终判定不是不通过时才写"通过"原因（避免覆盖不通过信息）
            if dims[dim] is not False:
                reasons[dim] = _fmt_reason(field, diff_str, "通过")
        else:
            reasons[dim] = f"{field}未检" if diff_str else ""

    # ★ 审查体系改造：不适用维度直接省略，不再填"需连手机截图"等占位文案
    #   配图/答错后/音频 在不适用维度时，reasons 留空 → 前端不渲染该维度

    # 2026-08-25 新增：App 实际题干 ↔ 脚本题干 匹配校验（口语训练核心需求：
    #   txt 流程第8步"答案与关键词是否正确/听力材料是否正确"→ 每道小题必须核对
    #   App 显示的题干是否与脚本一致，如"第1大题第1小题脚本=then，App 也必须显示 then"）
    # ★ 2026-08-26 巧记单词：抽题模式无序号对应，改为"内容匹配"——
    #   ① 题干类型一致：App 题干（如"听录音选释义"）在脚本同单元能找到同类题型
    #   ② 内容一致：App 关键词/选项/录音词 在脚本该题型下能找到对应单词
    def _script_match_qiaoji(_sq_list, _u_cur):
        """巧记单词内容匹配：返回 (匹配到的脚本题 or None, 相似度)。
        匹配策略：按 App 题干在脚本同单元找最相似题。
          - ratio ≥ 0.6 → 同题型，返回该题（上层写"✅题干类型一致"）
          - ratio < 0.6 → 无同类题（App 题型/题干与脚本不符），返回 None（上层报不通过）
        """
        import difflib as _dl
        _app_all = (stem_text + " " + option_text).strip()
        _cands = [s for s in _sq_list if _u_cur and s.unit == _u_cur]
        if not _cands:
            _cands = _sq_list
        _best, _best_ratio = None, 0.0
        for s in _cands:
            _ss = (getattr(s, "stem", "") or "").strip()
            if not _ss:
                continue
            _r = _dl.SequenceMatcher(None, _ss[:20], stem_text[:20]).ratio() if stem_text else 0.0
            if _r > _best_ratio:
                _best_ratio, _best = _r, s
        if _best is not None and _best_ratio >= 0.6:
            return _best, _best_ratio
        return None, _best_ratio

    try:
        _docx_cur = (_GLOBAL_DOCX_MAP or {}).get(_current_module_name, "")
        if _docx_cur and stem_text and stem_text not in ("(无题干文字)", "(无)"):
            _docx_path_cur = PROJECT_ROOT / "uploads" / _docx_cur
            if _docx_path_cur.exists():
                # ★ 脚本解析缓存（每题调用一次，100题脚本重复解析太慢）
                if _docx_cur not in _SCRIPT_CACHE:
                    from src.review_agent import ReviewAgent, ReviewConfig
                    _agent_cur = ReviewAgent(ReviewConfig(
                        docx_path=str(_docx_path_cur), unit=0, screenshot_dir="", verbose=False))
                    _SCRIPT_CACHE[_docx_cur] = _agent_cur.script_questions
                _sq_list = _SCRIPT_CACHE[_docx_cur]
                # ★ 当前单元号
                _u_cur = _inspection_state.get("unit", 0)
                try:
                    _u_cur = int(str(_u_cur).split("-")[0])
                except Exception:
                    _u_cur = 0
                # ★ 巧记单词：内容匹配（无序号，按题型+内容找）
                _qiaoji_ratio = 0.0
                if _current_module_name == "巧记单词":
                    _sq_cur, _qiaoji_ratio = _script_match_qiaoji(_sq_list, _u_cur)
                    if _sq_cur is None:
                        # ★ 题干类型与脚本不符（App 题型在脚本中找不到同类）
                        #   → 题干维度不通过，原因写清楚（用户要求：App看图填单词但脚本听录音填单词）
                        dims["stem"] = False
                        _script_types = sorted({getattr(s, "type_2", "") or "未知"
                                                for s in (_sq_list if _u_cur else _sq_list)
                                                if getattr(s, "unit", 0) == _u_cur and (getattr(s, "type_2", "") or "") not in ("", "/")})[:5]
                        reasons["stem"] = (f"❌ 题型与脚本不符：App 题干「{stem_text[:40]}」在脚本中无同类题"
                                           f"（脚本该单元题型: {_script_types or '未知'}）")
                        _sq_cur = None
                else:
                    _sq_cur = None
                    # ★ 按 (unit, big, stage_idx) 定位脚本小题（口语训练结构）
                    if big and _u_cur:
                        for _s in _sq_list:
                            if _s.unit == _u_cur and _s.big == big and _s.stage_idx == qidx:
                                _sq_cur = _s
                                break
                    if not _sq_cur:
                        for _s in _sq_list:
                            if _s.global_idx == qidx:
                                _sq_cur = _s
                                break
                if _sq_cur:
                    _script_stem = (getattr(_sq_cur, "stem", "") or "").strip()
                    _script_rec = (getattr(_sq_cur, "recording", "") or "").strip()
                    if _current_module_name == "巧记单词":
                        # ★ 巧记单词匹配成功：题干类型一致
                        if dims["stem"] is not False:
                            dims["stem"] = True
                            reasons["stem"] = (f"✅ 题干类型与脚本一致（脚本同类题「{_script_stem[:40]}」，"
                                               f"App：「{stem_text[:40]}」）")
                        # ★ 2026-08-26 扬声器播放内容 ↔ 脚本 recording 对比（ASR）
                        #   扬声器词 = 从题干后提取的英文单词（_last_speaker_word，由模块写入）
                        _spk = _inspection_state.get("_last_speaker_word", "")
                        if _script_rec:
                            try:
                                from common.asr import compare_against_script
                                _cmp = compare_against_script(_spk, _script_rec)
                                _audio_reason = _cmp.get("reason", "")
                                # 扬声器内容 append 到题干原因后（用户要求：放题干后面展示）
                                if _cmp.get("consistent") is False and dims["audio"] is not False:
                                    dims["audio"] = False
                                if dims["audio"] is None and _audio_reason:
                                    reasons["audio"] = _audio_reason
                                if dims["stem"] is not False and _spk:
                                    reasons["stem"] += f"（扬声器：{_spk}）"
                            except Exception:
                                pass
                            # 扬声器词为空（页面无英文词）→ 提示需接入 ASR 才能核对
                            if not _spk and dims["audio"] is None:
                                reasons["audio"] = (f"脚本录音「{_script_rec[:60]}」；"
                                                    f"扬声器内容需接入 ASR 才能核对（当前未识别到）")
                        elif dims["audio"] is None:
                            reasons["audio"] = f"脚本录音「{_script_rec[:60]}」"
                    elif _script_stem:
                        # ★ 相似度比对（App 提取文字可能带噪音，用包含/相似而非全等）
                        _app_stem = stem_text.strip()
                        _ok = (_script_stem in _app_stem) or (_app_stem in _script_stem)
                        if not _ok:
                            import difflib as _dl
                            _ratio = _dl.SequenceMatcher(None, _script_stem, _app_stem).ratio()
                            _ok = _ratio >= 0.85
                        if not _ok:
                            dims["stem"] = False   # ★ mismatch 优先：脚本比对不通过 → 题干不通过
                            reasons["stem"] = (f"❌ 题干与脚本不符：App 显示「{_app_stem[:60]}」，"
                                               f"脚本应为「{_script_stem[:60]}」（位置：第{big}大题·第{qidx}小题）")
                        else:
                            # ★ 通过也写比对原因（检查人员知道这题已核对过脚本）
                            if dims["stem"] is not False:
                                dims["stem"] = True
                                reasons["stem"] = f"✅ 题干与脚本一致（脚本：{_script_stem[:50]}，App：{stem_text[:50]}）"
                    if _script_rec and dims["audio"] is None:
                        # ★ 听力材料（脚本 recording）有内容但无音频证据 → 留空不渲染
                        pass
    except Exception:
        pass

    # ★ 综合判定：可查维度数（非None）≥3 且 无任何不通过 → 通过
    checked = [d for d in dims.values() if d is not None]
    failed = [d for d in dims.values() if d is False]
    passed_cnt = sum(1 for d in dims.values() if d is True)

    if not checked:
        overall = None   # 全未检 → 未审查
    elif failed:
        overall = False  # 任何不通过 → 不通过
    elif len(checked) >= 3:
        overall = True   # 至少查到3维且全通过 → 通过
    else:
        overall = None   # 检查不足3维 → 标未审

    # ★ 自动截图：错题（overall=False）自动抓当前页面截图，供错题报告嵌入。
    #   用户要求"错题日志中的每道错题都要有截图"；听力专项有答错截图，
    #   口语训练等模块没有 → 这里统一兜底（每题判定不通过即截图）。
    _auto_shot = ""
    if overall is False:
        try:
            # ★ 截图前等 0.8s：页面内容（排序项/选项）渲染有延迟，立即截图
            #   会截到空白/半渲染画面（用户实测"排序题内容还没出来"）
            time.sleep(0.8)
            _shot_dir = PROJECT_ROOT / "screenshots"
            _shot_dir.mkdir(parents=True, exist_ok=True)
            _auto_shot = f"auto_{_current_module_name or 'mod'}_q{qidx:03d}_{datetime.now().strftime('%H%M%S')}.png"
            _conn = _connect_device()
            if _conn is not None:
                for _r in range(3):
                    try:
                        _conn.screenshot(str(_shot_dir / _auto_shot))
                        break
                    except OSError:
                        if _r >= 2:
                            raise
                        time.sleep(0.5)
        except Exception:
            _auto_shot = ""

    total = len(_inspection_state.get("questions", {})) + 1
    # ★ 2026-08-25 修复：口语训练"第N大题·第M小题"结构下，不同大题的小题号重复
    #   （每个大题都有小题1-5）→ 旧 qid "auto-Q{小题号}" 会被后面大题的同号小题覆盖，
    #   只剩最后一大题。改为 "auto-Q{大题:02d}-{小题:02d}" 保证每题唯一；
    #   无 big（听力专项/单元自检等）保持旧格式 auto-Q{idx:03d}（与 summary 统计兼容）。
    if big:
        qid = f"auto-Q{big:02d}-{qidx:02d}"
    else:
        qid = f"auto-Q{qidx:03d}"

    # ★ 描述文字：让检查人员知道题目大概内容和位置
    # ★ 描述文字：让检查人员知道题目大概内容和位置
    if not stem_text:
        stem_text = f"第{qidx}题（{question_type}）"
    # ★ 大题号定位（口语训练"第2大题·第3小题"）：错题报告能定位到 大题·小题
    if big:
        stem_text = f"第{big}大题·{stem_text}" if not stem_text.startswith("第") else f"第{big}大题·{stem_text}"
        progress_str = f"大{big}-Q{qidx}"
    else:
        progress_str = f"Q{qidx}"
    _loc = {"big": big} if big else {}

    _inspection_state["questions"][qid] = {
        "idx": qidx,
        "total": total,
        "question_type": question_type,
        "screenshot": _auto_shot,  # ★ 错题自动截图（听力专项答错截图 + 通用兜底）
        "progress": progress_str,
        # ★ 分值信息（evidence 附加项，供检查人员参考）
        "score_info": _inspection_state.get("_last_score_info", ""),
        # ★ 模块定位信息（哪个模块·哪个子模块·模块内第几题）→ 错题报告按此分组
        "module": _current_module_name,
        "stage": _current_stage_name,
        "module_qno": _current_module_qno,
        **_loc,
        # 六维判定（旧格式，兼容已有消费者）
        "ai_stem": dims["stem"], "ai_content": dims["content"],
        "ai_image": dims["image"], "ai_answer": dims["answer"],
        "ai_audio": dims["audio"], "ai_post_error": dims["post_error"],
        "overall_passed": overall,
        "overall_score": round(passed_cnt / max(len(checked), 1), 2) if checked else 0.0,
        # 详细原因（自然语言）
        "stem_reason": reasons["stem"], "content_reason": reasons["content"],
        "image_reason": reasons["image"], "answer_reason": reasons["answer"],
        "audio_reason": reasons["audio"], "post_error_reason": reasons["post_error"],
        # ★ 审查体系改造：新结构——统一结论字典（旧字段兼容保留，迁移期并存）
        #   每个维度含 status(pass/fail/na) + detail(具体结论文本)
        #   na = 不适用（该维度无需检查），前端不渲染
        "conclusions": {
            "stem":        {"status": "pass" if dims["stem"] is True else ("fail" if dims["stem"] is False else "na"), "detail": reasons["stem"]},
            "content":     {"status": "pass" if dims["content"] is True else ("fail" if dims["content"] is False else "na"), "detail": reasons["content"]},
            "image":       {"status": "pass" if dims["image"] is True else ("fail" if dims["image"] is False else "na"), "detail": reasons["image"]},
            "answerable":  {"status": "pass" if dims["answer"] is True else ("fail" if dims["answer"] is False else "na"), "detail": reasons["answer"]},
            "answer_knowledge": {"status": "na", "detail": ""},
            "audio":       {"status": "pass" if dims["audio"] is True else ("fail" if dims["audio"] is False else "na"), "detail": reasons["audio"]},
            "flow":        {"status": "pass" if overall is True else ("fail" if overall is False else "na"), "detail": "流程无卡顿" if overall is True else ""},
            "report":      {"status": "na", "detail": ""},
        },
        # ★ 题目内容展示文字
        "stem": stem_text,
        "options": option_text,
        "script_answer": "",
        "note": "",
        "human_label": None,
        "human_note": "",
        "timestamp": datetime.now().isoformat(),
    }
    _inspection_state["current_question_idx"] = qidx
    # ★ 每 10 题保存一次到文件（供错题溯源 trace 页面读取，避免只存内存导致 trace 看不到）
    if total % 10 == 0:
        try:
            _save_inspection_state()
        except Exception:
            pass
    # ★ 检测过程中出现错题 → 触发实时错题报告（前端「📑 查看报告」实时更新）
    if overall is False:
        try:
            _live_regen_error_report()
        except Exception:
            pass


def _fmt_reason(field: str, diff: str, verdict: str) -> str:
    """生成自然语言原因描述，让检查人员一眼看懂"""
    prefix = {"题干": "界面题干", "选项": "可选项", "音频": "音频控件", "作答": "作答元素"}.get(field, field)
    if diff and len(diff) < 50:
        return f"{prefix}: {diff}"
    if verdict == "通过":
        return f"{prefix}正常 ✓"
    return f"{prefix}异常 ✗"


# ===== 快速检查(从当前页开始) =====
def run_quick_inspect_task(docx_file: str = "", unit: int = 0):
    """⚡ 快速检查：从手机当前页面直接开始逐题巡检（跳过启动/关广告/登录/导航）

    前提: 用户已手动将手机调到题目页 (如 听力专项-基础巩固-第1题)
    流程: 截图 → AI六维审查(若提供脚本) → 写入巡检状态 → 点选项 → 点底部按钮x2 → 循环
    坐标自动缩放 (适配不同分辨率手机)
    """
    try:
        set_running("quick_inspect")
        config = load_config()
        adb = get_adb()

        # 清空旧巡检状态
        _inspection_state["questions"] = {}
        _inspection_state["workflow_steps"] = []
        _inspection_state["current_question_idx"] = 0
        _save_inspection_state()

        # 加载脚本 + AI审查agent (可选)
        agent = None
        if docx_file:
            docx_path = str(UPLOAD_DIR / docx_file)
            if Path(docx_path).exists():
                try:
                    from src.review_agent import ReviewAgent, ReviewConfig
                    cfg = ReviewConfig(docx_path=docx_path, unit=int(unit or 0),
                                       screenshot_dir=str(PROJECT_ROOT / "screenshots"),
                                       verbose=False)
                    agent = ReviewAgent(cfg)
                    log_msg(f"📄 脚本已加载: {docx_file} (共{len(agent.script_questions)}题) → AI六维审查开启", "success")
                except Exception as e:
                    log_msg(f"⚠ 脚本加载失败: {e}，降级为仅截图", "warning")
            else:
                log_msg(f"⚠ 脚本不存在: {docx_path}，降级为仅截图", "warning")
        else:
            log_msg("⚡ 未指定脚本，仅截图+推进（选脚本可开启AI六维审查）", "info")

        log_msg("⚡ 快速检查启动（从当前页面开始，跳过导航）", "success")
        time.sleep(1)

        # 1. 尝试检测当前题目进度 (如 1/40)
        elements = adb.dump_ui(retries=2)
        cur = None
        total_q = 0
        for e in elements:
            m = re.match(r'^(\d+)/(\d+)$', (e.text or "").strip())
            if m:
                cur = int(m.group(1))
                total_q = int(m.group(2))
                break

        if cur:
            log_msg(f"✅ 检测到题目进度 {cur}/{total_q or 40}，开始逐题检查", "success")
        else:
            log_msg("⚠ 未检测到题目进度(如 1/40)。若您已在题目页将正常检查，", "warning")
            log_msg("   否则请先手动进入题目页再点「快速检查」", "warning")
            time.sleep(2)

        # 2. 逐题巡检
        last = 0
        q_count = 0
        for step in range(80):
            if not task_status["running"]:
                break

            idx = cur if cur else (last + 1)

            # 截图 (保存到 outputs/web, 前端可预览)
            shot = f"q{idx:02d}.png"
            adb.screenshot(shot)
            log_msg(f"📸 Q{idx} 已截图", "success")
            update_progress(idx, total_q or 40, f"Q{idx}")
            q_count += 1

            # ★ 实时题型识别（从 UI 判断当前是什么题）
            detected = None
            try:
                from src.type_detector import TypeDetector
                q_els = adb.dump_ui(retries=2)
                detected = TypeDetector(verbose=False).detect(q_els)
                if detected and detected.type_1 != "未知":
                    log_msg(f"  🔍 实时题型: {detected.describe()} (conf={detected.confidence:.0%})", "info")
            except Exception:
                detected = None

            # ====== AI 六维审查 (若有脚本) ======
            ai = {
                "ai_stem": None, "ai_content": None, "ai_image": None,
                "ai_answer": None, "ai_audio": None, "ai_post_error": None,
                "overall_passed": None, "overall_score": None,
                "stem_reason": "", "content_reason": "", "image_reason": "",
                "answer_reason": "", "audio_reason": "", "post_error_reason": "",
                "question_type": detected.full_type if detected and detected.full_type else "快速检查",
                "script_answer": "", "stem": "",
            }
            if agent:
                try:
                    # ★ 修复：截图保存在 screenshots/（adb.screenshot），
                    #   原代码读 outputs/web/ 导致找不到图 → 降级纯文字模式
                    shot_path = str(PROJECT_ROOT / "screenshots" / shot)
                    matching = [q for q in agent.script_questions if q.global_idx == idx]
                    if matching:
                        script_q = matching[0]
                        r = agent._review_one(script_q, shot_path, detected=detected)
                        ai = {
                            "ai_stem": r.stem_check.passed,
                            "ai_content": r.content_check.passed,
                            "ai_image": r.image_check.passed,
                            "ai_answer": r.answer_check.passed,
                            "ai_audio": r.audio_check.passed if r.audio_check is not None else None,
                            "ai_post_error": r.post_error_check.passed if r.post_error_check is not None else None,
                            "overall_passed": r.overall_passed,
                            "overall_score": round(r.overall_score, 2),
                            "stem_reason": r.stem_check.details[0][:100] if r.stem_check.details else "",
                            "content_reason": r.content_check.details[0][:100] if r.content_check.details else "",
                            "image_reason": r.image_check.details[0][:100] if r.image_check.details else "",
                            "answer_reason": r.answer_check.details[0][:100] if r.answer_check.details else "",
                            "audio_reason": r.audio_check.details[0][:100] if r.audio_check and r.audio_check.details else "",
                            "post_error_reason": r.post_error_check.details[0][:100] if r.post_error_check and r.post_error_check.details else "",
                            "question_type": (detected.full_type if detected and detected.full_type else (script_q.type_2 or "快速检查")),
                            "script_answer": script_q.answer or "",
                            "stem": (detected.stem if detected and detected.stem else script_q.stem)[:60] if (detected.stem if detected and detected.stem else script_q.stem) else "",
                        }
                        log_msg(f"  AI: 题干{ai['ai_stem']} 内容{ai['ai_content']} 配图{ai['ai_image']} 作答{ai['ai_answer']} 音频{ai['ai_audio']} 答错后{ai['ai_post_error']} 得分{ai['overall_score']}")
                    else:
                        log_msg(f"  ⚠ 脚本中无 Q{idx}，跳过AI审查", "warning")
                except Exception as e:
                    log_msg(f"  ⚠ AI审查失败 Q{idx}: {e}", "warning")

            # ====== 写入巡检状态 → 前端中栏显示 ======
            qid = f"quick-Q{idx:02d}"
            _inspection_state["questions"][qid] = {
                "idx": idx,
                "total": total_q or 40,
                "question_type": ai["question_type"],
                "screenshot": shot,
                "progress": f"{idx}/{total_q or 40}",
                "ai_stem": ai["ai_stem"], "ai_content": ai["ai_content"],
                "ai_image": ai["ai_image"], "ai_answer": ai["ai_answer"],
                "ai_audio": ai["ai_audio"], "ai_post_error": ai["ai_post_error"],
                "overall_passed": ai["overall_passed"],
                "overall_score": ai["overall_score"],
                "stem_reason": ai["stem_reason"], "content_reason": ai["content_reason"],
                "image_reason": ai["image_reason"], "answer_reason": ai["answer_reason"],
                "audio_reason": ai["audio_reason"], "post_error_reason": ai["post_error_reason"],
                "script_answer": ai["script_answer"],
                "stem": ai["stem"],
                "note": "",
                "human_label": None,
                "human_note": "",
                "timestamp": datetime.now().isoformat(),
            }
            _inspection_state["current_question_idx"] = idx
            _save_inspection_state()

            # 点选项: 优先找UI元素，找不到则点默认右侧选项区(缩放)
            elements = adb.dump_ui(retries=2)
            clicked_option = False
            for e in elements:
                if e.clickable and 700 < e.bounds[1] < 1700:
                    if e.center[0] > 450:
                        adb.tap(e.center[0], e.bounds[1] + 30)
                    else:
                        adb.tap((e.bounds[0] + e.bounds[2]) // 2,
                                (e.bounds[1] + e.bounds[3]) // 2)
                    clicked_option = True
                    break
            if not clicked_option:
                tx, ty = sc(970, 1500)   # 默认点右侧选项区
                adb.tap(tx, ty)
            time.sleep(1)

            # 点底部按钮 x2 (检查答案 → 下一题)
            bx, by = sc(540, 2174)
            adb.tap(bx, by)
            time.sleep(1.5)
            adb.tap(bx, by)
            time.sleep(1.5)

            # 检测进度变化
            elements_after = adb.dump_ui(retries=2)
            new_cur = None
            new_total = 0
            for e in elements_after:
                m = re.match(r'^(\d+)/(\d+)$', (e.text or "").strip())
                if m:
                    new_cur = int(m.group(1))
                    new_total = int(m.group(2))
                    break

            if new_cur and new_cur == last and cur:
                log_msg(f"  ⚠ Q{new_cur} 进度未变，重试一次", "warning")
                time.sleep(2)
                adb.tap(bx, by)
                time.sleep(1.5)
                continue

            if new_cur:
                last = new_cur
                cur = new_cur
                total_q = new_total or total_q
                if new_cur >= (total_q or 40):
                    log_msg(f"✅ {new_cur}题完成!", "success")
                    break
            else:
                last += 1
                log_msg(f"  未检测到进度文字, 继续下一题 (Q{last})", "info")

        # 3. 保存报告
        out = PROJECT_ROOT / "outputs" / "questions" / "quick_report.json"
        out.parent.mkdir(parents=True, exist_ok=True)
        with open(out, "w", encoding="utf-8") as f:
            json.dump({"questions": list(_inspection_state["questions"].values())},
                      f, ensure_ascii=False, indent=2)
        log_msg(f"✅ 快速检查完成: 共 {q_count} 题", "success")

    except Exception as e:
        log_msg(f"❌ 异常: {e}", "error")
        import traceback
        traceback.print_exc()
    finally:
        set_done()


# ===== 快速检查(从当前页开始) =====
# ============================================================
# 错题溯源 / 透明评分 整合（复用同学C的 TraceEngine）
# ============================================================
_TRACE_ENGINE = None

def _get_trace_engine():
    """惰性加载同学C的 TraceEngine（缺 Pillow/文件时返回 None，不阻塞）。"""
    global _TRACE_ENGINE
    if _TRACE_ENGINE is None:
        try:
            from trace_engine import TraceEngine
            _TRACE_ENGINE = TraceEngine(screenshots_dir=str(PROJECT_ROOT / "screenshots"))
        except Exception:
            _TRACE_ENGINE = False
    return _TRACE_ENGINE or None


def _score_breakdown(qd: dict) -> dict:
    """透明评分（★ 统一维度，让前端展示一致）：

    ★ 主评分（一致性保证）：仅 3 个核心维度 → 题干/选项/作答
       通过数 / 已检查数 × 100。这 3 维是任何题型必检的（基础完整性）。
       其他维度（配图/音频/答错后/报告）作为"附加检查项"展示，不参与主评分。
    ★ 这样所有题"维度"统一，不会出现"3 维 vs 6 维"不一致。
    """
    # ★ 主评分：3 维（基础完整性，所有题统一必检）
    main_dims = [
        ("题干", qd.get("ai_stem"), qd.get("stem_reason"), "medium"),
        ("选项", qd.get("ai_content"), qd.get("content_reason"), "high"),
        ("作答", qd.get("ai_answer"), qd.get("answer_reason"), "high"),
    ]
    # ★ 附加检查：4 维（按题型可选，不算主分）
    extra_dims = [
        ("配图", qd.get("ai_image"), qd.get("image_reason"), "high"),
        ("音频", qd.get("ai_audio"), qd.get("audio_reason"), "low"),
        ("答错后", qd.get("ai_post_error"), qd.get("post_error_reason"), "low"),
        ("报告", qd.get("ai_report"), qd.get("report_reason"), "low"),
    ]
    checked = [d for d in main_dims if d[1] is not None]
    passed = [d for d in checked if d[1] is True]
    total = len(checked)
    score = round(len(passed) / total * 100) if total else 0
    return {
        "score": score,
        "checked_count": total,
        "passed_count": len(passed),
        "failed_count": total - len(passed),
        # ★ 公式固定为 "3 维基础完整性检查"，让检查人员一眼明白评分逻辑
        "formula": f"{len(passed)}/{total}×100={score}（基础完整性 3 维评分）" if total else "未检查，无法评分",
        "dims": [{"name": d[0], "passed": d[1], "reason": d[2] or "", "severity": d[3]} for d in checked],
        # ★ 附加检查项（前端可单独展示，不参与主分）
        "extra_dims": [{"name": d[0], "passed": d[1], "reason": d[2] or "", "severity": d[3]}
                      for d in extra_dims if d[1] is not None],
    }


def _ensure_marked(qid: str, qd: dict, checks: list) -> str:
    """生成红框标注图（同学C的 draw_mark），输出到 outputs/web/，已存在则复用（缓存）。
    返回相对文件名（前端拼 /api/screenshot/<name>）；失败返回空串。"""
    try:
        shot = qd.get("screenshot", "")
        if not shot or not checks:
            return ""
        out_name = f"marked_{shot}"
        out_path = PROJECT_ROOT / "outputs" / "web" / out_name
        if out_path.exists():
            return out_name
        eng = _get_trace_engine()
        if not eng:
            return ""
        eng.draw_mark(shot, checks, str(out_path))
        return out_name if out_path.exists() else ""
    except Exception:
        return ""


def _enrich_inspect_state(state: dict):
    """给每道错题附加溯源/评分/红框图；通过题仅附评分（供参考）。"""
    eng = _get_trace_engine()
    for qid, qd in (state.get("questions") or {}).items():
        qd["_score"] = _score_breakdown(qd)
        if qd.get("overall_passed") is not False:
            continue  # 通过题不用溯源
        if eng is None:
            continue
        try:
            trace = eng.generate(qid, qd)
            qd["_trace"] = trace
            qd["_marked"] = _ensure_marked(qid, qd, trace.get("checks", []))
        except Exception:
            pass


@app.route("/api/inspect/state", methods=["GET"])
def api_inspect_state():
    """获取当前巡检状态 (流式)。
    ★ 集成同学C的溯源引擎：对每道错题附加 _trace（维度/原因/建议/严重度/坐标）、
    _score（透明评分明细，说明评分怎么来的）、_marked（红框标注图URL）。
    """
    # 浅拷贝一层，避免污染 _inspection_state（_score/_trace/_marked 只进响应）
    # ★ 必须带全 version/unit/stage/docx/live_report_path（前端"查看报告"/分组展示依赖）
    state = {
        "questions": {k: dict(v) for k, v in (_inspection_state.get("questions") or {}).items()},
        "workflow_steps": _inspection_state.get("workflow_steps", []),
        "current_question_idx": _inspection_state.get("current_question_idx", 0),
        "version": _inspection_state.get("version", ""),
        "unit": _inspection_state.get("unit", ""),
        "stage": _inspection_state.get("stage", ""),
        "docx": _inspection_state.get("docx", ""),
        "live_report_path": _inspection_state.get("live_report_path", ""),
    }
    try:
        _enrich_inspect_state(state)
    except Exception:
        pass
    return jsonify(state)


@app.route("/api/inspect/question-result", methods=["POST"])
def api_inspect_question_result():
    """接收AI对一道题的判断结果, 推送审查证据到前端日志"""
    data = request.get_json() or {}
    qid = data.get("qid", "")
    if not qid:
        return jsonify({"error": "缺少题号"}), 400

    _inspection_state["questions"][qid] = {
        **data,
        "human_label": None,  # 等待人工标注
        "human_note": "",
        "timestamp": datetime.now().isoformat(),
    }
    _inspection_state["current_question_idx"] = data.get("idx", 0)
    _save_inspection_state()

    # ★ 借鉴队友：该题不通过 → 实时重生成错题报告（前端「📑 查看报告」弹窗即时更新）
    if not data.get("overall_passed"):
        _live_regen_error_report()

    # ★ 推送审查证据到前端日志(收集各维度的 evidence)
    idx = data.get("idx", "?")
    overall = "通过" if data.get("overall_passed") else "不通过"
    level = "success" if data.get("overall_passed") else "error"
    all_evidence = []
    for dim in ("stem", "content", "image", "answer", "audio", "post_error"):
        check = data.get(f"{dim}_check", {})
        if isinstance(check, dict) and check.get("evidence"):
            all_evidence.extend(check["evidence"])
    log_msg(f"📋 Q{idx} 审查{overall} | 方法:{data.get('method','')}",
            level, evidence=all_evidence if all_evidence else None)
    return jsonify({"success": True, "qid": qid})


@app.route("/api/inspect/human-label", methods=["POST"])
def api_inspect_human_label():
    """人工标注某道题的对错 + 加入反馈循环"""
    data = request.get_json() or {}
    qid = data.get("qid", "")
    label = data.get("label", "")  # "通过" / "不通过"
    note = data.get("note", "")

    if not qid or label not in ["通过", "不通过"]:
        return jsonify({"error": "参数错误"}), 400

    # 更新巡检状态
    if qid in _inspection_state["questions"]:
        _inspection_state["questions"][qid]["human_label"] = label
        _inspection_state["questions"][qid]["human_note"] = note
    _save_inspection_state()

    # 加入反馈循环 (审查智能体学习)
    q_data = _inspection_state["questions"][qid]
    ai_judgment = "通过" if q_data.get("overall_passed") else "不通过"
    try:
        from src.feedback_loop import FeedbackSample
        store = FeedbackStore()
        store.add(FeedbackSample(
            question_id=qid,
            check_dimension=q_data.get("check_dimension", "all"),
            human_judgment=label,
            ai_judgment=ai_judgment,
            ai_reason=str(q_data.get("ai_reason", ""))[:200],
            human_note=note,
            question_type=q_data.get("question_type", ""),
            screenshot=q_data.get("screenshot", ""),
        ))
        return jsonify({
            "success": True,
            "qid": qid,
            "feedback_stored": True,
            "stats": store.get_stats(),
        })
    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        return jsonify({"success": True, "qid": qid, "feedback_error": str(e)})


@app.route("/api/inspect/workflow-step", methods=["POST"])
def api_inspect_workflow_step():
    """记录流程步骤状态"""
    data = request.get_json() or {}
    step_name = data.get("step", "")
    status = data.get("status", "")  # "running" / "done" / "error"
    detail = data.get("detail", "")

    # 替换或添加步骤
    found = False
    for s in _inspection_state["workflow_steps"]:
        if s["step"] == step_name:
            s["status"] = status
            s["detail"] = detail
            s["updated"] = datetime.now().isoformat()
            found = True
            break
    if not found:
        _inspection_state["workflow_steps"].append({
            "step": step_name,
            "status": status,
            "detail": detail,
            "updated": datetime.now().isoformat(),
        })
    _save_inspection_state()
    return jsonify({"success": True})


@app.route("/api/inspect/reset", methods=["POST"])
def api_inspect_reset():
    """清空当前巡检状态"""
    _inspection_state["questions"] = {}
    _inspection_state["workflow_steps"] = []
    _inspection_state["current_question_idx"] = 0
    _inspection_state["live_report_path"] = ""   # ★ 清空实时报告路径
    _save_inspection_state()
    return jsonify({"success": True})


@app.route("/api/modules")
def api_modules():
    """获取可用模块列表"""
    return jsonify({
        "教材精学": {
            "课本点读(左)": (203, 1191),
            "课本点读(中)": (540, 1191),
            "巧记单词":     (876, 1191),
            "语音评测":     (203, 1358),
        },
        "专项突破 (可见)": {
            "听课文":    (161, 1792),
            "课文动画":  (414, 1792),
            "基础训练":  (666, 1792),
            "一课一练":  (919, 1792),
            "课文配音":  (161, 2033),
            "口语训练":  (414, 2033),
            "复习回顾":  (666, 2033),
            "全脑记词":  (919, 2033),
        },
        "专项突破 (需滚动)": {
            "单元自检":     (414, 746),
            "单元练习计划": (666, 746),
        },
        "📢 听力专项": {
            "听力专项_五上": (414, 1250),
            "听力专项_五下": (414, 950),
            "听力专项_六上": (414, 1150),
            "听力专项_六下": (414, 1050),
        },
    })


@app.route("/api/grades", methods=["POST"])
def api_grades():
    """检测当前版本的可用年级列表（含新/老教材标记）"""
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409

    config = load_config()
    adb = get_adb()

    try:
        # 确保在主页，打开年级选择弹窗
        adb.tap(*sc(346, 275))  # 切换器
        time.sleep(3)

        elements = adb.dump_ui()
        if not elements:
            return jsonify({"error": "UI dump失败", "grades": []})

        # 解析年级数据
        grades = []
        seen = set()
        for elem in elements:
            text = elem.text
            if not text:
                continue
            # 检测是否含年级/册文字
            is_book = any(k in text for k in ['年级', '册'])
            is_new_tag = '新教材' in text
            if is_book or is_new_tag:
                if text not in seen:
                    seen.add(text)
                    # 找容器坐标（可点击区域）
                    container = None
                    for ce in elements:
                        if ce.clickable and elem.center[0] >= ce.bounds[0] and elem.center[0] <= ce.bounds[2] \
                                and elem.center[1] >= ce.bounds[1] and elem.center[1] <= ce.bounds[3]:
                            container = list(ce.center)
                            break

                    grades.append({
                        "text": text,
                        "center": list(elem.center),
                        "container": container,
                        "is_new_tag": is_new_tag,
                        "is_book": is_book,
                    })

        # 关闭弹窗
        adb.press_back()
        time.sleep(1)

        # 结构化输出
        book_grades = [g for g in grades if g["is_book"]]
        log_msg(f"✅ 检测到 {len(book_grades)} 个年级选项")

        return jsonify({"grades": grades, "book_grades": book_grades})

    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        log_msg(f"❌ 年级检测失败: {e}", "error")
        return jsonify({"error": str(e), "grades": []})


@app.route("/api/screenshot/<filename>")
def api_screenshot(filename):
    """获取截图"""
    screenshot_dir = PROJECT_ROOT / "outputs" / "web"
    if (screenshot_dir / filename).exists():
        return send_from_directory(str(screenshot_dir), filename)
    # 也查 outputs/screenshots
    alt_dir = PROJECT_ROOT / "outputs" / "screenshots"
    if (alt_dir / filename).exists():
        return send_from_directory(str(alt_dir), filename)
    # ★ 截图实际保存目录（快速检查 qXX.png / 答错截图 wrong_qXX.png 都在这）
    shot_dir = PROJECT_ROOT / "screenshots"
    if (shot_dir / filename).exists():
        return send_from_directory(str(shot_dir), filename)
    return jsonify({"error": "未找到截图"}), 404


@app.route("/api/screenshot/latest")
def api_screenshot_latest():
    """获取最新截图"""
    screenshot_dir = PROJECT_ROOT / "outputs" / "web"
    png_files = sorted(screenshot_dir.glob("*.png"), key=os.path.getmtime, reverse=True)
    if png_files:
        return send_from_directory(str(screenshot_dir), png_files[0].name)
    return jsonify({"error": "无截图"}), 404


# ★ 常驻实时截图线程：web_server 启动即开始，每 3 秒截一张 outputs/web/live.png，
#   供前端「手机画面」实时预览。设备未连接时静默重试（10s 间隔），不阻塞启动。
#   （原实现只在任务运行时启动截图循环，任务一结束前端就显示最后一张旧图）
_LIVE_SHOT = {"dev": None, "fail": 0}

def _start_live_screenshot_daemon():
    def _loop():
        shot_dir = PROJECT_ROOT / "outputs" / "web"
        try:
            shot_dir.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass
        while True:
            try:
                if _LIVE_SHOT["dev"] is None:
                    try:
                        _LIVE_SHOT["dev"] = _connect_device()
                        _LIVE_SHOT["fail"] = 0
                    except Exception:
                        _LIVE_SHOT["dev"] = None
                        _LIVE_SHOT["fail"] += 1
                if _LIVE_SHOT["dev"] is not None:
                    try:
                        _LIVE_SHOT["dev"].screenshot(str(shot_dir / "live.png"))
                        _LIVE_SHOT["fail"] = 0
                    except Exception:
                        _LIVE_SHOT["dev"] = None
                        _LIVE_SHOT["fail"] += 1
                time.sleep(3 if _LIVE_SHOT["fail"] == 0 else 10)
            except Exception:
                time.sleep(10)
    threading.Thread(target=_loop, daemon=True).start()


@app.route("/api/log")
def api_log():
    """获取日志"""
    return jsonify(task_status["log"])


@app.route("/api/hardware-info")
def api_hardware_info():
    """获取手机信息"""
    config = load_config()
    info = {"serial": config.device.serial, "screen": "1080x2400"}
    try:
        r = sp.run(["C:/Users/bunana/AppData/Local/Microsoft/WinGet/Packages/Google.PlatformTools_Microsoft.Winget.Source_8wekyb3d8bbwe/platform-tools/adb.exe", "-s", config.device.serial, "shell", "getprop", "ro.product.model"],
                   capture_output=True, text=True, timeout=5,
                   encoding="utf-8", errors="replace")
        if r.returncode == 0:
            info["model"] = r.stdout.strip()
        r2 = sp.run(["C:/Users/bunana/AppData/Local/Microsoft/WinGet/Packages/Google.PlatformTools_Microsoft.Winget.Source_8wekyb3d8bbwe/platform-tools/adb.exe", "-s", config.device.serial, "shell", "getprop", "ro.product.brand"],
                    capture_output=True, text=True, timeout=5,
                    encoding="utf-8", errors="replace")
        if r2.returncode == 0:
            info["brand"] = r2.stdout.strip()
    except Exception:
        pass
    return jsonify(info)


def _get_current_version_from_config() -> str:
    """获取当前目标版本（顶栏展示用；实际由 UI 下拉框决定，此处不再硬编码假值）"""
    return ""


# ============================================================
# 🔵 审查智能体 路由区 — 脚本上传 + 知识库 + 审查结果
# ============================================================

UPLOAD_DIR = PROJECT_ROOT / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)


@app.route("/api/upload", methods=["POST"])
def api_upload():
    """上传 DOCX/DOC 脚本文件（仅保存到 uploads 目录，供审查智能体使用）

    前端 uploadFile() 调用此接口上传文件
    """
    if "file" not in request.files:
        return jsonify({"error": "未选择文件"}), 400
    file = request.files["file"]
    ext = file.filename.lower().rsplit(".", 1)[-1] if "." in file.filename else ""
    if ext not in ("docx", "doc"):
        return jsonify({"error": "仅支持 .docx / .doc 文件"}), 400
    save_path = UPLOAD_DIR / file.filename
    file.save(str(save_path))
    log_msg(f"📤 已上传脚本: {file.filename}", "success")
    return jsonify({"success": True, "filename": file.filename})


@app.route("/api/upload-docx", methods=["POST"])
def api_upload_docx():
    """上传 DOCX/DOC 脚本文件, 自动导入知识库"""
    if "file" not in request.files:
        return jsonify({"error": "未选择文件"}), 400
    file = request.files["file"]
    ext = file.filename.lower().rsplit(".", 1)[-1] if "." in file.filename else ""
    if ext not in ("docx", "doc"):
        return jsonify({"error": "仅支持 .docx / .doc 文件"}), 400
    save_path = UPLOAD_DIR / file.filename
    file.save(str(save_path))
    try:
        from src.knowledge_base import KnowledgeBase
        kb = KnowledgeBase()
        stats = kb.add_bulk_from_docx(str(save_path))
        return jsonify({
            "success": True, "filename": file.filename,
            "knowledge_stats": {
                "questions_parsed": stats.get("questions", 0),
                "vocab_added": stats.get("vocab", 0),
                "patterns_added": stats.get("patterns", 0),
            },
            "message": f"已导入 {stats.get('questions',0)} 题到知识库",
        })
    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        return jsonify({"error": f"解析失败: {str(e)}"}), 500


@app.route("/api/order/parse", methods=["POST"])
def api_order_parse():
    """解析上传的任务清单文件 → 提取文本

    支持格式：.txt / .docx / .xlsx
    返回: {"text": ..., "filename": ..., "format": "txt|docx|xlsx"}
    前端拿到 text 后复用 applyOrder() 做版本/年级/模块识别
    """
    if "file" not in request.files:
        return jsonify({"error": "未选择文件"}), 400
    f = request.files["file"]
    name = f.filename or "task.txt"
    ext = name.lower().rsplit(".", 1)[-1] if "." in name else "txt"
    try:
        data = f.read()
        if ext == "txt":
            text = data.decode("utf-8", errors="replace")
        elif ext in ("docx",):
            import io as _io
            from docx import Document
            doc = Document(_io.BytesIO(data))
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for tb in doc.tables:
                for row in tb.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
        elif ext in ("xlsx", "xls"):
            import io as _io
            import openpyxl
            wb = openpyxl.load_workbook(_io.BytesIO(data), data_only=True)
            rows = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if cells:
                        rows.append(" | ".join(cells))
            text = "\n".join(rows)
        elif ext == "doc":
            return jsonify({
                "error": "老格式 .doc 暂不支持直接解析，请另存为 .docx 或 .txt 再上传",
                "filename": name,
            }), 400
        else:
            return jsonify({"error": f"不支持的文件格式 .{ext}（支持 .txt/.docx/.xlsx）"}), 400

        text = (text or "").strip()
        if not text:
            return jsonify({"error": "文件内容为空或未提取到文字", "filename": name}), 400
        log_msg(f"📋 已解析任务清单: {name}（{ext}，{len(text)} 字符）", "success")
        return jsonify({"text": text, "filename": name, "format": ext})
    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
        return jsonify({"error": "已停止"}), 500
    except Exception as e:
        log_msg(f"⚠ 任务清单解析失败 {name}: {e}", "warning")
        return jsonify({"error": f"解析失败: {str(e)[:100]}", "filename": name}), 500


@app.route("/api/upload/list")
def api_upload_list():
    """列出已上传的 DOCX 文件"""
    files = []
    for f in sorted(UPLOAD_DIR.glob("*"), key=os.path.getmtime, reverse=True):
        if f.suffix.lower() not in (".docx", ".doc"): continue
        if "_converted" in f.stem: continue  # 跳过转换生成的临时文件
        files.append({
            "name": f.name,
            "size": f.stat().st_size,
            "mtime": datetime.fromtimestamp(f.stat().st_mtime).isoformat(),
        })
    return jsonify({"files": files})


@app.route("/api/knowledge/status")
def api_knowledge_status():
    """知识库状态"""
    try:
        from src.knowledge_base import KnowledgeBase
        kb = KnowledgeBase()
        summary = kb.summary()
        total_entries = sum(s["vocab"] + s["patterns"] for s in summary)
        return jsonify({
            "total_entries": total_entries,
            "grades": [s["key"] for s in summary],
            "detail": summary,
        })
    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        return jsonify({"total_entries": 0, "error": str(e)})


_review_results = {}

@app.route("/api/review/run", methods=["POST"])
def api_review_run():
    """运行审查智能体"""
    data = request.get_json() or {}
    docx_file = data.get("docx", "")
    unit = data.get("unit", 0)
    stage = data.get("stage", "")
    docx_path = ""
    if docx_file:
        docx_path = str(UPLOAD_DIR / docx_file)
        if not Path(docx_path).exists():
            return jsonify({"error": f"文件 {docx_file} 不存在"}), 404
    else:
        files = [f for f in sorted(UPLOAD_DIR.glob("*"), key=os.path.getmtime, reverse=True)
                 if f.suffix.lower() in (".docx", ".doc")]
        if files: docx_path = str(files[0])
        else: return jsonify({"error": "未找到脚本文件"}), 404
    try:
        from src.review_agent import ReviewAgent, ReviewConfig
        cfg = ReviewConfig(docx_path=docx_path, unit=unit, stage=stage,
                           screenshot_dir=str(PROJECT_ROOT / "screenshots"), verbose=False)
        agent = ReviewAgent(cfg)
        results = agent.review()
        global _review_results
        _review_results = {
            "timestamp": datetime.now().isoformat(),
            "docx": Path(docx_path).name,
            "total": len(results),
            "passed": sum(1 for r in results if r.overall_passed),
            "avg_score": round(sum(r.overall_score for r in results) / len(results), 2) if results else 0,
            "results": [r.to_dict() for r in results],
        }
        return jsonify(_review_results)
    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500


@app.route("/api/review/results")
def api_review_results():
    return jsonify(_review_results or {"results": [], "total": 0})


@app.route("/api/review/export", methods=["POST"])
def api_review_export():
    try:
        py_content = ""
        if _review_results:
            r = _review_results
            py_content = f"# 英语宝审查报告\n\n生成时间: {r['timestamp']}\n脚本: {r['docx']}\n"
            py_content += f"总题数: {r['total']} | 通过: {r['passed']}/{r['total']} | 综合得分: {r['avg_score']}\n\n"
            py_content += "| # | 题型 | 题干 | 内容 | 配图 | 作答 | 总评 |\n|---|---|---|---|---|---|---|\n"
            for rr in r.get("results", []):
                def ic(p): return "Y" if p else "N"
                py_content += f"| Q{rr['idx']:02d} | {rr['type'][:10]} | {ic(rr['stem']['passed'])} | {ic(rr['content']['passed'])} | {ic(rr['image']['passed'])} | {ic(rr['answer']['passed'])} | {'PASS' if rr['overall_passed'] else 'FAIL'} ({rr['overall_score']}) |\n"
        return jsonify({"content": py_content})
    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/report/teacher")
def api_report_teacher():
    """生成教师可读的可视化HTML报告"""
    from datetime import datetime

    st = _inspection_state
    qs = st.get("questions", {})

    total = len(qs)
    passed = sum(1 for q in qs.values() if q.get("overall_passed"))
    labeled = sum(1 for q in qs.values() if q.get("human_label"))
    acc = 0
    if labeled > 0:
        corr = sum(1 for q in qs.values()
                   if q.get("human_label") == "通过" and q.get("overall_passed")
                   or q.get("human_label") == "不通过" and not q.get("overall_passed"))
        acc = int(corr / labeled * 100)

    # 构建HTML
    questions_html = ""
    for qid in sorted(qs.keys(), key=lambda x: int(re.findall(r'Q(\d+)', x)[0]) if re.findall(r'Q(\d+)', x) else 0):
        q = qs[qid]
        dims = [
            ("题干", q.get("ai_stem"), q.get("stem_reason", "")),
            ("内容", q.get("ai_content"), q.get("content_reason", "")),
            ("配图", q.get("ai_image"), q.get("image_reason", "")),
            ("作答", q.get("ai_answer"), q.get("answer_reason", "")),
        ]
        dim_html = "".join(
            f'<div class="dim {"pass" if p else "fail"}">'
            f'<div class="dim-icon">{ "✅" if p else "❌" }</div>'
            f'<div class="dim-label">{n}</div>'
            f'<div class="dim-reason">{r[:60] if r else "-"}</div>'
            f'</div>'
            for n, p, r in dims
        )
        hl = q.get("human_label", "")
        hl_html = f'<span class="hl {"hl-pass" if hl=="通过" else "hl-fail"}">人工: {hl}</span>' if hl else '<span class="hl hl-pending">待标注</span>'
        shot = q.get("screenshot", "")
        shot_html = f'<img src="/api/screenshot/{shot}" onclick="window.open(this.src)" style="cursor:pointer">' if shot else '<div class="no-shot">无截图</div>'
        questions_html += f'''
        <div class="q-card">
            <div class="q-hdr">
                <span class="q-id">Q{str(q.get("idx","?")).zfill(2)}</span>
                <span class="q-type">{q.get("question_type","?")}</span>
                <span class="q-ans">脚本答案: {q.get("script_answer","-")}</span>
                {hl_html}
                <span class="q-ovr {"ovr-pass" if q.get("overall_passed") else "ovr-fail"}">{"通过" if q.get("overall_passed") else "不通过"} ({q.get("overall_score",0):.2f})</span>
            </div>
            <div class="q-bdy">
                <div class="q-info">
                    <div class="info-line"><b>题干:</b> {q.get("stem","-")[:80]}</div>
                    <div class="info-line"><b>录音:</b> {q.get("recording","-")[:60]}</div>
                    <div class="info-line"><b>答案:</b> {q.get("script_answer","-")}</div>
                    {f'<div class="info-line kb"><b>知识库:</b> {q.get("knowledge_check","")[:80]}</div>' if q.get("knowledge_check") else ""}
                </div>
                <div class="q-shot">{shot_html}</div>
            </div>
            <div class="q-dims">{dim_html}</div>
        </div>'''

    docx_name = st.get("docx", st.get("version", "未知"))
    version = st.get("version", "未知")
    unit = st.get("unit", "?")
    stage = st.get("stage", "?")

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="utf-8"><title>审查报告</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:"Microsoft YaHei","PingFang SC",sans-serif;background:#f0f2f5;color:#1e293b;padding:20px;font-size:15px;line-height:1.8}}
.hdr{{background:linear-gradient(135deg,#4a6cf7,#2952e8);color:#fff;border-radius:14px;padding:24px 32px;margin-bottom:20px}}
.hdr h1{{font-size:24px}}
.hdr .meta{{display:flex;gap:30px;margin-top:12px;font-size:14px;opacity:.9}}
.stats{{display:flex;gap:16px;padding:16px 20px;background:#fff;border-radius:12px;box-shadow:0 1px 6px rgba(0,0,0,.06);margin-bottom:20px;flex-wrap:wrap}}
.stat{{display:flex;align-items:center;gap:8px;font-size:15px}}
.stat .n{{font-weight:700;font-size:22px}}
.stat .n.g{{color:#16a34a}}
.stat .n.r{{color:#dc2626}}
.stat .n.b{{color:#4a6cf7}}
.q-card{{background:#fff;border-radius:12px;margin-bottom:12px;border:1px solid #eef0f5;overflow:hidden}}
.q-hdr{{display:flex;align-items:center;gap:12px;padding:10px 16px;background:#fafbfc;border-bottom:1px solid #f0f2f6;flex-wrap:wrap}}
.q-id{{font-size:20px;font-weight:700;color:#4a6cf7;min-width:40px}}
.q-type{{font-size:12px;padding:2px 10px;border-radius:6px;background:#eef2ff;color:#4a6cf7}}
.q-ans{{font-size:12px;color:#64748b}}
.q-ovr{{margin-left:auto;font-size:12px;padding:2px 10px;border-radius:6px;font-weight:600}}
.ovr-pass{{background:#dcfce7;color:#166534}}
.ovr-fail{{background:#fee2e2;color:#991b1b}}
.hl{{font-size:12px;padding:2px 10px;border-radius:6px}}
.hl-pass{{background:#bbf7d0;color:#14532d}}
.hl-fail{{background:#fecaca;color:#7f1d1d}}
.hl-pending{{background:#fef3c7;color:#92400e}}
.q-bdy{{display:grid;grid-template-columns:1fr 180px;gap:14px;padding:12px 16px}}
.q-info .info-line{{font-size:14px;padding:2px 0;color:#475569}}
.q-info .info-line b{{color:#334155}}
.q-info .kb{{color:#6b7280;font-size:12px;font-style:italic}}
.q-shot img{{width:100%;border-radius:8px;border:1px solid #eef0f5;max-height:160px;object-fit:contain;background:#f8fafc}}
.no-shot{{height:100px;display:flex;align-items:center;justify-content:center;color:#94a3b8;font-size:12px;background:#f8fafc;border-radius:8px}}
.q-dims{{display:grid;grid-template-columns:repeat(4,1fr);gap:6px;padding:0 16px 12px}}
.dim{{text-align:center;padding:8px;border-radius:8px;border:1px solid #eef0f5;background:#fafbfc}}
.dim.pass{{background:#f0fdf4;border-color:#bbf7d0}}
.dim.fail{{background:#fef2f2;border-color:#fecaca}}
.dim-icon{{font-size:20px}}
.dim-label{{font-size:13px;font-weight:600;color:#64748b}}
.dim-reason{{font-size:11px;color:#94a3b8;margin-top:2px;word-break:break-all}}
.actions{{text-align:center;margin:20px 0}}
.actions button{{padding:10px 24px;border:none;border-radius:8px;background:#4a6cf7;color:#fff;font-size:15px;cursor:pointer}}
.actions button:hover{{background:#2952e8}}
@media print{{.actions{{display:none}} .q-card{{break-inside:avoid}}}}
</style></head>
<body>
<div class="hdr">
    <h1>试题审查报告</h1>
    <div class="meta">
        <span>脚本: {docx_name}</span>
        <span>版本: {version}</span>
        <span>Unit {unit} · {stage}</span>
        <span>{datetime.now().strftime("%Y-%m-%d %H:%M")}</span>
    </div>
</div>
<div class="stats">
    <div class="stat">总题数: <span class="n b">{total}</span></div>
    <div class="stat">AI通过: <span class="n g">{passed}</span></div>
    <div class="stat">AI不通过: <span class="n r">{total - passed}</span></div>
    <div class="stat">已人工标注: <span class="n b">{labeled}</span></div>
    <div class="stat">人机一致率: <span class="n g">{acc}%</span></div>
</div>
{questions_html}
<div class="actions"><button onclick="window.print()">打印 / 保存为PDF</button></div>
</body></html>'''

    return Response(html, mimetype="text/html")


# ============================================================
# 听力专项 · 练习+测试 全流程自动化（新引擎）
# ============================================================

_AUDIO_RUNNER = None  # 后台线程


@app.route("/api/audio/run", methods=["POST"])
def api_audio_run():
    """启动听力专项自动化（练习+测试）
    请求: {"mode": "all"|"practice"|"test", "units": [1], "test_units": [1]}
    """
    global _AUDIO_RUNNER
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409

    data = request.get_json() or {}
    mode = data.get("mode", "all")
    units = data.get("units", [1])
    test_units = data.get("test_units", units)

    def _run():
        _register_task_thread()  # 记录线程 id，供"立即停止"注入异常
        try:
            set_running(f"听力专项[{mode}]")
            log_msg(f"启动听力专项 {mode} 模式: 练习单元{units} 测试单元{test_units}")
            sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
            from modules.听力专项 import run_module, run_test_module
            from common.tools import dismiss_global_popups, close_ad, ensure_grade, settle_ads
            import uiautomator2 as u2

            d = _connect_device()
            log_msg("设备已连接")

            # 关广告 + 确认年级
            # ★ 用 settle_ads 循环「检测→关闭→再检测」消除"点空/广告刚弹出时误点"竞态
            settle_ads(d, wait_total=12)
            ok = ensure_grade(d, "五年级上册", "湘少版")
            log_msg("年级确认: 湘少版 五年级上册" if ok else "⚠ 年级切换失败，继续尝试")

            q1 = q2 = 0
            if mode in ("all", "practice"):
                log_msg("▶ 开始练习部分...")
                q1 = run_module(d)
                log_msg(f"练习部分完成: {q1} 题")
            if mode in ("all", "test"):
                log_msg("▶ 开始测试部分...")
                q2 = run_test_module(d)
                log_msg(f"测试部分完成: {q2} 题")

            log_msg(f"✅ 听力专项全部完成: 练习{q1} + 测试{q2} = {q1+q2} 题")
            set_done()
        except SystemExit:
            log_msg("⏹ 任务已被立即停止", "warning")
            _cleanup_task_state()  # 统一收尾: 清running + 重置停止标志
        except Exception as e:
            log_msg(f"❌ 任务异常: {e}", "error")
            set_done()

    _AUDIO_RUNNER = threading.Thread(target=_run, daemon=True)
    _AUDIO_RUNNER.start()
    return jsonify({"status": "started", "mode": mode,
                    "units": units, "test_units": test_units})


# ============================================================
# 口语训练 自动化（新引擎）
# ============================================================

_ORAL_RUNNER = None  # 后台线程


@app.route("/api/oral/run", methods=["POST"])
def api_oral_run():
    """启动口语训练自动化
    请求: {"units": [1,2,3,4]}
    """
    global _ORAL_RUNNER
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409

    data = request.get_json() or {}
    units = data.get("units", [1])

    def _run():
        _register_task_thread()  # 记录线程 id，供"立即停止"注入异常
        try:
            set_running(f"口语训练")
            log_msg(f"启动口语训练: 单元{units}")
            sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
            from modules.口语训练 import run_module
            from common.tools import dismiss_global_popups, close_ad, ensure_grade
            import uiautomator2 as u2

            d = _connect_device()
            log_msg("设备已连接")

            # 关广告 + 确认年级
            for _ in range(3):
                dismiss_global_popups(d)
            close_ad(d)
            ok = ensure_grade(d, "五年级上册", "湘少版")
            log_msg("年级确认: 湘少版 五年级上册" if ok else "⚠ 年级切换失败，继续尝试")

            q = run_module(d)
            log_msg(f"✅ 口语训练完成: {q} 题")
            set_done()
        except SystemExit:
            log_msg("⏹ 任务已被立即停止", "warning")
            _cleanup_task_state()  # 统一收尾: 清running + 重置停止标志
        except Exception as e:
            log_msg(f"❌ 任务异常: {e}", "error")
            set_done()

    _ORAL_RUNNER = threading.Thread(target=_run, daemon=True)
    _ORAL_RUNNER.start()
    return jsonify({"status": "started", "units": units})


# ============================================================
# 单元自检 自动化（新引擎）
# ============================================================

_UNIT_RUNNER = None  # 后台线程


@app.route("/api/unit/run", methods=["POST"])
def api_unit_run():
    """启动单元自检自动化
    请求: {"units": [1,2,3,4]}
    """
    global _UNIT_RUNNER
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409

    data = request.get_json() or {}
    units = data.get("units", [1])
    # ★ 2026-08-26：年级/版本从请求读取（默认继承主页当前），不再写死五年级——用户切到六上后可测六上卷
    _g = data.get("grade", "") or ""
    _v = data.get("version", "") or ""

    def _run():
        _register_task_thread()  # 记录线程 id，供"立即停止"注入异常
        try:
            set_running(f"单元自检")
            log_msg(f"启动单元自检: 单元{units}")
            sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
            from modules.单元自检 import run_module
            from common.tools import dismiss_global_popups, close_ad, ensure_grade
            import uiautomator2 as u2

            d = _connect_device()
            log_msg("设备已连接")

            # 关广告 + 确认年级（仅当请求传了 grade 才切换；默认继承当前，不强制）
            for _ in range(3):
                dismiss_global_popups(d)
            close_ad(d)
            if _g:
                ok = ensure_grade(d, _g, _v or "湘少版")
                log_msg(f"年级确认: {_v or '湘少版'} {_g}" if ok else "⚠ 年级切换失败，继续尝试")

            q = run_module(d, units=units)
            log_msg(f"✅ 单元自检完成: {q} 题")
            set_done()
        except SystemExit:
            log_msg("⏹ 任务已被立即停止", "warning")
            _cleanup_task_state()  # 统一收尾: 清running + 重置停止标志
        except Exception as e:
            log_msg(f"❌ 任务异常: {e}", "error")
            set_done()

    _UNIT_RUNNER = threading.Thread(target=_run, daemon=True)
    _UNIT_RUNNER.start()
    return jsonify({"status": "started", "units": units})


# ============================================================
# 知识过关 自动化（新引擎）
# ============================================================

_KNOWLEDGE_RUNNER = None  # 后台线程


@app.route("/api/knowledge/run", methods=["POST"])
def api_knowledge_run():
    """启动知识过关自动化（重点词汇+重点句型）
    请求: {"units": [1]}
    """
    global _KNOWLEDGE_RUNNER
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409

    data = request.get_json() or {}
    units = data.get("units", [1])

    def _run():
        _register_task_thread()  # 记录线程 id，供"立即停止"注入异常
        try:
            set_running(f"知识过关")
            log_msg(f"启动知识过关: 单元{units}")
            sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
            from modules.知识过关 import run_module
            from common.tools import dismiss_global_popups, close_ad, ensure_grade
            import uiautomator2 as u2

            d = _connect_device()
            log_msg("设备已连接")

            # 关广告 + 确认年级
            for _ in range(3):
                dismiss_global_popups(d)
            close_ad(d)
            ok = ensure_grade(d, "五年级上册", "湘少版")
            log_msg("年级确认: 湘少版 五年级上册" if ok else "⚠ 年级切换失败，继续尝试")

            q = run_module(d)
            log_msg(f"✅ 知识过关完成: {q} 题")
            set_done()
        except SystemExit:
            log_msg("⏹ 任务已被立即停止", "warning")
            _cleanup_task_state()  # 统一收尾: 清running + 重置停止标志
        except Exception as e:
            log_msg(f"❌ 任务异常: {e}", "error")
            set_done()

    _KNOWLEDGE_RUNNER = threading.Thread(target=_run, daemon=True)
    _KNOWLEDGE_RUNNER.start()
    return jsonify({"status": "started", "units": units})


# ============================================================
# 语音评测 自动化（题目未做好，仅进入模块）
# ============================================================

_VOICE_RUNNER = None  # 后台线程


@app.route("/api/voice/run", methods=["POST"])
def api_voice_run():
    """启动语音评测模块（半自动：脚本驱动+用户朗读；继承主页版本/年级）
    ★ 2026-08-23：语音评测从主页继承当前版本年级（如湘少版二年级上册），
      点主页"语音评测"文字(877,1987)进入 → 目录 Unit 1-7 → 单元展开 → 部分作答
    """
    global _VOICE_RUNNER
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409

    # 请求参数：units（如 [1,2]）、grade、version（默认继承主页当前）
    _data = request.get_json(silent=True) or {}
    _units = _data.get("units") or None
    _grade = _data.get("grade", "") or ""
    _version = _data.get("version", "") or ""

    def _run():
        _register_task_thread()  # 记录线程 id，供"立即停止"注入异常
        try:
            set_running(f"语音评测")
            log_msg(f"启动语音评测（半自动：用户朗读+脚本驱动）")
            sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
            from modules.语音评测 import run_module
            from common.tools import dismiss_global_popups, close_ad
            import uiautomator2 as u2

            d = _connect_device()
            log_msg("设备已连接")

            for _ in range(3):
                dismiss_global_popups(d)
            close_ad(d)
            # ★ 语音评测内部处理版本/年级：主页当前版本+传入期望值兜底
            r = run_module(d, units=_units,
                           expected_grade=_grade, expected_version=_version)
            log_msg(f"✅ 语音评测完成: {r} 个单元")
            set_done()
        except SystemExit:
            log_msg("⏹ 任务已被立即停止", "warning")
            _cleanup_task_state()  # 统一收尾: 清running + 重置停止标志
        except Exception as e:
            log_msg(f"❌ 任务异常: {e}", "error")
            set_done()

    _VOICE_RUNNER = threading.Thread(target=_run, daemon=True)
    _VOICE_RUNNER.start()
    return jsonify({"status": "started"})


# ============================================================
# 巧记单词 自动化（单词同步闯关）
# ============================================================

_QIAOJI_RUNNER = None  # 后台线程


@app.route("/api/qiaoji/run", methods=["POST"])
def api_qiaoji_run():
    """启动巧记单词自动化（单词同步闯关）
    请求: {"units": [1]}  # 单元数，默认 U1-U9 全跑
    """
    global _QIAOJI_RUNNER
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409

    data = request.get_json() or {}
    units = data.get("units", list(range(1, 10)))  # 默认全部单元

    def _run():
        _register_task_thread()  # 记录线程 id，供"立即停止"注入异常
        try:
            set_running(f"巧记单词")
            log_msg(f"启动巧记单词: 单元{units}")
            sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
            from modules.巧记单词 import run_module
            from common.tools import dismiss_global_popups, close_ad, ensure_grade
            import uiautomator2 as u2

            d = _connect_device()
            log_msg("设备已连接")

            for _ in range(3):
                dismiss_global_popups(d)
            close_ad(d)
            ok = ensure_grade(d, "五年级上册", "湘少版")
            log_msg("年级确认: 湘少版 五年级上册" if ok else "⚠ 年级切换失败，继续尝试")

            q = run_module(d)
            log_msg(f"✅ 巧记单词完成: {q} 题")
            set_done()
        except SystemExit:
            log_msg("⏹ 任务已被立即停止", "warning")
            _cleanup_task_state()  # 统一收尾: 清running + 重置停止标志
        except Exception as e:
            log_msg(f"❌ 任务异常: {e}", "error")
            set_done()

    _QIAOJI_RUNNER = threading.Thread(target=_run, daemon=True)
    _QIAOJI_RUNNER.start()
    return jsonify({"status": "started", "units": units})


# ============================================================
# 前提设置：切换版本 / 年级
# ============================================================
import importlib
_setup_mod = None

def _get_setup():
    global _setup_mod
    if _setup_mod is None:
        sys.path.insert(0, str(Path(__file__).parent / "scripts"))
        _setup_mod = importlib.import_module("common.setup")
    return _setup_mod


@app.route("/api/setup", methods=["POST"])
def api_setup():
    """切换教材版本 + 年级（前提功能，模块运行前先设置）

    请求: {"version": "湘少版", "grade": "五年级上册"}
    """
    global _SETUP_RUNNER
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行，请等待"}), 409

    data = request.get_json() or {}
    version = (data.get("version") or "").strip()
    grade = (data.get("grade") or "").strip()
    if not version or not grade:
        return jsonify({"error": "请填写版本和年级"}), 400

    def _run():
        _register_task_thread()  # 记录线程 id，供"立即停止"注入异常
        try:
            log_msg(f"前提设置: 切换到 {version} {grade}")
            setup = _get_setup()
            import uiautomator2 as u2
            d = _connect_device()
            ok = setup.switch_version_grade(d, version, grade)
            if ok:
                log_msg(f"切换成功: 当前 {version} {grade}", "success")
            else:
                log_msg("切换失败，请重试", "error")
        except SystemExit:
            log_msg("⏹ 任务已被立即停止", "warning")
            _cleanup_task_state()  # 统一收尾: 清running + 重置停止标志
        except Exception as e:
            log_msg(f"切换异常: {e}", "error")

    _SETUP_RUNNER = threading.Thread(target=_run, daemon=True)
    _SETUP_RUNNER.start()
    return jsonify({"status": "started", "version": version, "grade": grade})


# ============================================================
# 多模块检测：版本/年级 + 多个模块依次执行
# ============================================================
_LAST_MODULES_RESULT = None
# ★ 当前多模块检测的 模块→脚本 映射（_detect_content_dimension 按模块匹配脚本用）
_GLOBAL_DOCX_MAP = {}


@app.route("/api/modules/stop", methods=["POST"])
def api_modules_stop():
    """立即停止正在运行的任务（设置停止标志 + 注入异常强制中断线程）"""
    forced, message = _request_stop()
    if not task_status["running"] and not forced:
        return jsonify({"status": "idle", "message": message})
    return jsonify({"status": "stopping", "message": message})


# ============================================================
# 脚本对照 LLM 知识性审查（有脚本的模块：基础自动化后追加）
# ============================================================

def _qreview_to_state(module: str, docx: str, unit, r, stage: str = "", q=None) -> dict:
    """把 ReviewAgent 的 QuestionReview 转成六维面板 questions 记录。"""
    # ★ 2026-08-25 错题日志定位：构造"口语训练 · U3 · 第2大题·第4小题"位置
    #   （口语题 script_question 含 big/stage_idx；听力专项无 big 用 stage+unit）
    #   缺 q 时回退到"Q{idx:02d}"，不会因为脚本格式差异导致空位
    _pos = ""
    if q is not None:
        _u = getattr(q, "unit", 0) or unit or 0
        _big = getattr(q, "big", 0) or 0
        _no = getattr(q, "stage_idx", 0) or 0
        if _u and _big and _no:
            _pos = f"U{_u}·第{_big}大题·第{_no}小题"
        elif _u and _big:
            _pos = f"U{_u}·第{_big}大题"
        elif _u and _no:
            _pos = f"U{_u}·第{_no}小题"
        elif _u:
            _pos = f"U{_u}"
    if not _pos:
        _pos = f"Q{r.idx:02d}"
    def _cr(chk):
        if chk is None:
            return None, "", "skip"
        det = "；".join((chk.details or [])[:4])
        # ★ skip = 该维度无法核对（听音/图片题无截图），计为「未检」而非「不通过」
        if getattr(chk, "method", "") == "skip":
            return None, det[:300], "skip"
        # ★ uncertain = LLM 无法解析判定（待审/需人工复核）→ 也计「未检」而非「不通过」
        #   用户实测："待审"被当成不通过 → 明明几维过却综合判定不通过
        if getattr(chk, "method", "") == "uncertain":
            return None, det[:300], "uncertain"
        return chk.passed, det[:300], chk.method or ""
    dims = {
        "stem": _cr(r.stem_check), "content": _cr(r.content_check),
        "image": _cr(r.image_check), "answer": _cr(r.answer_check),
        "audio": _cr(r.audio_check), "post_error": _cr(r.post_error_check),
    }
    checked = [v for v, _, _ in dims.values() if v is not None]
    passed = [v for v, _, _ in dims.values() if v is True]
    overall = None
    if checked:
        # ★ 修复：用 checked（非 None 的实际判定维度）判断，None=未检维度不计入
        #   （原写法遍历全部 dims 含 None，`None is True` 恒 False → 有任一未检维度时
        #    整体永远判 None，维度全通过也显示不出来）
        if any(v is False for v in checked):
            overall = False
        elif all(v is True for v in checked):
            overall = True
        # 部分检查（无 False 但未全过，如仅 1 维）→ 保持 None（未定），不误报不通过
    # ★ module_qno：脚本 idx（后续由 report_exporter 按 stage 分组重置显示）
    # ★ progress：含具体位置（口语训练·U3·第2大题·第4小题），方便错题日志快速定位
    return {
        "idx": r.idx, "total": 0,
        "question_type": r.question_type or "脚本题",
        "module": module,            # ★ 模块（分组展示用）
        "stage": stage or "",        # ★ 子模块阶段（基础巩固/综合进阶/难点突破）
        "module_qno": r.idx,        # 写脚本 idx；report_exporter 按 stage 重新排号
        "screenshot": "",
        "progress": f"{_pos}",      # ★ 口语训练 U3·第2大题·第4小题 / 听力专项 基础巩固 1.1
        "position": _pos,           # ★ 单独存一份，供报告导出/下载错题日志按位置排序
        "ai_stem": dims["stem"][0], "ai_content": dims["content"][0],
        "ai_image": dims["image"][0], "ai_answer": dims["answer"][0],
        "ai_audio": dims["audio"][0], "ai_post_error": dims["post_error"][0],
        "stem_reason": dims["stem"][1], "content_reason": dims["content"][1],
        "image_reason": dims["image"][1], "answer_reason": dims["answer"][1],
        "audio_reason": dims["audio"][1], "post_error_reason": dims["post_error"][1],
        "overall_passed": overall,
        "overall_score": round(len(passed) / max(len(checked), 1), 2) if checked else 0.0,
        # ★ 题干显示实际内容（脚本题 stem；口语训练朗读题为多句子）
        "stem": (r.stem_check.actual or "").strip() if getattr(r.stem_check, "actual", "") else
                f"第{r.idx}题（{r.question_type or '脚本题'}）",
        "options": "", "script_answer": r.script_answer or "",
        "note": f"LLM 知识性审查 · 脚本 {docx}",
    }


def _run_llm_script_review(module: str, docx: str, version: str, unit, stage: str = ""):
    """用 ReviewAgent 对照脚本做 LLM 知识性审查（文字核对 + 有截图时图片识别），
    结果写入 _inspection_state["questions"]（前端六维面板实时展示）。

    ★ stage 参数：听力专项区分"练习/测试"，避免脚本错配
      - stage="练习" → 只审查 stage=练习 的题
      - stage="测试" → 只审查 stage=测试 的题
      - stage=""     → 不区分（其他模块）
    """
    global _inspection_state
    docx_path = PROJECT_ROOT / "uploads" / docx
    if not docx_path.exists():
        log_msg(f"❌ 脚本文件不存在: {docx}（请在页面上传脚本）", "error")
        return
    # ★ 2026-08-26 修复：后端也校验【脚本文件名 vs 模块名】是否匹配！
    #   前端单模块手动脚本兜底曾直接绑 reviewDocx 下拉的脚本（不校验模块名）→
    #   用户测"巧记单词"却拿到"听力专项U6"脚本做 LLM 审查（模块和脚本毫无关系）。
    #   这里做最终防线：文件名不含本模块关键词 → 视为无匹配脚本，仅基础完整性。
    try:
        _mod_kws = {
            "听力专项": ("听力专项", "听力"),
            "口语训练": ("口语", "朗读"),
            "单元自检": ("单元自检", "单元检测"),
            "知识过关": ("知识过关", "知识", "重点词汇", "重点句型"),
            "巧记单词": ("巧记", "单词"),
            "语音评测": ("语音评测", "语音"),
        }.get(module, (module,))
        if not any(k in docx for k in _mod_kws):
            log_msg(f"ℹ {module}：脚本「{docx}」不是{module}模块的脚本，无匹配脚本 → 仅做基础完整性检查", "info")
            return
    except Exception:
        pass
    _stage_label = f"·{stage}" if stage else ""
    # ★ 审查体系改造：索引复用检查（已审查且脚本未变 → 跳过 LLM，复用结论）
    _version = _inspection_state.get("version", "未知版本")
    _grade = _inspection_state.get("grade", "未知年级")
    _unit_key = str(unit).strip() if unit not in (None, "", 0, "0", "NONE") else "全部"
    _index_key = f"{_version}/{_grade}/{module}/{_unit_key}"
    _index = load_script_review_index()
    _docx_hash = _script_file_hash(docx_path)
    _cached = _index.get(_index_key, {})
    if _cached.get("status") == "done" and _cached.get("script_hash") == _docx_hash:
        _reuse_msg = f"{module}（{_unit_key}）已于 {_cached.get('reviewed_at','?')} 审查过，脚本未变，本次复用结论"
        log_msg(f"ℹ {_reuse_msg}", "info")
        _inspection_state["_reuse_info"] = _reuse_msg
        # 从缓存结果文件加载历史审查结论，合并到 _inspection_state["questions"]
        _cache_path = PROJECT_ROOT / "data" / "script_review_cache" / f"{_index_key.replace('/','_').replace(' ','')}.json"
        try:
            import json
            if _cache_path.exists():
                _cached_qs = json.loads(_cache_path.read_text(encoding="utf-8"))
                if isinstance(_cached_qs, dict):
                    _inspection_state["questions"].update(_cached_qs)
                    log_msg(f"✅ 已加载 {len(_cached_qs)} 题历史审查结论", "success")
        except Exception as _e:
            log_msg(f"⚠ 加载缓存审查结论失败: {_e}，将重新审查", "warning")
            # 缓存加载失败 → 回退到完整审查，不 return
        else:
            return  # 复用成功，跳过 LLM 审查
    log_msg(f"🧠 {module}{_stage_label} 正在用脚本「{docx}」做 LLM 知识性审查…（逐题核对题干/选项/答案，耗时较长）", "info")
    # ★ 听力专项多子模块：临时切换 set_current_module 上下文，确保 questions 用对的 stage 标记
    if stage:
        try:
            from common.logger import set_current_module
            set_current_module(module, stage)
        except Exception:
            pass
    try:
        sys.path.insert(0, str(PROJECT_ROOT))
        from src.review_agent import ReviewAgent, ReviewConfig, normalize_units
        # ★ 2026-08-25 修复：支持多单元（"2,4" / "2-4"），不再截断成第一个单元
        #   用户场景：选"口语训练五上 U2、U4"，脚本"口语训练湘少五上(全单元)"
        #   → 应提取 U2+U4 的题目审查；旧逻辑 int(unit.split("-")[0]) 只审 U2
        _u = str(unit).strip() if unit not in (None, "", 0, "0", "NONE") else 0
        agent = ReviewAgent(ReviewConfig(
            docx_path=str(docx_path),
            screenshot_dir=str(PROJECT_ROOT / "screenshots"),
            unit=_u, verbose=False,
        ))
        # ★ 单元不匹配提示：脚本不含所选单元（用户实测"单元根本没脚本"，
        #   静默显示"0题"误导人）→ 明确日志：脚本实际覆盖哪些单元
        # ★ 2026-08-25 修复：部分覆盖也要提示！旧逻辑只看 script_questions 是否为空
        #   （全不匹配才提示），选 U4+U8 而脚本只有 U6-U10 时，U4 不在范围被静默忽略。
        if not agent.script_questions:
            _units_ok = sorted(u for u in getattr(agent, "script_units", set()) if u > 0)
            # ★ 2026-08-26 精简提示：用户期望直接说"无覆盖脚本"即可，
            #   不再罗列脚本名/期望单元/实际单元等细节（检查人员只需要知道"没得审"）
            log_msg(f"ℹ {module}：无覆盖所选单元的脚本 → 跳过 LLM 脚本审查，仅基础完整性"
                    f"（脚本实际单元: {_units_ok if _units_ok else '未解析到'}）", "info")
            return
        # ★ 部分覆盖提示：所选单元里只有部分被脚本覆盖（如选 4,8 脚本只有 8）
        _units_want2 = normalize_units(_u)
        _units_ok2 = set(getattr(agent, "script_units", set()))
        if _units_want2:
            _missing = sorted(u for u in _units_want2 if u not in _units_ok2)
            if _missing:
                # ★ 2026-08-26 精简提示：只说哪些单元没覆盖（去掉脚本名/细节）
                log_msg(f"ℹ {module}：单元 {_missing} 无覆盖脚本，仅做基础完整性检查", "info")
        # ★ 不传 screenshots：自动化过程只有答错截图，旧截图与本次题目对不上会误导 LLM 判定
        #   → 统一走 LLM 脚本审查（_review_script_llm：脚本信息 → LLM 六维判定，理由具体有说服力）
        results = agent.review(screenshots={})
        n_pass = sum(1 for rr in results if rr.overall_passed)
        for q, rr in zip(agent.script_questions, results):
            qid = f"{module}-脚本-Q{rr.idx:02d}"
            _inspection_state["questions"][qid] = _qreview_to_state(
                module, docx, unit, rr, stage=getattr(q, "stage", "") or "", q=q)
        _save_inspection_state_merge()
        # ★ 触发实时错题报告（前端「📑 查看报告」）
        try:
            _live_regen_error_report()
        except Exception:
            pass
        log_msg(f"✅ {module} LLM 知识性审查完成: {len(results)}题（通过{n_pass}，未检维度不计入不通过）", "success")
        # ★ 审查体系改造：审查完成后写入索引 + 缓存文件
        try:
            _index = load_script_review_index()
            _index[_index_key] = {
                "script": docx, "script_hash": _docx_hash,
                "reviewed_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "questions": len(results),
                "status": "done",
            }
            save_script_review_index(_index)
            # 缓存审查结论到独立文件（后续复用可直接加载）
            _cache_dir = PROJECT_ROOT / "data" / "script_review_cache"
            _cache_dir.mkdir(parents=True, exist_ok=True)
            _cache_path = _cache_dir / f"{_index_key.replace('/','_').replace(' ','')}.json"
            _cache_path.write_text(
                json.dumps({k: _inspection_state["questions"].get(k) for k in
                           [f"{module}-脚本-Q{rr.idx:02d}" for rr in results]
                           if k in _inspection_state["questions"]},
                          ensure_ascii=False, indent=2),
                encoding="utf-8")
        except Exception as _e:
            pass
    except Exception as e:
        log_msg(f"❌ {module} LLM 知识性审查失败: {e}", "error")


@app.route("/api/modules/run", methods=["POST"])
def api_modules_run():
    """多模块检测：切换版本/年级后依次执行多个模块

    请求: {"version": "湘少版", "grade": "五年级上册", "modules": ["听力专项", "口语训练"],
           "units": {"听力专项": "1-3", "单元自检": "1-5"}}   # units 可选，指定各模块单元范围
    返回: {"status": "started", "modules": [...]}
    """
    global _MODULES_RUNNER
    if task_status["running"]:
        return jsonify({"error": "已有任务在运行"}), 409

    # 清空审查结果区（避免上次任务的题目残留）
    # ★ 2026-08-25 修复：必须同步清空磁盘文件 data/inspection_state.json！
    #   旧逻辑只清内存 _inspection_state["questions"]，磁盘仍保留上次运行的题
    #   （如五年级口语训练20题+4不通过），而 /api/errors/summary 读的是磁盘 →
    #   用户跑"三年级未上线0题"却显示"不通过4/20"（日志0题/前端旧数据矛盾）。
    #   快速检查入口(1247行)就是先清内存再 _save_inspection_state() 写盘，此处对齐。
    _inspection_state["questions"] = {}
    _inspection_state["current_question_idx"] = 0
    _save_inspection_state()

    data = request.get_json() or {}
    version = (data.get("version") or "湘少版").strip()
    grade = (data.get("grade") or "五年级上册").strip()
    unit_from = int(data.get("unit_from") or 0)
    unit_to = int(data.get("unit_to") or 0)
    modules = data.get("modules") or []
    units = data.get("units") or {}  # 可选：{模块名: 单元范围}
    docx_map = data.get("docxMap") or data.get("docx_map") or {}  # 可选：{模块名: 上传脚本文件名}
    if not isinstance(docx_map, dict):
        docx_map = {}

    # ★ 兼容两种 modules 格式：
    #   旧: ["听力专项", "口语训练"]
    #   新: [{"name": "听力专项", "units": "1-5"}, {"name": "口语训练", "units": ""}]
    module_names = []
    for m in modules:
        if isinstance(m, dict):
            name = (m.get("name") or "").strip()
            m_units = (m.get("units") or "").strip()
            if name:
                module_names.append(name)
                if m_units:
                    units.setdefault(name, m_units)
        else:
            module_names.append(str(m).strip())
    module_names = [n for n in module_names if n]
    if not module_names:
        return jsonify({"error": "请至少选择一个模块"}), 400

    # ★ 重置停止标志（新任务开始，清除上次的停止请求）
    global _STOP_REQUESTED
    with _STOP_LOCK:
        _STOP_REQUESTED = False

    # ★ 设备就绪检查：没连设备直接报错返回，不启动线程（避免 u2.connect 挂起卡住）
    try:
        sys.path.insert(0, str(Path(__file__).parent / "scripts"))
        from common.device import device_ok as _dev_ok
        cur_serial = os.environ.get("ANDROID_SERIAL") or ""
        if not _dev_ok(cur_serial or None):
            log_msg(f"❌ 设备未连接，无法开始检测（当前设备: {cur_serial or '未选择'}）", "error")
            return jsonify({"error": f"设备未连接（当前设备: {cur_serial or '未选择'}），请先连接设备"}), 400
        log_msg(f"✅ 设备就绪: {cur_serial or '(默认)'}", "success")
    except SystemExit:
        log_msg("⏹ 任务已被立即停止", "warning")
    except Exception as e:
        log_msg(f"❌ 设备检查失败: {e}", "error")
        return jsonify({"error": f"设备检查失败: {e}"}), 500

    def _run():
        _register_task_thread()  # 记录线程 id，供"立即停止"注入异常
        try:
            set_running("多模块检测")
            # ★ 记录版本/单元到巡检状态（实时错题报告标题/路径需要）
            _inspection_state["version"] = version
            _inspection_state["grade"] = grade
            # ★ 环境变量同步（模块代码读 YYB_VERSION/YYB_GRADE 用于脚本命名等，
            #   避免 sys.modules 不可靠）
            os.environ["YYB_VERSION"] = str(version)
            os.environ["YYB_GRADE"] = str(grade)
            if units:
                _first_unit = str(next(iter(units.values()), "?"))
                _inspection_state["unit"] = _first_unit
            units_desc = f"（单元: {units}）" if units else ""
            log_msg(f"多模块检测启动: {version} {grade} → {'、'.join(module_names)}{units_desc}")
            import importlib
            sys.path.insert(0, str(Path(__file__).parent / "scripts"))
            sched = importlib.import_module("scheduler")
            import uiautomator2 as u2
            # 线程内再次确认设备（防止入口检查后设备掉线）
            try:
                from common.device import device_ok as _dev_ok2
                if not _dev_ok2():
                    raise RuntimeError("设备已断开")
            except RuntimeError:
                raise
            except Exception:
                pass
            d = _connect_device()

            # ★ 逐模块分流：有匹配脚本 → 基础自动化 + 追加 LLM 知识性审查；
            #              无脚本 → 仅基础完整性检查
            # ★ 一次传全部模块给 run_all（内部统一循环）：只重启 App 一次，
            #   模块间 _back_to_home 切主页继续下一模块；LLM 审查经 on_module_done 回调
            global _GLOBAL_DOCX_MAP
            _GLOBAL_DOCX_MAP = docx_map or {}
            def _on_module_done(_mod, _res):
                _docx = (docx_map or {}).get(_mod, "")
                if not _docx:
                    log_msg(f"🔍 {_mod} 无匹配脚本 → 仅基础完整性检查", "info")
                    return
                _stage = (_res or {}).get("stage", "")
                log_msg(f"📖 {_mod}{'·' + _stage if _stage else ''} 有匹配脚本「{_docx}」→ 自动化后追加 LLM 知识性审查", "info")
                # ★ 修复模块间卡顿：LLM 审查改异步线程，不阻塞 run_all 的模块循环
                #   （原同步调用：听力专项完成后在主页面卡几分钟等逐题 LLM 审查，
                #     审查完才继续口语训练 → 用户实测"主页卡好久"）
                # ★ unit 解析：听力专项·测试单元在 units["听力专项_测试"]；练习单元在 units["听力专项"]
                #   ★ 2026-08-25 修复：传完整单元范围（"2,4"/"2-4"），不再截断成第一个单元
                #     （用户选 U2、U4 时旧逻辑只审 U2，U4 被静默丢弃）
                #   ★ 防御："NONE"哨兵（前端未勾选）→ 不是数字，忽略
                _u = 0
                if _stage == "测试":
                    _u_raw = units.get("听力专项_测试", "") or units.get(_mod, "")
                else:
                    _u_raw = units.get(_mod, "") or 0
                try:
                    _u_str = str(_u_raw).strip()
                    if _u_str and _u_str.upper() != "NONE" and _u_str[0].isdigit():
                        _u = _u_str            # ★ 完整范围（"2,4" / "2-4"），归一化交给 review_agent
                    else:
                        _u = 0
                except Exception:
                    _u = 0
                try:
                    _t = threading.Thread(
                        target=_run_llm_script_review,
                        args=(_mod, _docx, version, _u, _stage),
                        daemon=True,
                    )
                    _t.start()
                except Exception as _e:
                    log_msg(f"❌ {_mod} LLM 知识性审查启动失败: {_e}", "error")

            results = sched.run_all(
                module_names, d, version=version, grade=grade,
                units=units, stop_check=_is_stop_requested,
                on_module_done=_on_module_done,
            ) or {}
            # 保存结果供前端查询
            global _LAST_MODULES_RESULT
            _LAST_MODULES_RESULT = {
                "done": True,
                "version": version,
                "grade": grade,
                "results": results,
            }
            # 逐模块汇总日志
            for name, r in results.items():
                status = "成功" if r.get("ok") else "失败"
                lv = "success" if r.get("ok") else "error"
                log_msg(f"  [{status}] {name}: {r['q']}题 {r['t']}s", lv)
            if _is_stop_requested():
                log_msg("⏹ 任务已被手动停止", "warning")
            else:
                log_msg(f"多模块检测完成: {len(results)} 个模块", "success")
        except SystemExit:
            # 前端"立即停止"注入的异常：记录并正常收尾
            log_msg("⏹ 任务已被立即停止", "warning")
            _LAST_MODULES_RESULT = {
                "done": True,
                "version": version,
                "grade": grade,
                "stopped": True,
                "error": "任务已被手动停止",
                "results": {},
            }
        except Exception as e:
            log_msg(f"多模块检测异常: {e}", "error")
            _LAST_MODULES_RESULT = {
                "done": True,
                "version": version,
                "grade": grade,
                "error": str(e),
                "results": {},
            }
        finally:
            # 清理停止标志，供下次任务使用
            global _STOP_REQUESTED
            with _STOP_LOCK:
                _STOP_REQUESTED = False
            set_done()

    # ★ 修复：新任务启动前清空上次结果，否则前端轮询 /api/modules/result 会
    #   立刻拿到旧任务的 done=true 结果（表现为"2秒全部完成+旧题数"假象），
    #   把当前真正在跑的任务结果盖掉。
    global _LAST_MODULES_RESULT
    _LAST_MODULES_RESULT = None

    _MODULES_RUNNER = threading.Thread(target=_run, daemon=True)
    _MODULES_RUNNER.start()
    return jsonify({"status": "started", "version": version, "grade": grade, "modules": module_names})


@app.route("/api/modules/result", methods=["GET"])
def api_modules_result():
    """获取最近一次多模块检测的汇总结果"""
    return jsonify(_LAST_MODULES_RESULT or {"done": False})


# ============================================================
# 题目解析脚本：遍历检测时自动生成（单元答完 → LLM 知识点 → 写 docx）
# 前端只「列出 + 下载」，不手动触发生成（避免 LLM 重复做无用功）
# ============================================================
@app.route("/api/script/list", methods=["GET"])
def api_script_list():
    """列出已生成的题目解析脚本（gen_scripts/ + outputs/reports/ 下的脚本 docx）"""
    files = []
    for _d in (PROJECT_ROOT / "gen_scripts", PROJECT_ROOT / "outputs" / "reports"):
        if _d.exists():
            for f in sorted(_d.glob("*.docx"), reverse=True):
                # 只收解析脚本（新规范名：日期+版本+年级+模块；旧名含"生成脚本"）
                if "生成脚本" not in f.name and not re.match(r"^\d{6}.{0,20}(单元自检|听力专项|口语训练|知识过关|巧记单词|语音评测)", f.name):
                    continue
                files.append({
                    "name": f.name,
                    "path": str(f),
                    "size": f.stat().st_size,
                    "mtime": datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M"),
                    "url": f"/api/export/download/{f.name}",
                })
    return jsonify({"success": True, "files": files})


# ============================================================
# 启动
# ============================================================

if __name__ == "__main__":
    print("=" * 50)
    print("  英语宝模块检测 - Web 控制面板")
    print("=" * 50)
    print(f"  项目路径: {PROJECT_ROOT}")
    print(f"  启动: http://localhost:5000")
    print("=" * 50)
        # 自动检测屏幕分辨率并缩放坐标
    config = load_config()
    detect_screen_resolution(config.device.serial)
    scale_all_coords()

    # ★ 常驻实时截图（前端「手机画面」随时可见，不再只在任务运行时才有）
    _start_live_screenshot_daemon()

    # ★ 2026-08-25：支持 PORT 环境变量（默认 5000；多实例验证/端口冲突时可换端口）
    _port = int(os.environ.get("PORT", "5000"))
    app.run(host="0.0.0.0", port=_port, debug=False, threaded=True)
