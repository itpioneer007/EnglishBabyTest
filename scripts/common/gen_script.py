"""题目解析脚本生成器 — 把答题过程中收集的题目（题干+选项+答案）汇总成脚本 docx

用法（在模块答题循环中）:
    from common.gen_script import QuestionCollector
    coll = QuestionCollector(module="单元自检", version="湘少版", grade="五年级上册")
    # 每题作答时:
    coll.add(qno=q, stem=..., options=[...], answer=opt, qtype=..., unit=6)
    # 单元答完:
    coll.finish_unit(unit=6)   # 写入 docx 脚本文件

生成格式（参考"260707新湘鲁六上听力专项(试编).docx"）:
    1. 题干
    A. xx  B. xx  C. xx
    答案：X
    一级题型：xxx
    二级题型：xxx
    难度：基础
    关键词：模块#版本U单元#...
    知识点：...
    技能：...
    认知目标：识记
    适用年级：xxx
    编写年份：2026年
"""
import os
import re
import sys
import time
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False

_KNOWLEDGE_KWS = ("语法", "句型", "词汇", "单词", "固定搭配", "时态", "发音",
                  "am/is/are", "can", "like", "want", "what", "where", "who",
                  "When", "How", "Do ", "Does ", "Is ", "Are ")

# ★★★ 固定解析提示词模板（用户要求：规范解析方式与准度，不能乱讲，符合年级要求）
#   解析必须满足：考点准确 + 理由对应正确答案 + 年级适配 + 精简有说服力
#   任何解析生成（前端/后端/批量）都走这个模板，保证一致性。
KNOWLEDGE_PROMPT_TEMPLATE = (
    "你是{g_grade}（{g_version}）英语老师，正在给学生讲解一道错题的考点。\n"
    "题目：{g_stem}\n"
    "选项：{g_options}\n"
    "正确答案：{g_answer}\n\n"
    "请给出本题的知识点解析，必须遵守以下规范：\n"
    "1. 【准确】考点必须是该题真正的语法/词汇/句型点，结合选项和正确答案分析，"
    "不得臆造考点，不得泛泛而谈。\n"
    "2. 【年级适配】解释方式必须符合{g_grade}学生的认知水平"
    "（如be动词、一般现在时、名词复数这类该年级正在学的知识），"
    "不能讲超纲概念（如虚拟语气、定语从句）。\n"
    "3. 【对应答案】必须明确说明为什么{g_answer}是对的，其他选项为什么错/不同类。\n"
    "4. 【精简】30-70字，直接说结论，不要铺垫、不要'首先/其次'。\n"
    "5. 【格式】用中文解释，英文单词保留原文；只输出解析本身，不要标题、编号、前缀。\n"
    "6. 【禁止】禁止'涉及XX相关知识点''本题考查XX'这类模糊套话，禁止编造教材没有的内容。"
)


def _extract_ui_question(xml, answer="", qtype_hint=""):
    """从答题页 XML 提取一道题的（题干/选项/题型），供脚本生成收集。
    - 题干：分级提取——题干节点(question_title_tv/tv_caption/tv_title/tv_stem等)优先，
      兜底长文本（排除按钮/进度/计时器/选项噪音）；多段题干合并（指令+内容）
    - 选项：字母开头项（A. xx）或选项容器（option_cb 等）或纯字母（T/F）
    - 纯听音无题干（题干为空/仅"听录音"指令）→ stem 返回空，上层跳过
    """
    if not xml:
        return {"stem": "", "options": [], "qtype": qtype_hint}

    noise = ("下一题", "上一题", "检查", "提交", "查看报告", "练习报告",
             "恭喜", "回答正确", "回答错误", "很遗憾", "练习结束还剩", "还剩",
             "得分", "用时", "点击录音", "点击结束", "继续答题", "退出训练",
             "答题卡", "单元自检", "口语训练", "听力专项", "知识过关",
             "巧记单词", "语音评测", "全脑记词", "单词听写", "查看解析",
             "正确答案", "错误答案", "本大题", "本小题", "分值", "满分",
             "考前突破", "当前版本", "练习记录", "训练规则说明", "好的",
             "播放", "暂停", "重听", "上一页", "下一页")

    def _is_noise(t):
        if t in noise or any(n in t for n in noise):
            return True
        if re.match(r"^\d{1,2}:\d{2}$|^\d+(\.\d+)?%?$|^\d+\s*/\s*\d+$|^\d+分$", t):
            return True
        if re.match(r"^[A-F][.、．]\s*", t) or re.match(r"^[TF]\s*$", t):
            return True
        return False

    def _node_parts(_block):
        """从 node 块提取 (text, resource_id)"""
        _mt = re.search(r'text="([^"]*)"', _block)
        _mr = re.search(r'resource-id="([^"]*)"', _block)
        return (_mt.group(1).strip() if _mt else "",
                _mr.group(1) if _mr else "")

    # ① 题干：优先题干节点（resource-id 含题干语义）——这些一定是题干文字
    stem_nodes = []
    for _nd in re.finditer(r'<node\b[^>]*>', xml):
        _block = _nd.group(0)
        t, rid = _node_parts(_block)
        if not t or len(t) < 2 or len(t) >= 100:
            continue
        if rid and re.search(r'(question_title|tv_caption|tv_title|tv_stem|tv_question|question_content|title_tv)', rid):
            stem_nodes.append(t)
    # ② 兜底：非选项/非噪音长文本（含短题干 ≥2 字符）
    if not stem_nodes:
        for _nd in re.finditer(r'<node\b[^>]*>', xml):
            _block = _nd.group(0)
            t, rid = _node_parts(_block)
            if not t or len(t) < 2 or len(t) >= 100:
                continue
            # 排除选项容器/按钮/交互节点
            if rid and re.search(r'(option_cb|option_iv|radio_btn|check_box|tv_option|btn_|tv_progress|tv_time|tv_index|iv_|img_)', rid):
                continue
            if _is_noise(t):
                continue
            stem_nodes.append(t)
    # ③ 合并：去重 + 去包含（保留完整），最多 3 段
    _merged = []
    for t in stem_nodes:
        t = re.sub(r"[（(]\s*共\d+分\s*[）)]", "", t).strip()  # 清洗分值
        if not t:
            continue
        if any(t == s or t in s or s in t for s in _merged):
            continue
        _merged.append(t)
        if len(_merged) >= 3:
            break
    stem = " ".join(_merged).strip()
    # ★ 纯听音指令判定：题干仅为"听录音/听音选X"等指令性文字 → 视为无题干（跳过）
    if stem.strip() and re.fullmatch(
        r"(听录音|听音|听一听|Listen|听录音[，,。 ]?[选判勾][出择].*|"
        r"听对话|听短文|听句子|听单词|听材料|听问题|请听).*",
        stem.strip(), re.IGNORECASE):
        stem = ""
    # ② 选项：字母开头项（A. xx）优先
    opts = []
    for m in re.finditer(r'text="([A-E][.、．]\s*[^"]{1,60})"', xml):
        t = m.group(1).strip()
        if t and t not in opts:
            opts.append(t)
    if not opts:
        # ★ 选项是 CheckBox/RadioButton（option_cb 等，文本无字母前缀，如 parrot/notebook）
        #   按出现顺序收集（单元自检等模块的选项就是这种独立单词）
        for m in re.finditer(
                r'<node[^>]*resource-id="[^"]*(?:option_cb|option_iv|radio_btn|check_box)[^"]*"[^>]*text="([^"]{1,40})"|'
                r'<node[^>]*text="([^"]{1,40})"[^>]*resource-id="[^"]*(?:option_cb|option_iv|radio_btn|check_box)[^"]*"', xml):
            t = (m.group(1) or m.group(2) or "").strip()
            if t and t not in opts:
                opts.append(t)
    if not opts:
        # 纯字母（T/F/A-E 判断题）
        for m in re.finditer(r'text="([TFABCDE])"', xml):
            t = m.group(1)
            if t not in opts:
                opts.append(t + ". ")
    return {"stem": stem.strip(), "options": opts[:6], "qtype": qtype_hint}


class QuestionCollector:
    """答题过程中的题目收集器（每道有题干的题收集一条，单元答完汇总写脚本）"""

    def __init__(self, module="单元自检", version="", grade="", save_root=None):
        self.module = module
        self.version = version or "湘少版"
        self.grade = grade or "五年级上册"
        # 保存目录：项目根/gen_scripts（默认，scripts/common/gen_script.py 上三级）；可传 save_root 覆盖
        self.save_root = save_root or os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "gen_scripts")
        os.makedirs(self.save_root, exist_ok=True)
        self.questions = []       # 当前单元收集的题（含题干可解析的）
        self.skipped_no_stem = 0  # 无题干跳过的题数（纯听音等）

    # ------------------------------------------------------------
    def add(self, qno, stem="", options=None, answer="", qtype="",
            unit=0, recording="", knowledge="", big=None):
        """收集一道题。仅当有题干（非纯听音）才保留，否则跳过。

        Args:
            qno: 题号（单元内序号）
            stem: 题干文字（空 = 纯听音/无题干 → 跳过）
            options: 选项列表 ['A. xx', 'B. xx', ...]
            answer: 答案（点击的选项，如 'B'）
            qtype: 题型（如 '单选', '判断'）
            unit: 单元号
            recording: 录音原文（听力题可能有）
            knowledge: 知识点（可选，LLM 生成或留空）
            big: 大题号（口语训练等"第N大题"定位用）
        """
        stem = (stem or "").strip()
        if not stem or stem in ("(无题干文字)", "(无)", "听录音", "听录音，选择正确答案"):
            self.skipped_no_stem += 1
            return
        # ★ 听音题判定（用户明确要求）：题干含"听"字即跳过——
        #   "听词汇/听句子/听对话/听录音" 全是听力题，题目实质在音频里，
        #   脚本里没有可解析的题干文字，跳过不做脚本审查
        #   （保留含"听"的非听音词：如"听说读写"等极少数情况由下方白名单保护）
        _LISTEN_WHITELIST = ("听说读写", "听力理解能力", "听说")
        if "听" in stem and not any(w in stem for w in _LISTEN_WHITELIST):
            self.skipped_no_stem += 1
            return
        # 兜底：英语 Listen/listen 开头（"Listen and choose"等）
        if re.match(r"^\s*(Listen|listen|Listening)\b", stem):
            self.skipped_no_stem += 1
            return
        opts = [str(o).strip() for o in (options or []) if str(o).strip()]
        # 清洗选项：去掉多余空白（"A.  sport" → "A. sport"）
        opts = [re.sub(r"\s+", " ", o) for o in opts]
        self.questions.append({
            "qno": int(qno) if str(qno).isdigit() else qno,
            "stem": stem, "options": opts, "answer": str(answer).strip(),
            "qtype": qtype or "", "unit": int(unit) if str(unit).isdigit() else unit,
            "recording": (recording or "").strip(),
            "knowledge": (knowledge or "").strip(),
            "big": big,
            "time": datetime.now().strftime("%H:%M:%S"),
        })

    # ------------------------------------------------------------
    def _llm_knowledge(self, q):
        """用大模型生成知识点解析：符合该版本年级的知识水平，精简有说服力。
        像"为什么1+1=2"要用小学方法解释——不模糊、直接给出该年级该懂的结论。
        """
        try:
            sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
            from src.reviewer_common import LLMClient
            llm = LLMClient.from_config()
            opt_str = " ".join(q["options"]) if q["options"] else "（无文字选项）"
            # ★ 固定模板（KNOWLEDGE_PROMPT_TEMPLATE），保证解析规范一致
            prompt = KNOWLEDGE_PROMPT_TEMPLATE.format(
                g_grade=self.grade, g_version=self.version,
                g_stem=q["stem"], g_options=opt_str, g_answer=q["answer"])
            ans = ""
            for _try in range(2):  # 不合格重试一次
                ans = (llm.ask(prompt) or "").strip()
                if ans and len(ans) > 5 and "LLM 调用失败" not in ans and self._check_knowledge_ok(ans):
                    break
                ans = ""  # 不合格 → 重试
            if ans:
                return ans[:120]
        except Exception as _e:
            print(f"  ⚠ 知识点 LLM 生成失败: {_e}")
        return self._auto_knowledge(q)

    # ------------------------------------------------------------
    @staticmethod
    def _check_knowledge_ok(ans):
        """解析质量校验：模糊套话/超纲词 → 不合格（防'乱讲'）"""
        fuzzy = ("涉及", "相关知识点", "本题考查", "本道题考查", "一般考查", "基础知识点")
        if any(f in ans for f in fuzzy):
            return False
        # 超纲概念（五年级不应出现）
        over = ("虚拟语气", "定语从句", "现在完成时", "过去完成时", "被动语态", "非谓语")
        if any(o in ans for o in over):
            return False
        return True

    # ------------------------------------------------------------
    def finish_unit(self, unit):
        """单元答完：把收集到的题写成脚本 docx（模板格式）。
        返回生成的文件路径；无题可写返回 None。
        """
        if not self.questions:
            print(f"  ⚠ 单元 {unit} 无可生成解析的题目（{self.skipped_no_stem} 题无题干跳过）")
            return None
        if not _HAS_DOCX:
            print("  ❌ python-docx 不可用，无法生成脚本")
            return None
        # ★ 知识点用大模型生成（每题一次调用，串行；该年级知识点有说服力）
        print(f"  🧠 正在为 {len(self.questions)} 题生成知识点解析（LLM）…")
        for q in self.questions:
            if not q.get("knowledge"):
                q["knowledge"] = self._llm_knowledge(q)
        _u = f"U{unit}"
        # ★ 规范文件名：日期+版本缩写+年级缩写+模块[+单元]，如
        #   "260817湘少五上单元自检U6.docx"（参考官方脚本"260714新湘鲁六上听力专项"格式）
        fname = f"{datetime.now().strftime('%y%m%d')}{self._short_version()}{self._short_grade()}{self.module}{_u}.docx"
        path = os.path.join(self.save_root, fname)
        doc = Document()
        # ★ 统一字体：不用内置 Heading/Title 样式（会随主题变蓝色Calibri→
        #   与正文字体颜色/大小/粗细不一致，用户反馈"同一个标题里不一致"）
        #   全部用 Normal + 手动设置：中文黑体、英文Arial、黑色、固定大小
        _normal = doc.styles["Normal"]
        _normal.font.name = "Arial"
        _normal.font.size = Pt(11)
        _normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        _normal.font.bold = False
        try:
            _normal.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        except Exception:
            pass

        def _add_line(text="", bold=False, size=11, space_after=4):
            """统一行：Normal 样式 + 指定加粗/字号（全文档一致，黑体Arial）"""
            _p = doc.add_paragraph()
            _p.paragraph_format.space_after = Pt(space_after)
            _r = _p.add_run(text)
            _r.font.name = "Arial"
            try:
                _r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            except Exception:
                pass
            _r.font.size = Pt(size)
            _r.font.bold = bold
            _r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            return _p

        # 文档标题（大号黑色加粗，不用 Title 样式）
        _add_line(f"{self.module} · {self.grade} {_u} · 题目解析脚本", bold=True, size=16, space_after=6)
        _add_line(f"版本：{self.version}　|　单元：{_u}　|　生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", size=10, space_after=2)
        _add_line(f"共 {len(self.questions)} 题（无题干/纯听音 {self.skipped_no_stem} 题跳过）", size=10, space_after=10)
        for i, q in enumerate(self.questions, 1):
            # 题号（标题，简洁）
            _add_line(f"第{i}题", bold=True, size=13, space_after=2)
            # 题干（完整内容：指令+具体内容，如"请找出与下面选项不符的一项"）
            _add_line(f"题干：{q['stem']}", size=11, space_after=2)
            _add_line(f"位置：{self._loc_str(q)}", size=11, space_after=2)
            _add_line(f"选项：{'　　'.join(q['options']) if q['options'] else '（无文字选项）'}", size=11, space_after=2)
            _add_line(f"正确答案：{q['answer']}", size=11, space_after=2)
            _add_line(f"知识点：{q['knowledge'] or self._auto_knowledge(q)}", size=11, space_after=10)
        doc.save(path)
        print(f"  ✅ 生成脚本: {path}（{len(self.questions)} 题，跳过无题干 {self.skipped_no_stem}）")
        return path

    # ------------------------------------------------------------
    def _loc_str(self, q):
        """App 位置描述：模块 · 单元 · 第N题（含大题时带大题号）"""
        _parts = [self.module]
        if q.get("unit"):
            _parts.append(f"U{q['unit']}")
        if q.get("big"):
            _parts.append(f"第{q['big']}大题")
        if q.get("qno"):
            _parts.append(f"第{q['qno']}题")
        return " · ".join(_parts)

    # ------------------------------------------------------------
    def _short_version(self):
        """版本缩写：湘少版→湘少；新湘鲁版→新湘鲁；其余原样（去'版'字）"""
        v = (self.version or "").strip()
        v = v.replace("（", "").replace("）", "").replace("(", "").replace(")", "")
        for k in ("新湘鲁", "湘鲁", "湘少", "人教", "外研", "译林"):
            if k in v:
                return k
        return v.replace("版", "") or "通用"

    def _short_grade(self):
        """年级缩写：五年级上册→五上；六年级下册→六下；其余压缩"""
        g = (self.grade or "").strip()
        m = re.match(r"([一二三四五六七八九])\s*年级\s*(上|下)\s*册", g)
        if m:
            return f"{m.group(1)}{m.group(2)}"
        # 兜底：去掉"年级/册"字
        return g.replace("年级", "").replace("册", "") or ""

    # ------------------------------------------------------------
    @staticmethod
    def _auto_knowledge(q):
        """从题干/答案自动推断知识点（简单启发式；复杂可后续接 LLM）"""
        txt = f"{q['stem']} {q['answer']} {' '.join(q['options'])}"
        for kw in _KNOWLEDGE_KWS:
            if kw in txt:
                return f"涉及{kw}相关知识点"
        return "基础知识点"

    # ------------------------------------------------------------
    def summary(self, unit):
        return {
            "unit": unit,
            "generated": len(self.questions),
            "skipped": self.skipped_no_stem,
        }
